# -*- coding: utf-8 -*-
"""
agent.py
========
统一林业智能分析核心入口 —— 并行感知 Neuro-Symbolic ReAct Agent 

"""

import os
import sys
import io
import json
import time
import yaml
from concurrent.futures import ThreadPoolExecutor, as_completed
from dotenv import load_dotenv
from neo4j import GraphDatabase
from provider import chat_with_tools

# 修复 Windows 控制台中文输出编码
try:
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stdout, 'buffer') and not isinstance(sys.stdout, io.TextIOWrapper):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)
    if hasattr(sys.stderr, 'buffer') and not isinstance(sys.stderr, io.TextIOWrapper):
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
except Exception:
    pass

from forestry_spatial_tools import (
    tool_calc_stand_structure_metrics,
    tool_calc_tree_morphology_metrics,
    tool_calc_species_diversity_metrics,
    tool_calc_volume_metrics,
    tool_calc_deadwood_metrics,
    tool_calc_shrub_metrics,
    tool_calc_hegyi_competition,
    tool_calc_carbon_and_hydrology_metrics,
    tool_assess_climate_and_biosecurity_risks,
    tool_scan_subplots_risk_summary,
    FORESTRY_SPATIAL_SCHEMAS,
)
try:
    from forestry_visualization_engine import (
        plot_subplot_grid_heatmap,
        plot_size_class_distribution,
        plot_species_composition,
        plot_tree_relationship_scatter,
        plot_group_comparison_boxplot,
        plot_tree_spatial_map,
        plot_subplot_percentile_profile,
        plot_climate_time_series,
    )
    _FORESTRY_VIS_ENGINE_AVAILABLE = True
    _FORESTRY_VIS_ENGINE_IMPORT_ERROR = None
except Exception as e:
    plot_subplot_grid_heatmap = None
    plot_size_class_distribution = None
    plot_species_composition = None
    plot_tree_relationship_scatter = None
    plot_group_comparison_boxplot = None
    plot_tree_spatial_map = None
    plot_subplot_percentile_profile = None
    plot_climate_time_series = None
    _FORESTRY_VIS_ENGINE_AVAILABLE = False
    _FORESTRY_VIS_ENGINE_IMPORT_ERROR = str(e)
from formula_execution_engine import NeuroSymbolicFormulaEngine


def _visualization_engine_error_response(tool_name: str) -> str:
    return json.dumps({
        "error": (
            f"可视化引擎不可用：{tool_name} 无法执行。"
            f" forestry_visualization_engine 模块导入失败，原因：{_FORESTRY_VIS_ENGINE_IMPORT_ERROR}。"
            " 请安装 plotly、kaleido 等依赖后重试。"
        )
    }, ensure_ascii=False)

# 1. 兼容性加载环境变量 (基于当前文件所在目录，摆脱硬编码)
base_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(base_dir, ".env"))

URI = os.getenv("NEO4J_URI", "bolt://localhost:7687")
USER_DB = os.getenv("NEO4J_USER", "neo4j")
PASSWORD = os.getenv("NEO4J_PASSWORD", "1820401753")

_driver = GraphDatabase.driver(URI, auth=(USER_DB, PASSWORD))
_ns_engine = NeuroSymbolicFormulaEngine()


# ==============================================================================
# 工具 1: 动态本体契约检索器 (增加灵活性，规避硬编码)
# ==============================================================================
def tool_inspect_ontology_schema(query_concept: str) -> str:
    """
    动态概念本体与声明式规则规范检索器。
    查询本地《qilian_ontology.yaml》与《forestry_knowledge_registry.yaml》，
    精准提取类、属性、SQLite 字段映射以及神经符号规则契约切片，避免智能体发生属性猜测或标签幻觉。
    """
    concept = str(query_concept).strip().lower()
    results = {"query_concept": query_concept, "matched_schema_slices": []}
    
    # 1. 检索类定义与关系契约 (qilian_ontology.yaml)
    ont_path = os.path.join(os.path.dirname(__file__), "ontology", "qilian_ontology.yaml")
    if os.path.exists(ont_path):
        try:
            with open(ont_path, "r", encoding="utf-8") as f:
                ont_data = yaml.safe_load(f) or {}
            for cls in ont_data.get("classes", []):
                cls_id = str(cls.get("class_id", "")).lower()
                cls_lbl = str(cls.get("label", "")).lower()
                cls_desc = str(cls.get("description", "")).lower()
                if concept in cls_id or concept in cls_lbl or concept in cls_desc:
                    results["matched_schema_slices"].append({"type": "OntologyClass", "schema": cls})
            for field in ont_data.get("canonical_tree_observation_fields", []):
                if concept in str(field.get("target", "")).lower() or any(concept in str(ex).lower() for ex in field.get("source_examples", [])):
                    results["matched_schema_slices"].append({"type": "SQLiteColumnMapping", "mapping": field})
        except Exception as e:
            results["ontology_error"] = str(e)

    # 2. 检索声明式计算规则与公式契约库 (forestry_knowledge_registry.yaml)
    reg_path = os.path.join(os.path.dirname(__file__), "ontology", "forestry_knowledge_registry.yaml")
    if os.path.exists(reg_path):
        try:
            with open(reg_path, "r", encoding="utf-8") as f:
                reg_data = yaml.safe_load(f) or {}
            for key in ["variables", "indicators", "formulas", "models", "diagnostic_rules"]:
                for item in reg_data.get(key, []):
                    item_str = json.dumps(item, ensure_ascii=False).lower()
                    item_id = str(item.get("variable_id", item.get("indicator_id", item.get("knowledge_id", "")))).lower()
                    item_name = str(item.get("name_cn", "")).lower()
                    if concept in item_str or concept in item_id or concept in item_name:
                        results["matched_schema_slices"].append({
                            "type": f"Registry{key.capitalize()[:-1] if key.endswith('s') else key.capitalize()}",
                            "contract": item
                        })
        except Exception as e:
            results["registry_error"] = str(e)

    # 3. 拦截林业保护政策规则校验
    if any(k in concept for k in ["间伐", "疏伐", "抚育", "采伐", "大径材", "经营"]):
        results["matched_schema_slices"].append({
            "type": "NationalParkPolicyAxiom",
            "rule_id": "ZERO_INTERVENTION_LAW",
            "principle": "零人为干预最高保护原则 (Zero-thinning Ecological Sanctuary Axiom)",
            "description": "祁连山国家公园严格执行生态休养法保护规范。绝对不可在此样区推荐或规划间伐、采伐等经济获利行为。高竞争与被压木一律解释为天然林自疏演替规律；监测行动应当以无损应力波、激光雷达或镜检替代物理采伐操作。"
        })

    if not results["matched_schema_slices"]:
        results["hint"] = "未检索到精确匹配的关联概念，建议查询基类：MonitoringPlot, Subplot, TreeIndividual, Taxon, DiagnosticRule, VolumeTableParameter"

    return json.dumps(results, ensure_ascii=False, indent=2)


# ==============================================================================
# 工具 2: 数理与神经符号规则引擎快捷绑定
# ==============================================================================
def tool_execute_neuro_symbolic_formula(knowledge_id: str, context_json: str = "{}") -> str:
    """神经-符号公式统一执行接口：参数一为公式ID (如 F_HEIGHT_DIAMETER_RATIO), 参数二为 JSON 变量字典"""
    try:
        ctx = json.loads(context_json) if isinstance(context_json, str) else context_json
        res = _ns_engine.execute_formula(knowledge_id, ctx)
        return json.dumps(res, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": f"神经符号公式执行失败: {e}"}, ensure_ascii=False)

def tool_evaluate_diagnostic_rules(subplot_id: str) -> str:
    """级联评估并诊断目标样方的森林健康，自动匹配阈值并读取 `forestry_knowledge_registry.yaml` 中的规则建议"""
    try:
        res = _ns_engine.evaluate_diagnostic_rules(subplot_id)
        return json.dumps(res, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": f"领域诊断规则推理评估失败: {e}"}, ensure_ascii=False)

def tool_math_calculator(expression: str) -> str:
    """安全数学计算器，仅接收纯数字的四则运算表达式"""
    try:
        if any(c not in set("0123456789+-*/(). ") for c in expression):
            return "错误：包含非法字符，只能包含数字与标准四则算术符"
        return f"计算结果: {eval(expression)}"
    except Exception as e:
        return f"计算错误: {e}"


def tool_convert_markdown_to_docx(markdown_path: str, output_path: str = None) -> str:
    """将 Markdown 文件转换为 Word (.docx)。"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return json.dumps({"error": "缺少依赖，请安装 python-docx: pip install python-docx"}, ensure_ascii=False)

    if not os.path.exists(markdown_path):
        return json.dumps({"error": f"Markdown 文件不存在: {markdown_path}"}, ensure_ascii=False)

    output_path = output_path or os.path.splitext(markdown_path)[0] + ".docx"
    doc = Document()
    with open(markdown_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    table_buffer = []
    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = [ [cell.strip() for cell in row.split("|") if cell.strip() != ""] for row in table_buffer ]
        if len(rows) >= 2:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    table.rows[i].cells[j].text = cell_text
        table_buffer = []

    for raw in lines:
        line = raw.rstrip("\n")
        if line.startswith("# "):
            flush_table()
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            flush_table()
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            flush_table()
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("![") and "](" in line and line.endswith(")"):
            flush_table()
            start = line.find("![") + 2
            end = line.find("](")
            alt = line[start:end]
            path = line[end+2:-1]
            if os.path.exists(path):
                try:
                    doc.add_paragraph(alt)
                    doc.add_picture(path, width=Inches(6))
                except Exception:
                    doc.add_paragraph(f"[图片无法插入: {path}]")
            else:
                doc.add_paragraph(f"[图片未找到: {path}]")
        elif line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
        elif line.strip().startswith(("- ", "* ", "+ ")):
            flush_table()
            doc.add_paragraph(line.strip(), style="List Bullet")
        elif line.strip().isdigit() and raw.strip().endswith("."):
            flush_table()
            doc.add_paragraph(line.strip(), style="List Number")
        else:
            if line.strip() == "":
                flush_table()
                doc.add_paragraph("")
            else:
                flush_table()
                doc.add_paragraph(line)
    flush_table()
    doc.save(output_path)
    return json.dumps({"success": True, "output_path": output_path}, ensure_ascii=False)


def tool_convert_markdown_to_pdf(markdown_path: str, output_path: str = None) -> str:
    """将 Markdown 文件转换为 PDF。如果可用，将先生成 docx 再转 pdf。"""
    if not os.path.exists(markdown_path):
        return json.dumps({"error": f"Markdown 文件不存在: {markdown_path}"}, ensure_ascii=False)

    output_path = output_path or os.path.splitext(markdown_path)[0] + ".pdf"
    try:
        import pypandoc
        pypandoc.convert_file(markdown_path, "pdf", outputfile=output_path)
        return json.dumps({"success": True, "output_path": output_path}, ensure_ascii=False)
    except ImportError:
        pass
    try:
        from docx2pdf import convert
        intermediate_docx = os.path.splitext(output_path)[0] + ".docx"
        convert_result = json.loads(tool_convert_markdown_to_docx(markdown_path, intermediate_docx))
        if not convert_result.get("success"):
            return json.dumps({"error": "无法先转换为 DOCX，无法生成 PDF"}, ensure_ascii=False)
        convert(intermediate_docx, output_path)
        return json.dumps({"success": True, "output_path": output_path}, ensure_ascii=False)
    except ImportError:
        return json.dumps({"error": "缺少依赖。请安装 pypandoc 或 docx2pdf，再重试。"}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": f"PDF 转换失败: {e}"}, ensure_ascii=False)


# ==============================================================================
# 工具 3: Cypher 图数据库提取器
# ==============================================================================
def tool_query_kg(cypher: str) -> str:
    """执行大模型自主生成的 Cypher 检索知识图谱，返回 JSON 字符串结果"""
    try:
        with _driver.session() as session:
            records = [dict(rec) for rec in session.run(cypher)]
            if not records:
                return json.dumps({"message": "Cypher 遍历为空，请确认节点标签是否符合本体"}, ensure_ascii=False)
            return json.dumps(records[:40], ensure_ascii=False, default=str)
    except Exception as e:
        if "Failed to establish connection" in str(e) or "Neo4j is not connected" in str(e) or "WinError 10061" in str(e):
            return json.dumps({"message": "本地未启动 Neo4j 图数据库，系统已自动切换为纯计算模式，请忽略图谱信息，直接基于本地样地实测数据报告结果。"}, ensure_ascii=False)
        return json.dumps({"error": f"Cypher 执行失败: {e}"}, ensure_ascii=False)


# ==============================================================================
# 工具 OpenAPI Schema 定义
# ==============================================================================
ONTOLOGY_INSPECT_SCHEMA = [{
    "type": "function",
    "function": {
        "name": "tool_inspect_ontology_schema",
        "description": "动态概念本体与公式契约检索器。当你需要查询特定的类（如 Subplot, Taxon）、字段映射（如 tree_dbh_cm）或者知识公式定义（如 HDR, Hegyi）时，自主提取对应规范切片以避免猜测属性或混淆表名！",
        "parameters": {
            "type": "object",
            "properties": {
                "query_concept": {"type": "string", "description": "要查询的领域概念、指标或公式名，如 'Subplot', 'TreeIndividual', 'HDR', 'Hegyi', 'Taxon'"}
            },
            "required": ["query_concept"]
        }
    }
}]

KG_QUERY_SCHEMA = [{
    "type": "function",
    "function": {
        "name": "tool_query_kg",
        "description": (
            "执行 Cypher 查询以抽取 Neo4j 知识图谱信息。图谱设计必须严格按《qilian_ontology.yaml》概念本体执行，包含的主要标签：\n"
            "- 类实体标签：:MonitoringPlot (样地), :Subplot (样方), :TreeIndividual (单木), :Taxon 或 :TreeSpecies (青海云杉等分类树种), :VolumeTableParameter (二元参数)\n"
            "- 核心关系：(Subplot)-[:HAS_TREE|containsTree]->(TreeIndividual)-[:BELONGS_TO_TAXON|IS_SPECIES]->(Taxon:TreeSpecies)\n"
            "⛔ 严禁臆测不存在的标签名。样方具体的公顷密度、高径比和单木坐标实测指标，系统已封装为專有物理算子，应直接调取算子工具而不是写 Cypher 在图里遍历海量单木！"
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "cypher": {"type": "string", "description": "标准本体标签的 Cypher 查询指令"}
            },
            "required": ["cypher"]
        }
    }
}]

NS_FORMULA_SCHEMA = [
    {
        "type": "function",
        "function": {
            "name": "tool_execute_neuro_symbolic_formula",
            "description": "调用神经-符号代数求值引擎，解析并计算声明式公式。这能对特定的林学变量求值，如 F_HEIGHT_DIAMETER_RATIO (高径比)、F_QUADRATIC_MEAN_DBH (平方平均胸径)。",
            "parameters": {
                "type": "object",
                "properties": {
                    "knowledge_id": {"type": "string", "description": "公式知识 ID，例如 'F_STAND_DENSITY', 'F_QUADRATIC_MEAN_DBH', 'F_HEIGHT_DIAMETER_RATIO'"},
                    "context_json": {"type": "string", "description": "传入计算所需的变量环境参数 json"}
                },
                "required": ["knowledge_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_evaluate_diagnostic_rules",
            "description": "对目标样方进行自动化诊断规则评估，自动匹配逻辑条件并输出对应的营林诊断信号与处方行动建议。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "小样方 ID 编号，如 '2816', '2901'"}
                },
                "required": ["subplot_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_math_calculator",
            "description": "安全数学计算器，用来辅助智能体完成对复杂数值的复合四则运算。",
            "parameters": {
                "type": "object",
                "properties": {
                    "expression": {"type": "string", "description": "纯数字的数学表达式，如 '(4.5 + 5.6) / 2'"}
                },
                "required": ["expression"]
            }
        }
    }
]

VISUALIZATION_TOOLS = [
    
    {
        "type": "function",
        "function": {
            "name": "tool_plot_subplot_grid_heatmap",
            "description": "生成 600 个样方空间格网热力图，用于展示样地指标的宏观空间分布。",
            "parameters": {
                "type": "object",
                "properties": {
                    "metric": {"type": "string", "description": "指标名称，例如 total_volume_m3、volume_per_ha、density_per_ha、shannon_index、mean_hdr、high_hdr_ratio_pct"}
                },
                "required": ["metric"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_size_class_distribution",
            "description": "生成样方胸径级分布图，展示林分径级结构与蓄积贡献。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"},
                    "target_type": {"type": "string", "description": "样本类型，例如 Subplot 或 All"}
                },
                "required": ["subplot_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_species_composition",
            "description": "生成样方树种组成对比图，展示株数、断面积和蓄积三维优势度。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"},
                    "target_type": {"type": "string", "description": "样本类型，例如 Subplot 或 All"}
                },
                "required": ["subplot_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_tree_relationship_scatter",
            "description": "生成单木胸径与树高关系散点图，并可显示高径比风险边界线。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"},
                    "x_var": {"type": "string", "description": "X 轴变量，例如 tree_dbh_cm"},
                    "y_var": {"type": "string", "description": "Y 轴变量，例如 tree_height_m"}
                },
                "required": ["subplot_id", "x_var", "y_var"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_group_comparison_boxplot",
            "description": "生成分组箱线图，比较不同组别间的变量分布和异常值。",
            "parameters": {
                "type": "object",
                "properties": {
                    "variable": {"type": "string", "description": "待比较变量，例如 hdr、tree_dbh_cm 或 tree_height_m"},
                    "group_by": {"type": "string", "description": "分组字段，例如 species"}
                },
                "required": ["variable", "group_by"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_tree_spatial_map",
            "description": "生成样方单木空间分布图，标注重点高风险树木。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"}
                },
                "required": ["subplot_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_subplot_percentile_profile",
            "description": "生成样方指标百分位画像图，展示目标样方在600个样方中的相对位置。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"}
                },
                "required": ["subplot_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_climate_time_series",
            "description": "生成区域气候时序图，展示生长季降水量和温度距平变化。",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_convert_markdown_to_docx",
            "description": "将 Markdown 报告转换为 Word 文档 (.docx)。",
            "parameters": {
                "type": "object",
                "properties": {
                    "markdown_path": {"type": "string", "description": "Markdown 文件路径"},
                    "output_path": {"type": "string", "description": "可选的 DOCX 输出路径"}
                },
                "required": ["markdown_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_convert_markdown_to_pdf",
            "description": "将 Markdown 报告转换为 PDF 文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "markdown_path": {"type": "string", "description": "Markdown 文件路径"},
                    "output_path": {"type": "string", "description": "可选的 PDF 输出路径"}
                },
                "required": ["markdown_path"]
            }
        }
    }
]

ALL_TOOLS = ONTOLOGY_INSPECT_SCHEMA + KG_QUERY_SCHEMA + NS_FORMULA_SCHEMA + FORESTRY_SPATIAL_SCHEMAS + VISUALIZATION_TOOLS


# ==============================================================================
# 智能体运行系统 System Prompt (祁连山国家公园天然林诊断专用)
# ==============================================================================
SYSTEM_PROMPT = """# 角色
你是祁连山国家公园森林生态智能体。你基于24公顷大样地的乔木调查数据（含坐标、胸径、树高、冠幅），回答生态查询、执行专业计算、诊断林分结构问题。

# 生态常识（你理解数据的底层逻辑）
1. 祁连山属于高寒半干旱生态区，水热条件和空间竞争是限制森林生长的核心因子。
2. 关键树种的生态性格：
   - 青海云杉：顶极群落优势种，耐阴，长寿。
   - 白桦/山杨：先锋树种，喜光，速生，短寿。
   - 祁连圆柏：顶极伴生种，极耐旱，喜阳。
   - 乌柳/花楸：湿生环境指示种，多分布于沟谷溪边。
3. 林分诊断需区分“客观观测值”（胸径/树高）与“派生诊断指标”（竞争指数/高径比）。

# 可用工具
以下工具由后台执行，你只需决定何时调用，无需关心如何计算。
1. **查询工具**：提取单木或样方的基础数据（数量、最大值、归属关系）。
2. **统计工具**：计算林分结构性指标（平均胸径、树高、密度、径级分布、蓄积量）。
3. **空间工具**：计算目标树周围邻近树木的距离和竞争关系。
4. **多样性工具**：计算 Shannon 指数、Simpson 指数、优势树种占比。
5. **诊断工具**：综合对比参照系，输出结构/更新/活力/风险四个维度的诊断信号。
6. **可视化工具**：生成径级分布柱状图、空间格局散点图、多维生态对比雷达图。
7. **文献检索工具**：当需要引用标准或学术依据时，检索本地知识网络。

# 推理流程
1. 解析用户意图：事实查询、对比分析、综合诊断还是报告生成。
2. 判断所需工具和数据维度，优先采用【并发】机制一次拉取分析所需的多维数据，严禁低效死循环。
3. 将计算结果与生态常识结合，给出深度解读（禁止机械罗列数字，必须结合树种习性推演演替动力）。
4. 若为诊断类问题，给出置信度（高/中/低）和符合国家公园零干预法则的建议。

# 输出风格
- 结论有据：注明确切的数据来源或计算依据。
- 图文结合：面对宏观报告或对比指令，自主调用可视化作图算子，并将图片 Markdown 嵌入回答对应位置实现视觉化。
- 承认边界：无法回答或无测算工具支持时明确说明，绝不推测。
"""


# ==============================================================================
# 并行分发后台调度器
# ==============================================================================
def _execute_tool(fn: str, args: dict) -> str:
    """将 LLM 请求的函数映射分发到本地物理执行单元，返回规范 JSON。"""
    sid = str(args.get("subplot_id", "")).strip()

    if fn == "tool_inspect_ontology_schema":
        return tool_inspect_ontology_schema(args.get("query_concept", ""))
    elif fn == "tool_query_kg":
        return tool_query_kg(args.get("cypher", ""))
    elif fn == "tool_execute_neuro_symbolic_formula":
        return tool_execute_neuro_symbolic_formula(args.get("knowledge_id", ""), args.get("context_json", "{}"))
    elif fn == "tool_evaluate_diagnostic_rules":
        return tool_evaluate_diagnostic_rules(sid)
    elif fn == "tool_math_calculator":
        return tool_math_calculator(args.get("expression", ""))
    elif fn == "tool_calc_stand_structure_metrics":
        return tool_calc_stand_structure_metrics(sid)
    elif fn == "tool_calc_tree_morphology_metrics":
        return tool_calc_tree_morphology_metrics(sid, args.get("target_tree_id"))
    elif fn == "tool_calc_species_diversity_metrics":
        return tool_calc_species_diversity_metrics(sid, args.get("survey_event_id", "EVENT_2023"))
    elif fn == "tool_calc_volume_metrics":
        return tool_calc_volume_metrics(sid)
    elif fn == "tool_calc_deadwood_metrics":
        return tool_calc_deadwood_metrics(sid)
    elif fn == "tool_calc_shrub_metrics":
        return tool_calc_shrub_metrics(sid)
    elif fn == "tool_calc_hegyi_competition":
        return tool_calc_hegyi_competition(sid, args.get("target_tree_id", ""), args.get("radius_m", 6.0))
    elif fn == "tool_calc_carbon_and_hydrology_metrics":
        return tool_calc_carbon_and_hydrology_metrics(sid)
    elif fn == "tool_assess_climate_and_biosecurity_risks":
        return tool_assess_climate_and_biosecurity_risks(sid)
    elif fn == "tool_scan_subplots_risk_summary":
        return tool_scan_subplots_risk_summary(args.get("start_subplot_id", "0101"), args.get("end_subplot_id", "0120"))
    elif fn == "tool_plot_ecosystem_radar_chart":
        return tool_plot_ecosystem_radar_chart(args.get("subplot_ids_str", sid))
    elif fn == "tool_plot_subplot_grid_heatmap":
        if not _FORESTRY_VIS_ENGINE_AVAILABLE:
            return _visualization_engine_error_response(fn)
        return json.dumps(plot_subplot_grid_heatmap(args.get("metric", "total_volume_m3")), ensure_ascii=False)
    elif fn == "tool_plot_size_class_distribution":
        if not _FORESTRY_VIS_ENGINE_AVAILABLE:
            return _visualization_engine_error_response(fn)
        return json.dumps(plot_size_class_distribution(args.get("subplot_id", sid), args.get("target_type", "Subplot")), ensure_ascii=False)
    elif fn == "tool_plot_species_composition":
        if not _FORESTRY_VIS_ENGINE_AVAILABLE:
            return _visualization_engine_error_response(fn)
        return json.dumps(plot_species_composition(args.get("subplot_id", sid), args.get("target_type", "Subplot")), ensure_ascii=False)
    elif fn == "tool_plot_tree_relationship_scatter":
        if not _FORESTRY_VIS_ENGINE_AVAILABLE:
            return _visualization_engine_error_response(fn)
        return json.dumps(plot_tree_relationship_scatter(args.get("subplot_id", sid), args.get("x_var", "tree_dbh_cm"), args.get("y_var", "tree_height_m")), ensure_ascii=False)
    elif fn == "tool_plot_group_comparison_boxplot":
        if not _FORESTRY_VIS_ENGINE_AVAILABLE:
            return _visualization_engine_error_response(fn)
        return json.dumps(plot_group_comparison_boxplot(args.get("variable", "hdr"), args.get("group_by", "species")), ensure_ascii=False)
    elif fn == "tool_plot_tree_spatial_map":
        if not _FORESTRY_VIS_ENGINE_AVAILABLE:
            return _visualization_engine_error_response(fn)
        return json.dumps(plot_tree_spatial_map(args.get("subplot_id", sid)), ensure_ascii=False)
    elif fn == "tool_plot_subplot_percentile_profile":
        if not _FORESTRY_VIS_ENGINE_AVAILABLE:
            return _visualization_engine_error_response(fn)
        return json.dumps(plot_subplot_percentile_profile(args.get("subplot_id", sid)), ensure_ascii=False)
    elif fn == "tool_plot_climate_time_series":
        if not _FORESTRY_VIS_ENGINE_AVAILABLE:
            return _visualization_engine_error_response(fn)
        return json.dumps(plot_climate_time_series(), ensure_ascii=False)
    elif fn == "tool_convert_markdown_to_docx":
        return tool_convert_markdown_to_docx(args.get("markdown_path", ""), args.get("output_path", None))
    elif fn == "tool_convert_markdown_to_pdf":
        return tool_convert_markdown_to_pdf(args.get("markdown_path", ""), args.get("output_path", None))
    else:
        return json.dumps({"error": f"未知工具名: {fn}"}, ensure_ascii=False)


def _execute_tool_calls_parallel(tool_calls: list) -> list:
    """多线程并发执行一系列工具调用，返回封装好的 tool messages"""
    results = {}

    def run_wrapper(call):
        fn = call["function"]["name"]
        try:
            args = json.loads(call["function"].get("arguments", "{}"))
        except Exception:
            args = {}
        t_start = time.time()
        res_str = _execute_tool(fn, args)
        t_elapsed = time.time() - t_start
        return call["id"], fn, args, res_str, t_elapsed

    # 并行执行所有工具
    with ThreadPoolExecutor(max_workers=min(len(tool_calls), 8)) as executor:
        futures = [executor.submit(run_wrapper, c) for c in tool_calls]
        for f in as_completed(futures):
            cid, fn, args, res_str, elapsed = f.result()
            results[cid] = (fn, args, res_str, elapsed)

    # 规范编排回传报文
    tool_messages = []
    for call in tool_calls:
        cid = call["id"]
        fn, args, res_str, elapsed = results[cid]
        preview = res_str[:250] + "..." if len(res_str) > 250 else res_str
        print(f"    ├─ [{cid}] 并行执行 {fn}({args}) → {elapsed:.2f}s → {preview.strip()}")
        tool_messages.append({
            "role": "tool",
            "tool_call_id": cid,
            "name": fn,
            "content": res_str
        })
    return tool_messages


# ==============================================================================
# 新 ReAct 主循环 (Parallel-Aware)
# ==============================================================================
def run_agent(question: str, report_file: str = "report.txt", max_rounds: int = 15):
    """
    启动 ReAct 主循环。
    LLM 将根据 System Prompt 决定工具组合，单轮能同时运行 N 个工具，大幅提升诊断自适应能力。
    """
    print(f"\n=================================================================")
    print(f" [Forestry Agent] 正在启动祁连山天然林智能诊断引擎 (ReAct Mode)")
    print(f" 问题: {question}")
    print(f"=================================================================\n")

    messages = [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": question}
    ]

    total_tools_called = 0
    t_engine_start = time.time()

    for rnd in range(1, max_rounds + 1):
        print(f"  ── [第 {rnd} 轮] 诊断思考中 ──")
        t_call_start = time.time()
        
        response = chat_with_tools(messages, tools=ALL_TOOLS)
        llm_elapsed = time.time() - t_call_start

        content = response.get("content") or ""
        tool_calls = response.get("tool_calls") or []

        if content:
            print(f"  🧠 AI 思考阐释 ({llm_elapsed:.1f}s): {content[:300]}{'...' if len(content) > 300 else ''}")

        if tool_calls:
            n_tools = len(tool_calls)
            total_tools_called += n_tools
            print(f"  🔧 发射 {n_tools} 个并行工具调用:")
            
            # 追加上一轮的 Assistant 应答
            messages.append({
                "role": "assistant",
                "content": content,
                "tool_calls": tool_calls
            })

            t_exec = time.time()
            tool_msgs = _execute_tool_calls_parallel(tool_calls)
            exec_elapsed = time.time() - t_exec
            
            messages.extend(tool_msgs)
            print(f"  ✔ 本轮 {n_tools} 个并行算子返回成功，执行历时 {exec_elapsed:.2f}s\n")
        else:
            # 没有 tool_calls，代表推理链条收敛，生成了最终分析报告
            total_elapsed = time.time() - t_engine_start
            print(f"=================================================================")
            print(f" ✅ 智能体分析报告收敛。历时: {total_elapsed:.1f}s | 总计调用工具: {total_tools_called} 次")
            print(f"=================================================================\n")

            with open(report_file, "w", encoding="utf-8") as f:
                f.write(content)
            print(f" [OK] 报告成功存入本地文件: `{report_file}`\n")
            return content

    # 循环溢出恢复机制
    print(f"\n [Warning] 已突破上限最大轮次 ({max_rounds})，启动强制收敛收拢报告。")
    messages.append({
        "role": "user",
        "content": "已达最大分析时钟。请立即用已有数据停止决策工具，直接输出完整的排队对比诊断报告 markdown 文本。"
    })
    forced_resp = chat_with_tools(messages, tools=[])
    final_report = forced_resp.get("content", "（因外部接口调用强制收尾失败）")
    with open(report_file, "w", encoding="utf-8") as f:
        f.write(final_report)
    print(f" [OK] 溢出阶段强制生成报告并存入: `{report_file}`")
    return final_report


if __name__ == "__main__":
    if len(sys.argv) > 1:
        user_q = " ".join(sys.argv[1:])
    else:
        user_q = (
            "请分析样方0101-0120，请全面分析这些样方的森林质量，我需要去对其中风险最高的5个样方进行踏查，由于时间有限，每个样方只能复测5棵树，请给我一个包含图片和表格的报告，用于我有针对性地去野外调查。。"
        )
    run_agent(
        question=user_q,
        report_file="Forest_Quality_6.md",
        max_rounds=30
    )
    _driver.close()
