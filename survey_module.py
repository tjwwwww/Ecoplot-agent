# -*- coding: utf-8 -*-
"""
survey_module.py
================
野外调查模块 — AI 驱动的调查方案生成、现场记录与报告导出。

核心功能：
1. 根据自然语言需求，结合数据库已有数据，自动生成调查方案
2. 方案包含具体树木、样地、调查原因和建议行动
3. 现场逐条记录观察结果
4. 调查完成后生成对比报告
"""

import json
import math
import os
import sqlite3
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

BASE_DIR = Path(__file__).resolve().parent
try:
    from dotenv import load_dotenv
    load_dotenv(BASE_DIR / ".env")
except Exception:
    pass
DATA_DIR = BASE_DIR / "data"
DB_PATH = Path(os.getenv("FORESTRY_SQLITE_DB", str(DATA_DIR / "qilian_forest.db"))).expanduser().resolve()
REPORT_DIR = BASE_DIR / "reports"

# =============================================================================
# 数据库初始化
# =============================================================================

SURVEY_TABLES_SQL = """
CREATE TABLE IF NOT EXISTS field_survey_plans (
    plan_id INTEGER PRIMARY KEY AUTOINCREMENT,
    title TEXT NOT NULL,
    user_request TEXT NOT NULL,
    ai_analysis TEXT,
    status TEXT DEFAULT 'draft' CHECK(status IN ('draft','active','completed','cancelled')),
    summary TEXT,
    tree_count INTEGER DEFAULT 0,
    subplot_count INTEGER DEFAULT 0,
    completed_count INTEGER DEFAULT 0,
    created_at TEXT DEFAULT (datetime('now','localtime')),
    updated_at TEXT DEFAULT (datetime('now','localtime')),
    deleted_at TEXT,
    delete_reason TEXT
);

CREATE TABLE IF NOT EXISTS survey_recommendations (
    rec_id INTEGER PRIMARY KEY AUTOINCREMENT,
    plan_id INTEGER NOT NULL,
    tree_id TEXT,
    subplot_id TEXT,
    species TEXT,
    priority TEXT DEFAULT 'medium' CHECK(priority IN ('high','medium','low')),
    category TEXT,
    reason TEXT NOT NULL,
    suggested_actions TEXT,
    evidence_data TEXT,
    status TEXT DEFAULT 'pending' CHECK(status IN ('pending','completed','skipped')),
    completed_at TEXT,
    obs_id INTEGER,
    sort_order INTEGER DEFAULT 0,
    FOREIGN KEY (plan_id) REFERENCES field_survey_plans(plan_id)
);

CREATE TABLE IF NOT EXISTS field_observations (
    obs_id INTEGER PRIMARY KEY AUTOINCREMENT,
    plan_id INTEGER NOT NULL,
    rec_id INTEGER,
    tree_id TEXT,
    subplot_id TEXT,
    species TEXT,
    notes TEXT,
    health_status TEXT CHECK(health_status IN ('good','fair','poor','dead',NULL)),
    pest_signs TEXT CHECK(pest_signs IN ('yes','no','suspected',NULL)),
    phenophase TEXT CHECK(phenophase IN ('budding','flowering','fruiting','leaf_changing','dormant',NULL)),
    photo_paths TEXT,
    latitude REAL,
    longitude REAL,
    recorded_at TEXT DEFAULT (datetime('now','localtime')),
    updated_at TEXT DEFAULT (datetime('now','localtime')),
    FOREIGN KEY (plan_id) REFERENCES field_survey_plans(plan_id),
    FOREIGN KEY (rec_id) REFERENCES survey_recommendations(rec_id)
);

CREATE TABLE IF NOT EXISTS survey_site_briefs (
    brief_id INTEGER PRIMARY KEY AUTOINCREMENT,
    brief_key TEXT NOT NULL UNIQUE,
    analysis_text TEXT NOT NULL,
    payload_json TEXT,
    created_at TEXT DEFAULT (datetime('now','localtime')),
    updated_at TEXT DEFAULT (datetime('now','localtime'))
);
"""


def init_survey_db() -> None:
    """初始化调查相关的数据库表"""
    if not DB_PATH.exists():
        print(f"[survey] 数据库不存在: {DB_PATH}")
        return
    with sqlite3.connect(DB_PATH) as conn:
        conn.executescript(SURVEY_TABLES_SQL)
        existing_columns = {
            row[1] for row in conn.execute("PRAGMA table_info(field_survey_plans)").fetchall()
        }
        if "deleted_at" not in existing_columns:
            conn.execute("ALTER TABLE field_survey_plans ADD COLUMN deleted_at TEXT")
        if "delete_reason" not in existing_columns:
            conn.execute("ALTER TABLE field_survey_plans ADD COLUMN delete_reason TEXT")
        if "latest_report_text" not in existing_columns:
            conn.execute("ALTER TABLE field_survey_plans ADD COLUMN latest_report_text TEXT")
        if "latest_report_file" not in existing_columns:
            conn.execute("ALTER TABLE field_survey_plans ADD COLUMN latest_report_file TEXT")
        if "latest_report_mode" not in existing_columns:
            conn.execute("ALTER TABLE field_survey_plans ADD COLUMN latest_report_mode TEXT")
        if "latest_report_generated_at" not in existing_columns:
            conn.execute("ALTER TABLE field_survey_plans ADD COLUMN latest_report_generated_at TEXT")
        conn.commit()


def _dict_factory(cursor: sqlite3.Cursor, row: sqlite3.Row) -> dict:
    """将查询结果转换为字典"""
    return {col[0]: row[idx] for idx, col in enumerate(cursor.description)}


def _get_conn() -> sqlite3.Connection:
    if not DB_PATH.exists():
        raise FileNotFoundError(f"SQLite database not found: {DB_PATH}")
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = _dict_factory
    return conn


# =============================================================================
# 数据采集 — 为 AI 提供上下文
# =============================================================================

def _detect_species_column(conn) -> Optional[str]:
    rows = conn.execute("PRAGMA table_info(tree_observations)").fetchall()
    columns = [row["name"] for row in rows]
    candidates = ["species", "species_name", "species_cn", "accepted_name_cn", "taxon_name", "tree_species"]
    for column in candidates:
        if column in columns:
            count = conn.execute(
                f"SELECT COUNT(*) AS n FROM tree_observations WHERE {column} IS NOT NULL AND TRIM({column}) <> ''"
            ).fetchone()["n"]
            if count:
                return column
    return None


def _gather_species_overview() -> List[Dict[str, Any]]:
    """Get species overview with schema-tolerant species column detection."""
    with _get_conn() as conn:
        species_col = _detect_species_column(conn)
        if not species_col:
            return []
        rows = conn.execute(f"""
            SELECT
                {species_col} AS species,
                COUNT(*) AS count,
                ROUND(AVG(tree_dbh_cm), 1) AS avg_dbh,
                ROUND(AVG(tree_height_m), 1) AS avg_height,
                ROUND(AVG(CASE WHEN health_status = 'good' THEN 1.0 ELSE 0 END) * 100, 1) AS health_good_pct,
                ROUND(AVG(CASE WHEN health_status = 'poor' THEN 1.0 ELSE 0 END) * 100, 1) AS health_poor_pct
            FROM tree_observations
            WHERE {species_col} IS NOT NULL AND TRIM({species_col}) <> ''
            GROUP BY {species_col}
            ORDER BY count DESC
        """).fetchall()
    return rows


def _gather_health_anomaly_trees(limit: int = 30) -> List[Dict[str, Any]]:
    """获取健康异常的树木（健康状态为差/枯死，或高径比异常）"""
    with _get_conn() as conn:
        rows = conn.execute("""
            SELECT
                tree_id, subplot_id, species, tree_dbh_cm, tree_height_m,
                health_status, remarks,
                ROUND(100.0 * tree_height_m / tree_dbh_cm, 2) AS hdr
            FROM tree_observations
            WHERE (health_status IN ('倒木', '病木', '断头木')
                   OR (tree_dbh_cm > 0 AND tree_height_m > 0
                       AND 100.0 * tree_height_m / tree_dbh_cm > 100))
            AND species IS NOT NULL AND TRIM(species) <> ''
            ORDER BY
                CASE health_status WHEN '倒木' THEN 0 WHEN '病木' THEN 1 WHEN '断头木' THEN 2 ELSE 3 END,
                hdr DESC
            LIMIT ?
        """, (limit,)).fetchall()
    return rows


def _gather_climate_context() -> Dict[str, Any]:
    """获取气候背景"""
    context = {}
    with _get_conn() as conn:
        # 最近气候数据
        recent = conn.execute("""
            SELECT
                COUNT(*) AS days,
                ROUND(AVG(mean_temperature_c), 1) AS avg_temp,
                ROUND(SUM(precipitation_mm), 1) AS total_precip,
                ROUND(AVG(precipitation_mm), 1) AS avg_daily_precip
            FROM climate_daily_normalized
            WHERE observation_date >= date('now', '-90 days')
        """).fetchone()
        if recent and recent["days"] and recent["days"] > 0:
            context["recent_90d"] = recent

        # 去年同期对比
        last_year = conn.execute("""
            SELECT
                ROUND(SUM(precipitation_mm), 1) AS total_precip,
                ROUND(AVG(mean_temperature_c), 1) AS avg_temp
            FROM climate_daily_normalized
            WHERE observation_date BETWEEN date('now', '-1 year', '-90 days')
                                      AND date('now', '-1 year')
        """).fetchone()
        if last_year and last_year.get("total_precip"):
            context["same_period_last_year"] = last_year

        # 年降水量
        yearly = conn.execute("""
            SELECT strftime('%Y', observation_date) AS year,
                   ROUND(SUM(precipitation_mm), 1) AS annual_precip,
                   ROUND(AVG(mean_temperature_c), 1) AS annual_temp
            FROM climate_daily_normalized
            GROUP BY year
            ORDER BY year DESC
            LIMIT 3
        """).fetchall()
        if yearly:
            context["recent_years"] = yearly

    return context


def _table_exists(conn, table_name: str) -> bool:
    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type IN ('table','view') AND name=?",
        (table_name,),
    ).fetchone()
    return row is not None


def _gather_topography_context() -> List[Dict[str, Any]]:
    """Get topography overview from available table/view."""
    with _get_conn() as conn:
        source = "topography_observations" if _table_exists(conn, "topography_observations") else None
        if source is None and _table_exists(conn, "vw_tree_topography_context"):
            source = "vw_tree_topography_context"
        if source is None:
            return []
        rows = conn.execute(f"""
            SELECT
                subplot_id,
                ROUND(AVG(elevation_m), 1) AS mean_elevation,
                ROUND(AVG(slope_degree), 1) AS mean_slope,
                ROUND(AVG(aspect_degree), 1) AS mean_aspect
            FROM {source}
            WHERE subplot_id IS NOT NULL
            GROUP BY subplot_id
            LIMIT 600
        """).fetchall()
    return rows


def _gather_subplot_summary(subplot_ids: Optional[List[str]] = None) -> List[Dict[str, Any]]:
    """获取样地汇总"""
    where = ""
    params: List[Any] = []
    if subplot_ids:
        placeholders = ",".join("?" for _ in subplot_ids)
        where = f"WHERE subplot_id IN ({placeholders})"
        params = subplot_ids

    with _get_conn() as conn:
        rows = conn.execute(f"""
            SELECT
                subplot_id,
                COUNT(*) AS tree_count,
                COUNT(DISTINCT species) AS species_count,
                ROUND(AVG(tree_dbh_cm), 1) AS avg_dbh,
                ROUND(AVG(tree_height_m), 1) AS avg_height,
                ROUND(AVG(100.0 * tree_height_m / NULLIF(tree_dbh_cm, 0)), 2) AS avg_hdr,
                SUM(CASE WHEN health_status = '__never_match_dead__' THEN 1 ELSE 0 END) AS dead_count,
                SUM(CASE WHEN health_status = '__never_match_poor__' THEN 1 ELSE 0 END) AS poor_count,
                SUM(CASE WHEN health_status IN ('倒木', '病木', '断头木') THEN 1 ELSE 0 END) AS health_attention_count,
                SUM(CASE WHEN health_status = '倒木' THEN 1 ELSE 0 END) AS fallen_status_count,
                SUM(CASE WHEN health_status = '病木' THEN 1 ELSE 0 END) AS diseased_status_count,
                SUM(CASE WHEN health_status = '断头木' THEN 1 ELSE 0 END) AS broken_top_status_count
            FROM tree_observations
            {where}
            GROUP BY subplot_id
            ORDER BY subplot_id
        """, params).fetchall()
    return rows



def _row_to_dict(row: Any) -> Dict[str, Any]:
    return dict(row) if row is not None else {}


_SITE_BRIEF_CACHE: Optional[Dict[str, Any]] = None


def _read_site_brief_cache(brief_key: str = "default") -> Optional[Dict[str, Any]]:
    try:
        with _get_conn() as conn:
            row = conn.execute(
                "SELECT analysis_text, payload_json FROM survey_site_briefs WHERE brief_key=?",
                (brief_key,),
            ).fetchone()
    except Exception:
        return None
    if not row:
        return None
    try:
        payload = json.loads(row.get("payload_json") or "{}")
    except Exception:
        payload = {}
    if not isinstance(payload, dict):
        payload = {}
    payload["analysis_text"] = row.get("analysis_text") or payload.get("analysis_text") or ""
    return payload if payload.get("analysis_text") else None


def _write_site_brief_cache(brief: Dict[str, Any], brief_key: str = "default") -> None:
    analysis_text = str(brief.get("analysis_text") or "").strip()
    if not analysis_text:
        return
    payload_json = json.dumps(brief, ensure_ascii=False, default=str)
    with _get_conn() as conn:
        conn.execute(
            """
            INSERT INTO survey_site_briefs (brief_key, analysis_text, payload_json)
            VALUES (?, ?, ?)
            ON CONFLICT(brief_key) DO UPDATE SET
                analysis_text=excluded.analysis_text,
                payload_json=excluded.payload_json,
                updated_at=datetime('now','localtime')
            """,
            (brief_key, analysis_text, payload_json),
        )
        conn.commit()





def _fallback_site_analysis(site_summary: Dict[str, Any], top_species: List[Dict[str, Any]], terrain_summary: Dict[str, Any]) -> str:
    species_text = "\u3001".join(
        f"{item.get('species')}\uff08{item.get('count')}\u682a\uff09"
        for item in top_species[:4]
        if item.get("species")
    ) or "\u6811\u79cd\u7edf\u8ba1\u6682\u4e0d\u5b8c\u6574"
    if terrain_summary.get("min_elevation_m") is not None:
        terrain_text = (
            f"\u5730\u5f62\u8bb0\u5f55\u663e\u793a\uff0c\u6837\u5730\u8986\u76d6\u6d77\u62d4\u7ea6{terrain_summary.get('min_elevation_m')}-"
            f"{terrain_summary.get('max_elevation_m')}m\uff0c\u5e73\u5747\u5761\u5ea6\u7ea6{terrain_summary.get('mean_slope_degree')}\u00b0\u3002"
        )
    else:
        terrain_text = "\u5730\u5f62\u8bb0\u5f55\u6682\u4e0d\u5b8c\u6574\uff0c\u9700\u8981\u7ed3\u5408\u6837\u65b9\u4f4d\u7f6e\u548c\u73b0\u573a\u89c2\u5bdf\u8865\u5145\u5224\u65ad\u3002"
    return (
        f"\u5f53\u524d\u6837\u5730\u5df2\u8bb0\u5f55\u4e54\u6728{site_summary.get('tree_count', 0)}\u682a\u3001"
        f"\u6837\u65b9{site_summary.get('subplot_count', 0)}\u4e2a\u3001\u6811\u79cd{site_summary.get('species_count', 0)}\u79cd\u3002"
        f"\u4e3b\u8981\u6811\u79cd\u5305\u62ec{species_text}\u3002{terrain_text}"
        "\u4ece\u8c03\u67e5\u7406\u89e3\u4e0a\u770b\uff0c\u7b2c\u4e00\u8f6e\u5e94\u91cd\u70b9\u628a\u6797\u5206\u7ed3\u6784\u3001\u6811\u79cd\u7ec4\u6210\u3001\u5730\u5f62\u68af\u5ea6\u548c\u5f02\u5e38\u5355\u6728\u653e\u5728\u540c\u4e00\u5f20\u8bc1\u636e\u94fe\u91cc\u770b\uff0c"
        "\u4e0d\u8981\u53ea\u770b\u5355\u4e2a\u6307\u6807\u3002\u5efa\u8bae\u4f18\u5148\u6838\u67e5\u4ee3\u8868\u6027\u6837\u65b9\u3001\u8fb9\u7f18\u6216\u7279\u6b8a\u5730\u5f62\u6837\u65b9\u3001\u4ee5\u53ca\u5f62\u6001\u6216\u8bb0\u5f55\u5f02\u5e38\u7684\u5355\u6728\uff0c"
        "\u7528\u73b0\u573a\u89c2\u5bdf\u8865\u8db3\u51a0\u5c42\u53d7\u538b\u3001\u66f4\u65b0\u72b6\u51b5\u3001\u571f\u58e4\u6e7f\u5ea6\u3001\u75c5\u866b\u5bb3\u75d5\u8ff9\u548c\u4eba\u4e3a\u5e72\u6270\u7b49\u6570\u636e\u5e93\u65e0\u6cd5\u76f4\u63a5\u786e\u8ba4\u7684\u4fe1\u606f\u3002"
    )


def _generate_site_analysis_with_agent(snapshot: Dict[str, Any]) -> Optional[str]:
    try:
        from agent import run_agent_chat
    except Exception:
        return None
    prompt = (
        "\u4f60\u662f\u6837\u5730\u91ce\u5916\u8c03\u67e5\u4e13\u5bb6\u3002\u8bf7\u57fa\u4e8e\u4e0b\u9762\u7684\u6570\u636e\u5feb\u7167\uff0c\u5199\u4e00\u6bb5\u9762\u5411\u8c03\u67e5\u4eba\u5458\u7684\u6837\u5730\u8ba4\u77e5\u6458\u8981\u3002"
        "\u91cd\u70b9\u8bf4\u660e\u8fd9\u4e2a\u6837\u5730\u76ee\u524d\u5448\u73b0\u51fa\u7684\u6797\u5206\u3001\u6811\u79cd\u3001\u5730\u5f62\u548c\u6c14\u5019\u80cc\u666f\u7279\u70b9\uff0c\u4ee5\u53ca\u4e3a\u4e86\u66f4\u597d\u7406\u89e3\u8fd9\u4e2a\u6837\u5730\uff0c"
        "\u73b0\u573a\u8c03\u67e5\u5e94\u8be5\u91cd\u70b9\u89c2\u5bdf\u54ea\u4e9b\u73b0\u8c61\u3001\u8865\u5145\u54ea\u4e9b\u8bc1\u636e\u3002\u4e0d\u8981\u53ea\u7f57\u5217\u6570\u636e\u5e93\u6307\u6807\uff0c\u4e0d\u8981\u8bf4'\u82e5\u660e\u5929\u65f6\u95f4\u6709\u9650'\uff0c"
        "\u4e0d\u8981\u8f93\u51faJSON\u3002\u7528\u4e2d\u6587\uff0c2\u52303\u6bb5\uff0c\u5177\u4f53\u3001\u514b\u5236\u3001\u6709\u73b0\u573a\u6307\u5bfc\u610f\u4e49\u3002\n\n"
        f"\u6570\u636e\u5feb\u7167\uff1a{json.dumps(snapshot, ensure_ascii=False, default=str)}"
    )
    try:
        result = run_agent_chat(
            question=prompt,
            session_id="survey_site_brief_cache",
            client_id="survey_site_brief",
            context={"current_page": "survey_site_brief", "context_policy": "auto"},
            options={"max_tool_rounds": 0, "history_limit": 0},
        )
        answer = str(result.get("answer") or "").strip()
        return answer or None
    except Exception as exc:
        print(f"[survey] site brief agent failed: {exc}")
        return None


def get_site_survey_brief(force_refresh: bool = False) -> Dict[str, Any]:
    """Return AI-oriented site brief for field survey planning."""
    global _SITE_BRIEF_CACHE
    if _SITE_BRIEF_CACHE is not None and not force_refresh:
        return {"status": "success", "brief": _SITE_BRIEF_CACHE, "cached": True, "cache_source": "memory"}
    if not force_refresh:
        persisted = _read_site_brief_cache()
        if persisted:
            _SITE_BRIEF_CACHE = persisted
            return {"status": "success", "brief": persisted, "cached": True, "cache_source": "database"}

    context = _gather_survey_context()
    species = [dict(row) for row in context.get("species") or []]
    anomalies = [dict(row) for row in context.get("anomalies") or []]
    subplots = [dict(row) for row in context.get("subplots") or []]
    topography = [dict(row) for row in context.get("topography") or []]
    climate = context.get("climate") or {}

    with _get_conn() as conn:
        species_col = _detect_species_column(conn)
        species_count_expr = f"COUNT(DISTINCT CASE WHEN {species_col} IS NOT NULL AND TRIM({species_col}) <> '' THEN {species_col} END)" if species_col else "0"
        totals = conn.execute(f"""
            SELECT
                COUNT(*) AS tree_count,
                COUNT(DISTINCT subplot_id) AS subplot_count,
                {species_count_expr} AS species_count
            FROM tree_observations
        """).fetchone()

    site_summary = {
        "tree_count": int(totals.get("tree_count") or 0) if totals else 0,
        "species_count": int(totals.get("species_count") or 0) if totals else len(species),
        "subplot_count": int(totals.get("subplot_count") or 0) if totals else len({item.get("subplot_id") for item in subplots if item.get("subplot_id")}),
    }
    top_species = species[:5]
    high_attention = anomalies[:8]

    elevations = [float(item.get("mean_elevation")) for item in topography if item.get("mean_elevation") is not None]
    slopes = [float(item.get("mean_slope")) for item in topography if item.get("mean_slope") is not None]
    terrain_summary = {
        "subplot_samples": len(topography),
        "min_elevation_m": round(min(elevations), 1) if elevations else None,
        "max_elevation_m": round(max(elevations), 1) if elevations else None,
        "mean_slope_degree": round(sum(slopes) / len(slopes), 1) if slopes else None,
    }

    suggested_questions = [
        "\u57fa\u4e8e\u8fd9\u4e2a\u6837\u5730\u7279\u70b9\uff0c\u5e2e\u6211\u751f\u6210\u4e00\u6b21\u7efc\u5408\u8c03\u67e5\u65b9\u6848\u3002",
        "\u6211\u60f3\u91cd\u70b9\u4e86\u89e3\u4e3b\u8981\u6811\u79cd\u7684\u7a7a\u95f4\u5dee\u5f02\uff0c\u5e94\u8be5\u600e\u4e48\u8c03\u67e5\uff1f",
        "\u54ea\u4e9b\u6837\u65b9\u6216\u5355\u6728\u6700\u9002\u5408\u4f5c\u4e3a\u5f02\u5e38\u590d\u6838\u5bf9\u8c61\uff1f",
        "\u5982\u679c\u8981\u9a8c\u8bc1\u5730\u5f62\u68af\u5ea6\u4e0e\u6797\u5206\u7ed3\u6784\u5dee\u5f02\uff0c\u5e94\u8be5\u600e\u4e48\u5b89\u6392\uff1f",
    ]
    snapshot = {
        "site_summary": site_summary,
        "top_species": top_species,
        "terrain_summary": terrain_summary,
        "climate_context": climate,
        "attention_examples": high_attention,
    }
    analysis_text = _generate_site_analysis_with_agent(snapshot) or _fallback_site_analysis(site_summary, top_species, terrain_summary)

    _SITE_BRIEF_CACHE = {
        "analysis_text": analysis_text,
        "suggested_questions": suggested_questions,
        "data_snapshot": site_summary,
        "top_species": top_species,
        "terrain_summary": terrain_summary,
        "climate_context": climate,
        "attention_examples": high_attention,
    }
    _write_site_brief_cache(_SITE_BRIEF_CACHE)
    return {"status": "success", "brief": _SITE_BRIEF_CACHE, "cached": False, "cache_source": "generated"}


# =============================================================================
# AI 方案生成 — 使用智能体（ReAct Agent）分析生成
# =============================================================================

def generate_survey_plan(user_request: str) -> Dict[str, Any]:
    """
    根据用户的自然语言需求，使用智能体（ReAct Agent）生成调查方案。

    流程：
    1. 调用 agent.run_agent_chat()，智能体使用工具分析数据
    2. 智能体逐步分析：查树种分布 → 健康异常 → 气候背景 → 竞争压力
    3. 智能体结合用户需求综合判断，输出针对性方案
    4. 解析 JSON → 验证 → 保存到数据库

    与之前的硬编码方案不同，智能体可以自主决定查什么数据、怎么分析。
    """
    print(f"[survey] 🧠 使用智能体生成调查方案，需求: {user_request}")

    try:
        from agent import run_agent_chat
    except Exception as exc:
        print(f"[survey] 无法导入 agent: {exc}")
        return {"status": "error", "message": f"智能体模块不可用: {exc}"}

    # 构造智能体提示词 — 核心思路：先分析数据，再结合需求
    # Build a compact local evidence package first. The agent may still call tools,
    # but the plan must be grounded in the available survey data rather than a fixed template.
    planning_context = _gather_survey_context()
    planning_payload = json.dumps(planning_context, ensure_ascii=False, default=str)
    prompt_template = "You are a field ecology survey planning agent for a large forest monitoring plot.\nYour job is to create an executable and evidence-grounded field survey plan from the user's request and the local evidence package.\n\nPlanning principles:\n1. Understand the user's real objective before choosing targets.\n2. Use the local evidence package first; call tools only if additional verification is needed.\n3. Select only real tree_id or subplot_id values found in data or tool results.\n4. Balance three target types: priority checks, representative samples, and controls.\n5. Do not turn screening signals into causal conclusions. Use wording such as \"needs field verification\", \"comparison target\", \"attention signal\", or \"data quality check\".\n6. Respect user constraints such as time, people, species, subplot, route, and task limit.\n7. If the request is vague, proactively design a high-value plan from the data instead of asking the user to fill a form.\n\nUser request:\n__USER_REQUEST__\n\nLocal evidence package JSON:\n```json\n__PLANNING_PAYLOAD__\n```\n\nReturn STRICT JSON only. No Markdown. No extra explanation.\nAll values in title, summary, ai_analysis, reason, and suggested_actions MUST be written in Simplified Chinese.\n\nSchema:\n{\n  \"title\": \"Chinese plan title\",\n  \"summary\": \"2-4 Chinese sentences explaining what this plan covers and why\",\n  \"ai_analysis\": \"Chinese explanation of user intent, data facts, screening basis, selected priorities, and field verification boundaries\",\n  \"recommendations\": [\n    {\n      \"tree_id\": \"real tree_id or null\",\n      \"subplot_id\": \"real subplot_id\",\n      \"species\": \"species name or null\",\n      \"priority\": \"high/medium/low\",\n      \"category\": \"health_check/morphology/competition/topography/species_observation/control/data_quality\",\n      \"reason\": \"Chinese evidence-based reason for selecting this target; avoid causal claims\",\n      \"suggested_actions\": \"Chinese field instructions: how to locate it, what to check, what to measure, what to record\"\n    }\n  ]\n}\n\nQuality constraints:\n- recommendations should normally contain 5-15 items; do not exceed 25 unless the user explicitly asks for a comprehensive plan.\n- high = priority check; medium = representative sample; low = control or supplementary check.\n- If tree_id is provided, it must be real. For subplot-level tasks, tree_id may be null, but subplot_id must be real.\n- suggested_actions must be practical for field workers; do not write only \"further investigation\"."
    agent_prompt = (
        prompt_template
        .replace("__USER_REQUEST__", str(user_request or ""))
        .replace("__PLANNING_PAYLOAD__", planning_payload)
        .strip()
    )

    try:
        # 调用智能体 — 使用完整的 ReAct 循环（工具调用 + 推理）
        result = run_agent_chat(
            question=agent_prompt,
            session_id=None,
            client_id="survey_plan_generator",
            context={"current_page": "survey_planning", "survey_mode": True},
            options={
                "max_tool_rounds": 10,
                "history_limit": 0,
            },
        )

        answer = result.get("answer", "")
        used_tools = result.get("used_tools", [])

        print(f"[survey] 智能体使用了 {len(used_tools)} 个工具")
        print(f"[survey] 智能体回答前200字: {answer[:200]}")

        # 解析 JSON
        plan_data = _parse_llm_json(answer)

        if not plan_data:
            print("[survey] 智能体输出不是合法 JSON，尝试用 LLM 直接重试…")
            # 兜底：把智能体的回答发给 LLM 提取 JSON
            plan_data = _retry_extract_json(answer, user_request)

        if plan_data and plan_data.get("recommendations"):
            print(f"[survey] 成功解析方案: {plan_data.get('title', '无标题')}, "
                  f"{len(plan_data['recommendations'])} 条建议")
            return _save_plan_to_db(user_request, plan_data)

        # 智能体返回了内容但没有成功解析 JSON → 用确定性逻辑作为最终兜底
        print("[survey] 智能体未能输出结构化方案，使用规则兜底")
        return _generate_deterministic_plan(user_request, _gather_survey_context())

    except Exception as exc:
        print(f"[survey] 智能体调用异常: {exc}")
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": f"智能体分析失败: {exc}"}


def _gather_survey_context() -> Dict[str, Any]:
    """收集调查上下文数据（供兜底方案使用）"""
    species = _gather_species_overview()
    anomalies = _gather_health_anomaly_trees(40)
    climate = _gather_climate_context()
    subplots = _gather_subplot_summary()
    topography = _gather_topography_context()
    return {
        "species_overview": species[:15],
        "health_anomalies": anomalies,
        "climate_context": climate,
        "subplot_summaries": subplots[:15],
        "topography_samples": topography[:10],
        "total_trees_in_db": sum(s["count"] for s in species) if species else 0,
        "total_species": len(species),
        "total_subplots": len(subplots),
    }


def _retry_extract_json(agent_answer: str, user_request: str) -> Optional[Dict[str, Any]]:
    """当智能体回答不是纯 JSON 时，用 LLM 提取 JSON"""
    try:
        from provider import get_ai_response
        prompt = f"""从下面的文本中提取调查方案的 JSON 数据。
文本是 AI 林业助手对调查需求的分析回复，其中应该包含了调查建议。

用户需求：{user_request}

AI 回复：
{agent_answer}

请提取 JSON 格式的方案，格式要求：
{{
  "title": "方案标题",
  "summary": "方案概述",
  "ai_analysis": "分析说明",
  "recommendations": [
    {{"tree_id": "编号或null", "subplot_id": "样地", "species": "树种", "priority": "high/medium/low", "category": "类别", "reason": "原因", "suggested_actions": "行动"}}
  ]
}}

只输出 JSON，不要其他文字。"""
        text = get_ai_response(content=prompt, prompt="提取 JSON", temperature=0.1)
        return _parse_llm_json(text)
    except Exception:
        return None


def _parse_llm_json(text: str) -> Optional[Dict[str, Any]]:
    """从 LLM 回复中提取 JSON"""
    if not text:
        return None

    # 尝试直接解析
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # 尝试提取 ```json ... ``` 块
    import re
    match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass

    # 尝试从中提取第一个 { ... } 块
    brace_start = text.find('{')
    brace_end = text.rfind('}')
    if brace_start >= 0 and brace_end > brace_start:
        try:
            return json.loads(text[brace_start:brace_end + 1])
        except json.JSONDecodeError:
            pass

    return None


def _save_plan_to_db(user_request: str, plan_data: Dict[str, Any]) -> Dict[str, Any]:
    """验证并保存方案到数据库"""
    title = str(plan_data.get("title", "野外调查方案")).strip()
    summary = str(plan_data.get("summary", "")).strip()
    ai_analysis = str(plan_data.get("ai_analysis", "")).strip()
    recommendations = plan_data.get("recommendations", plan_data.get("items", []))

    if not recommendations:
        return {"status": "error", "message": "方案中没有调查建议"}

    # 验证推荐的真实性
    valid_recs = _validate_recommendations(recommendations)

    if not valid_recs:
        return {"status": "error", "message": "验证失败：生成的建议中没有有效的树木/样地数据"}

    # 写入数据库
    with _get_conn() as conn:
        cursor = conn.execute(
            "INSERT INTO field_survey_plans (title, user_request, ai_analysis, summary, status, tree_count, subplot_count) "
            "VALUES (?, ?, ?, ?, 'active', ?, ?)",
            (
                title or "野外调查方案",
                user_request,
                ai_analysis,
                summary,
                sum(1 for r in valid_recs if r.get("tree_id")),
                len(set(r["subplot_id"] for r in valid_recs if r.get("subplot_id"))),
            ),
        )
        plan_id = cursor.lastrowid

        for idx, rec in enumerate(valid_recs):
            conn.execute(
                "INSERT INTO survey_recommendations "
                "(plan_id, tree_id, subplot_id, species, priority, category, reason, suggested_actions, evidence_data, status, sort_order) "
                "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'pending', ?)",
                (
                    plan_id,
                    rec.get("tree_id"),
                    rec.get("subplot_id"),
                    rec.get("species"),
                    rec.get("priority", "medium"),
                    rec.get("category", "health_check"),
                    rec.get("reason", ""),
                    rec.get("suggested_actions", ""),
                    rec.get("evidence_data"),
                    idx + 1,
                ),
            )
        conn.commit()

    # 返回完整方案
    return _get_plan_full(plan_id)


def _validate_recommendations(recommendations: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """验证建议中的树木是否真实存在于数据库"""
    valid = []
    tree_ids_to_check = []
    subplot_ids_to_check = set()

    for rec in recommendations:
        tid = str(rec.get("tree_id") or "").strip() if rec.get("tree_id") else ""
        spid = str(rec.get("subplot_id") or "").strip() if rec.get("subplot_id") else ""

        if tid:
            tree_ids_to_check.append(tid)
        if spid:
            subplot_ids_to_check.add(spid)

    # 批量查询树是否存在
    existing_trees = set()
    existing_subplots = set()

    if tree_ids_to_check:
        with _get_conn() as conn:
            placeholders = ",".join("?" for _ in tree_ids_to_check)
            rows = conn.execute(
                f"SELECT DISTINCT tree_id FROM tree_observations WHERE tree_id IN ({placeholders})",
                tree_ids_to_check,
            ).fetchall()
            existing_trees = set(r["tree_id"] for r in rows)

    if subplot_ids_to_check:
        with _get_conn() as conn:
            placeholders = ",".join("?" for _ in subplot_ids_to_check)
            rows = conn.execute(
                f"SELECT DISTINCT subplot_id FROM tree_observations WHERE subplot_id IN ({placeholders})",
                list(subplot_ids_to_check),
            ).fetchall()
            existing_subplots = set(r["subplot_id"] for r in rows)
            # 也检查 topography
            rows2 = conn.execute(
                f"SELECT DISTINCT subplot_id FROM topography_observations WHERE subplot_id IN ({placeholders})",
                list(subplot_ids_to_check),
            ).fetchall()
            existing_subplots.update(r["subplot_id"] for r in rows2)

    for rec in recommendations:
        tid = str(rec.get("tree_id") or "").strip() if rec.get("tree_id") else ""
        spid = str(rec.get("subplot_id") or "").strip() if rec.get("subplot_id") else ""

        # 如果有 tree_id，必须在数据库中
        if tid and tid not in existing_trees:
            print(f"[survey] 跳过不存在的树: {tid}")
            continue

        # 如果有 subplot_id 但没有 tree_id，验证样地
        if spid and not tid and spid not in existing_subplots:
            print(f"[survey] 跳过不存在的样地: {spid}")
            continue

        # 补充物种信息
        species = rec.get("species", "")
        if not species and tid:
            with _get_conn() as conn:
                row = conn.execute(
                    "SELECT species FROM tree_observations WHERE tree_id=? LIMIT 1",
                    (tid,),
                ).fetchone()
                if row:
                    species = row["species"]

        valid.append({
            "tree_id": tid or None,
            "subplot_id": spid or None,
            "species": species or rec.get("species", ""),
            "priority": rec.get("priority", "medium"),
            "category": rec.get("category", "health_check"),
            "reason": rec.get("reason", ""),
            "suggested_actions": rec.get("suggested_actions", ""),
        })

    return valid


# =============================================================================
# 确定性方案生成（LLM 不可用时的后备）
# =============================================================================

def _generate_deterministic_plan(user_request: str, context: Dict[str, Any]) -> Dict[str, Any]:
    """使用确定性规则生成调查方案（无需 LLM）"""
    print("[survey] 使用确定性规则生成方案（LLM 后备）")

    anomalies = context.get("health_anomalies", [])
    species_list = context.get("species_overview", [])
    subplots = context.get("subplot_summaries", [])

    recommendations = []
    seen_trees = set()

    # 1. 高优先级：健康异常的树木
    for tree in anomalies:
        if len(recommendations) >= 25:
            break
        tid = tree["tree_id"]
        if tid in seen_trees:
            continue
        seen_trees.add(tid)

        priority = "high" if tree.get("health_status") == "dead" else "high" if tree.get("health_status") == "poor" else "medium"
        category = "health_check"
        reason = f"健康状态为「{tree.get('health_status')}」"
        actions = "检查存活状态、树干是否完整、冠层状况、有无病虫害迹象"
        if tree.get("hdr") and tree["hdr"] > 100:
            reason += f"，高径比 HDR={tree['hdr']:.1f}（超过100阈值）"
            actions += "；检查树干是否倾斜、树冠是否稀疏"
            category = "morphology"

        recommendations.append({
            "tree_id": tid,
            "subplot_id": tree.get("subplot_id"),
            "species": tree.get("species"),
            "priority": priority,
            "category": category,
            "reason": reason,
            "suggested_actions": actions,
        })

    # 2. 按树种选健康对照
    species_counts = {}
    for s in species_list:
        species_counts[s["species"]] = s["count"]

    # 每个主要树种选1棵健康样木作为对照
    major_species = [s for s in species_list if s["count"] >= 10][:5]
    with _get_conn() as conn:
        for sp in major_species:
            if len(recommendations) >= 28:
                break
            # 找健康的该树种
            healthy = conn.execute("""
                SELECT tree_id, subplot_id, tree_dbh_cm, tree_height_m
                FROM tree_observations
                WHERE species=? AND health_status='good' AND tree_dbh_cm > 0
                ORDER BY RANDOM() LIMIT 1
            """, (sp["species"],)).fetchone()
            if healthy and healthy["tree_id"] not in seen_trees:
                seen_trees.add(healthy["tree_id"])
                recommendations.append({
                    "tree_id": healthy["tree_id"],
                    "subplot_id": healthy["subplot_id"],
                    "species": healthy["species"],
                    "priority": "low",
                    "category": "control",
                    "reason": f"健康对照样木（{sp['species']}）",
                    "suggested_actions": "记录正常状态，与异常个体对比",
                })

    # 3. 样地级建议
    for sp in subplots[:5]:
        if len(recommendations) >= 30:
            break
        attention_count = sp.get("health_attention_count", 0) or 0
        attention_pct = round(attention_count / max(sp["tree_count"], 1) * 100, 1)
        if attention_pct > 10:
            recommendations.append({
                "tree_id": None,
                "subplot_id": sp["subplot_id"],
                "species": None,
                "priority": "medium",
                "category": "health_check",
                "reason": f"\u6837\u65b9 {sp['subplot_id']} \u7684\u5065\u5eb7\u5173\u6ce8\u5bf9\u8c61\u6bd4\u4f8b\u4e3a {attention_pct}%",
                "suggested_actions": "\u73b0\u573a\u6838\u67e5\u6837\u65b9\u5185\u975e\u5065\u5eb7\u6728\u3001\u67af\u7acb\u6728\u3001\u5012\u6728\u548c\u9700\u8981\u8865\u5145\u8bb0\u5f55\u7684\u5bf9\u8c61\u3002",
            })

    title = "智能生成调查方案"
    summary = f"基于数据库分析，共生成 {len(recommendations)} 条调查建议，覆盖 {len(seen_trees)} 株树木。"
    ai_analysis = f"检测到 {len(anomalies)} 株健康异常树木，{len(major_species)} 个主要树种。"

    plan_data = {
        "title": title,
        "summary": summary,
        "ai_analysis": ai_analysis,
        "recommendations": recommendations,
    }
    return _save_plan_to_db(user_request, plan_data)


# =============================================================================
# 查询接口
# =============================================================================

def _build_field_check_items(rec: Dict[str, Any]) -> List[Dict[str, str]]:
    """Build practical field checklist items from a recommendation."""
    text = " ".join(str(rec.get(k) or "") for k in ("category", "reason", "suggested_actions")).lower()
    items = [
        {"key": "identity", "label": "\u6838\u5bf9\u5bf9\u8c61", "prompt": "\u6838\u5bf9\u6811\u53f7/\u6837\u65b9\u53f7\u662f\u5426\u4e0e\u63a8\u8350\u4efb\u52a1\u4e00\u81f4\uff0c\u5fc5\u8981\u65f6\u62cd\u6444\u6811\u724c\u6216\u4f4d\u7f6e\u7167\u7247\u3002"},
        {"key": "location", "label": "\u5b9a\u4f4d\u8bb0\u5f55", "prompt": "\u8bb0\u5f55\u5f53\u524d\u4f4d\u7f6e\u3001\u6837\u65b9\u5185\u76f8\u5bf9\u65b9\u4f4d\uff0c\u4ee5\u53ca\u662f\u5426\u5bb9\u6613\u518d\u6b21\u627e\u5230\u3002"},
    ]

    def add(key: str, label: str, prompt: str) -> None:
        if not any(item["key"] == key for item in items):
            items.append({"key": key, "label": label, "prompt": prompt})

    if any(word in text for word in ("health", "\u5065\u5eb7", "\u67af", "\u6b7b", "\u75c5", "\u866b", "\u53d7\u5bb3")):
        add("health", "\u5065\u5eb7\u72b6\u6001", "\u89c2\u5bdf\u53f6\u8272\u3001\u67af\u679d\u3001\u65ad\u68a2\u3001\u75c5\u866b\u75d5\u8ff9\u3001\u673a\u68b0\u635f\u4f24\uff0c\u5e76\u533a\u5206\u786e\u5b9a\u89c2\u6d4b\u4e0e\u5f85\u786e\u8ba4\u73b0\u8c61\u3002")
    if any(word in text for word in ("morph", "\u5f62\u6001", "hdr", "\u9ad8\u5f84", "\u80f8\u5f84", "\u6811\u9ad8", "\u51a0", "\u503e\u659c")):
        add("morphology", "\u5f62\u6001\u590d\u6838", "\u590d\u6838\u80f8\u5f84\u3001\u6811\u9ad8\u3001\u51a0\u5e45\u3001\u679d\u4e0b\u9ad8\u3001\u503e\u659c\u6216\u5f2f\u66f2\u60c5\u51b5\uff0c\u6807\u660e\u660e\u663e\u6d4b\u91cf\u5f02\u5e38\u3002")
    if any(word in text for word in ("competition", "\u7ade\u4e89", "\u53d7\u538b", "\u90bb\u6728", "\u90c1\u95ed")):
        add("competition", "\u90bb\u6728\u7ade\u4e89", "\u89c2\u5bdf\u5468\u56f4\u90bb\u6728\u5bc6\u5ea6\u3001\u51a0\u5c42\u906e\u6321\u3001\u540c\u79cd/\u5f02\u79cd\u90bb\u6728\uff0c\u4ee5\u53ca\u5bf9\u8c61\u6728\u662f\u5426\u660e\u663e\u53d7\u538b\u3002")
    if any(word in text for word in ("topography", "\u5730\u5f62", "\u6d77\u62d4", "\u5761", "\u6c9f", "\u810a", "\u5761\u5411")):
        add("topography", "\u5fae\u5730\u5f62", "\u8bb0\u5f55\u5761\u4f4d\u3001\u5761\u5411\u3001\u88f8\u5ca9\u3001\u51b2\u5237\u3001\u79ef\u6c34\u6216\u571f\u58e4\u6e7f\u5ea6\u7b49\u5fae\u5730\u5f62\u7ebf\u7d22\u3002")
    if any(word in text for word in ("climate", "\u6c14\u5019", "\u5e72\u65f1", "\u964d\u6c34", "\u4f4e\u6e29", "\u971c", "\u98ce", "\u96ea")):
        add("climate_exposure", "\u6c14\u5019\u66b4\u9732\u75d5\u8ff9", "\u89c2\u5bdf\u5e72\u65f1\u3001\u51bb\u5bb3\u3001\u98ce\u96ea\u538b\u3001\u843d\u679d\u843d\u53f6\u7b49\u66b4\u9732\u75d5\u8ff9\uff1b\u4e0d\u8981\u76f4\u63a5\u5f52\u56e0\u4e3a\u6c14\u5019\u3002")

    add("notes", "\u73b0\u573a\u5907\u6ce8", "\u8bb0\u5f55\u4e0e\u63a8\u8350\u7406\u7531\u76f8\u5173\u7684\u5173\u952e\u73b0\u8c61\u3001\u7167\u7247\u7f16\u53f7\u548c\u9700\u8981\u8865\u6d4b\u7684\u5b57\u6bb5\u3002")
    return items


def _get_plan_full(plan_id: int) -> Optional[Dict[str, Any]]:
    """Return a full survey plan with recommendations and saved observations."""
    with _get_conn() as conn:
        plan = conn.execute(
            "SELECT * FROM field_survey_plans WHERE plan_id=? AND deleted_at IS NULL",
            (plan_id,),
        ).fetchone()
        if not plan:
            return None

        recommendation_rows = conn.execute(
            "SELECT * FROM survey_recommendations WHERE plan_id=? ORDER BY sort_order",
            (plan_id,),
        ).fetchall()
        observations = conn.execute(
            "SELECT * FROM field_observations WHERE plan_id=?",
            (plan_id,),
        ).fetchall()

    obs_by_id = {obs.get("obs_id"): dict(obs) for obs in observations if obs.get("obs_id") is not None}
    obs_by_rec = {obs.get("rec_id"): dict(obs) for obs in observations if obs.get("rec_id") is not None}
    recommendations = []
    for row in recommendation_rows:
        rec = dict(row)
        observation = obs_by_id.get(rec.get("obs_id")) or obs_by_rec.get(rec.get("rec_id"))
        rec["observation"] = observation
        rec["field_check_items"] = _build_field_check_items(rec)
        recommendations.append(rec)

    result = dict(plan)
    result["recommendations"] = recommendations
    if not result.get("latest_report_text"):
        report_candidates = sorted(
            REPORT_DIR.glob(f"survey_report_{plan_id}_*.md"),
            key=lambda file_path: file_path.stat().st_mtime if file_path.exists() else 0,
            reverse=True,
        )
        if report_candidates:
            latest_path = report_candidates[0]
            try:
                result["latest_report_text"] = latest_path.read_text(encoding="utf-8")
                result["latest_report_file"] = latest_path.name
                result["latest_report_mode"] = result.get("latest_report_mode") or "file"
                result["latest_report_generated_at"] = result.get("latest_report_generated_at") or datetime.fromtimestamp(latest_path.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            except OSError:
                pass
    return {"status": "success", "plan": result}


def list_plans(limit: int = 20) -> Dict[str, Any]:
    """列出所有调查方案"""
    with _get_conn() as conn:
        plans = conn.execute(
            "SELECT * FROM field_survey_plans WHERE deleted_at IS NULL ORDER BY created_at DESC LIMIT ?",
            (limit,),
        ).fetchall()
    return {"status": "success", "plans": plans, "count": len(plans)}


def get_plan(plan_id: int) -> Dict[str, Any]:
    """获取方案详情"""
    result = _get_plan_full(plan_id)
    if not result:
        return {"status": "error", "message": f"方案不存在: plan_id={plan_id}"}
    return result


def revise_survey_plan(plan_id: int, instruction: str) -> Dict[str, Any]:
    """Create a revised survey plan from an existing plan and a natural-language instruction."""
    instruction = str(instruction or "").strip()
    if not instruction:
        return {"status": "error", "message": "revision instruction is required"}

    current = get_plan(plan_id)
    if current.get("status") != "success" or not current.get("plan"):
        return current

    plan = current["plan"]
    recommendations = plan.get("recommendations") or []
    compact_recs = []
    for rec in recommendations[:80]:
        compact_recs.append({
            "tree_id": rec.get("tree_id"),
            "subplot_id": rec.get("subplot_id"),
            "species": rec.get("species"),
            "priority": rec.get("priority"),
            "category": rec.get("category"),
            "reason": rec.get("reason"),
            "suggested_actions": rec.get("suggested_actions"),
        })

    revision_request = (
        "Revise this existing field survey plan according to the user's new instruction. "
        "Keep only real tree_id/subplot_id values already present in the database. "
        "Return a practical updated plan, not an explanation.\n\n"
        f"Original plan id: {plan_id}\n"
        f"Original title: {plan.get('title') or ''}\n"
        f"Original user request: {plan.get('user_request') or ''}\n"
        f"Original summary: {plan.get('summary') or ''}\n"
        f"Original recommendations JSON: {json.dumps(compact_recs, ensure_ascii=False)}\n\n"
        f"User revision instruction: {instruction}\n"
    )
    revised = generate_survey_plan(revision_request)
    if revised.get("status") == "success" and revised.get("plan"):
        new_plan_id = revised["plan"].get("plan_id")
        if new_plan_id:
            with _get_conn() as conn:
                conn.execute(
                    "UPDATE field_survey_plans SET summary=? WHERE plan_id=?",
                    (
                        f"Revised from plan {plan_id}. Instruction: {instruction}",
                        new_plan_id,
                    ),
                )
                conn.commit()
            revised = get_plan(int(new_plan_id))
            if revised.get("status") == "success":
                revised["revised_from_plan_id"] = plan_id
                revised["revision_instruction"] = instruction
    return revised


def update_plan_status(plan_id: int, status: str) -> Dict[str, Any]:
    """更新方案状态"""
    valid_statuses = {"draft", "active", "completed", "cancelled"}
    if status not in valid_statuses:
        return {"status": "error", "message": f"无效状态: {status}，可选: {valid_statuses}"}

    with _get_conn() as conn:
        conn.execute(
            "UPDATE field_survey_plans SET status=?, updated_at=datetime('now','localtime') WHERE plan_id=?",
            (status, plan_id),
        )
        conn.commit()
    return get_plan(plan_id)


# =============================================================================
# 现场记录
# =============================================================================

def _update_plan_completed_count(plan_id: int) -> None:
    """更新方案的完成进度"""
    with _get_conn() as conn:
        stats = conn.execute("""
            SELECT
                COUNT(*) AS total,
                SUM(CASE WHEN status='completed' THEN 1 ELSE 0 END) AS completed
            FROM survey_recommendations WHERE plan_id=?
        """, (plan_id,)).fetchone()
        if stats:
            conn.execute(
                "UPDATE field_survey_plans SET completed_count=?, updated_at=datetime('now','localtime') WHERE plan_id=?",
                (stats["completed"] or 0, plan_id),
            )
            # 如果全部完成，自动更新状态
            if stats["total"] and stats["completed"] and stats["total"] == stats["completed"]:
                conn.execute(
                    "UPDATE field_survey_plans SET status='completed', updated_at=datetime('now','localtime') WHERE plan_id=?",
                    (plan_id,),
                )
        conn.commit()



def delete_survey_plan(plan_id: int, reason: str = "") -> Dict[str, Any]:
    """
    Soft-delete a generated survey plan.

    This only hides the generated plan from the app. It does not delete any
    original tree, subplot, climate, topography, recommendation, or observation
    rows from the database.
    """
    reason = str(reason or "").strip()[:500]
    with _get_conn() as conn:
        plan = conn.execute(
            "SELECT plan_id, deleted_at FROM field_survey_plans WHERE plan_id=?",
            (plan_id,),
        ).fetchone()
        if not plan:
            return {"status": "not_found", "message": f"\u65b9\u6848\u4e0d\u5b58\u5728: plan_id={plan_id}"}
        if plan.get("deleted_at"):
            return {"status": "success", "plan_id": plan_id, "already_deleted": True}

        conn.execute(
            """
            UPDATE field_survey_plans
            SET status='cancelled',
                deleted_at=datetime('now','localtime'),
                delete_reason=?,
                updated_at=datetime('now','localtime')
            WHERE plan_id=?
            """,
            (reason or "\u7528\u6237\u5220\u9664\u751f\u6210\u7684\u8c03\u67e5\u65b9\u6848", plan_id),
        )
        conn.commit()

    return {
        "status": "success",
        "plan_id": plan_id,
        "message": "\u8c03\u67e5\u65b9\u6848\u5df2\u5220\u9664\u3002\u539f\u59cb\u6837\u5730\u3001\u5355\u6728\u3001\u6c14\u5019\u548c\u5730\u5f62\u6570\u636e\u672a\u88ab\u5220\u9664\u3002",
    }



def _refresh_survey_plan_counts(plan_id: int) -> None:
    """Refresh generated plan counters from survey_recommendations only."""
    with _get_conn() as conn:
        stats = conn.execute(
            """
            SELECT
                COUNT(*) AS total,
                COUNT(DISTINCT CASE WHEN subplot_id IS NOT NULL AND TRIM(subplot_id) <> '' THEN subplot_id END) AS subplot_count,
                SUM(CASE WHEN status='completed' THEN 1 ELSE 0 END) AS completed
            FROM survey_recommendations
            WHERE plan_id=?
            """,
            (plan_id,),
        ).fetchone()
        conn.execute(
            """
            UPDATE field_survey_plans
            SET tree_count=?, subplot_count=?, completed_count=?, updated_at=datetime('now','localtime')
            WHERE plan_id=?
            """,
            (
                (stats or {}).get("total") or 0,
                (stats or {}).get("subplot_count") or 0,
                (stats or {}).get("completed") or 0,
                plan_id,
            ),
        )
        conn.commit()


def purge_survey_plan(plan_id: int, delete_reports: bool = True) -> Dict[str, Any]:
    """
    Permanently delete a generated survey plan and its generated child records.

    Safe boundary: this never deletes original forestry data tables such as
    tree_observations, subplot data, climate data, topography data, or Neo4j data.
    """
    with _get_conn() as conn:
        plan = conn.execute(
            "SELECT plan_id FROM field_survey_plans WHERE plan_id=?",
            (plan_id,),
        ).fetchone()
        if not plan:
            return {"status": "not_found", "message": f"\u65b9\u6848\u4e0d\u5b58\u5728: plan_id={plan_id}"}

        obs_cursor = conn.execute("DELETE FROM field_observations WHERE plan_id=?", (plan_id,))
        rec_cursor = conn.execute("DELETE FROM survey_recommendations WHERE plan_id=?", (plan_id,))
        plan_cursor = conn.execute("DELETE FROM field_survey_plans WHERE plan_id=?", (plan_id,))
        conn.commit()

    removed_reports = []
    if delete_reports:
        REPORT_DIR.mkdir(parents=True, exist_ok=True)
        for file_path in REPORT_DIR.glob(f"survey_report_{plan_id}_*"):
            if file_path.is_file() and file_path.suffix.lower() in {".md", ".docx", ".pdf"}:
                try:
                    file_path.unlink()
                    removed_reports.append(file_path.name)
                except OSError:
                    pass

    return {
        "status": "success",
        "plan_id": plan_id,
        "deleted_plan_rows": plan_cursor.rowcount,
        "deleted_recommendation_rows": rec_cursor.rowcount,
        "deleted_observation_rows": obs_cursor.rowcount,
        "removed_reports": removed_reports,
        "message": "\u5df2\u6c38\u4e45\u5220\u9664\u751f\u6210\u7684\u8c03\u67e5\u65b9\u6848\uff1b\u539f\u59cb\u6837\u5730\u3001\u5355\u6728\u3001\u6c14\u5019\u548c\u5730\u5f62\u6570\u636e\u672a\u88ab\u5220\u9664\u3002",
    }


def delete_survey_recommendation(rec_id: int, delete_observation: bool = True) -> Dict[str, Any]:
    """
    Remove one generated recommendation/tree task from a survey plan.

    Safe boundary: this removes only the generated task row in the current plan.
    It never deletes the original tree record from tree_observations or the KG.
    """
    with _get_conn() as conn:
        rec = conn.execute(
            "SELECT rec_id, plan_id, tree_id, subplot_id, species FROM survey_recommendations WHERE rec_id=?",
            (rec_id,),
        ).fetchone()
        if not rec:
            return {"status": "not_found", "message": f"\u63a8\u8350\u9879\u4e0d\u5b58\u5728: rec_id={rec_id}"}

        plan_id = rec.get("plan_id")
        obs_deleted = 0
        if delete_observation:
            obs_cursor = conn.execute("DELETE FROM field_observations WHERE rec_id=?", (rec_id,))
            obs_deleted = obs_cursor.rowcount

        rec_cursor = conn.execute("DELETE FROM survey_recommendations WHERE rec_id=?", (rec_id,))
        conn.commit()

    _refresh_survey_plan_counts(int(plan_id))
    return {
        "status": "success",
        "plan_id": plan_id,
        "rec_id": rec_id,
        "tree_id": rec.get("tree_id"),
        "subplot_id": rec.get("subplot_id"),
        "species": rec.get("species"),
        "deleted_recommendation_rows": rec_cursor.rowcount,
        "deleted_observation_rows": obs_deleted,
        "message": "\u5df2\u4ece\u5f53\u524d\u8c03\u67e5\u65b9\u6848\u4e2d\u5220\u9664\u8be5\u6811/\u4efb\u52a1\uff1b\u539f\u59cb\u5355\u6728\u6570\u636e\u672a\u88ab\u5220\u9664\u3002",
    }


def save_observation(
    plan_id: int,
    rec_id: Optional[int] = None,
    tree_id: Optional[str] = None,
    subplot_id: Optional[str] = None,
    species: Optional[str] = None,
    notes: str = "",
    health_status: Optional[str] = None,
    pest_signs: Optional[str] = None,
    phenophase: Optional[str] = None,
    photo_paths: Optional[str] = None,
    latitude: Optional[float] = None,
    longitude: Optional[float] = None,
) -> Dict[str, Any]:
    """保存一条野外观察记录，并更新对应建议的状态"""
    # 如果有关联的建议，获取树信息
    if rec_id and not tree_id:
        with _get_conn() as conn:
            rec = conn.execute(
                "SELECT tree_id, subplot_id, species FROM survey_recommendations WHERE rec_id=?",
                (rec_id,),
            ).fetchone()
            if rec:
                tree_id = rec.get("tree_id") or tree_id
                subplot_id = rec.get("subplot_id") or subplot_id
                species = rec.get("species") or species

    with _get_conn() as conn:
        cursor = conn.execute(
            "INSERT INTO field_observations "
            "(plan_id, rec_id, tree_id, subplot_id, species, notes, health_status, pest_signs, phenophase, photo_paths, latitude, longitude) "
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (
                plan_id, rec_id, tree_id, subplot_id, species, notes,
                health_status, pest_signs, phenophase, photo_paths, latitude, longitude,
            ),
        )
        obs_id = cursor.lastrowid

    # 更新对应建议为已完成
    if rec_id:
        with _get_conn() as conn:
            conn.execute(
                "UPDATE survey_recommendations SET status='completed', completed_at=datetime('now','localtime'), obs_id=? WHERE rec_id=?",
                (obs_id, rec_id),
            )
            conn.commit()
        _update_plan_completed_count(plan_id)

    return get_observation(obs_id)


def update_observation(obs_id: int, **kwargs) -> Dict[str, Any]:
    """更新观察记录"""
    valid_fields = {"notes", "health_status", "pest_signs", "phenophase", "photo_paths"}
    updates = {k: v for k, v in kwargs.items() if k in valid_fields and v is not None}
    if not updates:
        return get_observation(obs_id)

    updates["updated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    set_clause = ", ".join(f"{k}=?" for k in updates)
    values = list(updates.values())

    with _get_conn() as conn:
        conn.execute(
            f"UPDATE field_observations SET {set_clause} WHERE obs_id=?",
            (*values, obs_id),
        )
        conn.commit()

    return get_observation(obs_id)


def get_observation(obs_id: int) -> Dict[str, Any]:
    """获取观察记录"""
    with _get_conn() as conn:
        obs = conn.execute(
            "SELECT * FROM field_observations WHERE obs_id=?",
            (obs_id,),
        ).fetchone()
    if not obs:
        return {"status": "error", "message": f"记录不存在: obs_id={obs_id}"}
    return {"status": "success", "observation": obs}


def get_plan_observations(plan_id: int) -> Dict[str, Any]:
    """获取方案的所有观察记录"""
    with _get_conn() as conn:
        observations = conn.execute(
            "SELECT * FROM field_observations WHERE plan_id=? ORDER BY recorded_at",
            (plan_id,),
        ).fetchall()
    return {"status": "success", "observations": observations, "count": len(observations)}


def update_recommendation_status(rec_id: int, status: str, obs_id: Optional[int] = None) -> Dict[str, Any]:
    """Update recommendation status without accidentally clearing saved observations."""
    with _get_conn() as conn:
        if obs_id is None:
            conn.execute(
                "UPDATE survey_recommendations SET status=?, completed_at=CASE WHEN ? THEN datetime('now','localtime') ELSE NULL END WHERE rec_id=?",
                (status, status == "completed", rec_id),
            )
        else:
            conn.execute(
                "UPDATE survey_recommendations SET status=?, completed_at=CASE WHEN ? THEN datetime('now','localtime') ELSE NULL END, obs_id=? WHERE rec_id=?",
                (status, status == "completed", obs_id, rec_id),
            )
        conn.commit()

    with _get_conn() as conn:
        rec = conn.execute(
            "SELECT plan_id FROM survey_recommendations WHERE rec_id=?",
            (rec_id,),
        ).fetchone()
    if rec:
        _update_plan_completed_count(rec["plan_id"])

    return {"status": "success", "rec_id": rec_id, "new_status": status}


init_survey_db()
print("[survey] 野外调查模块已加载")

# =============================================================================
# Leadership-oriented survey report generator (overrides earlier generate_report)
# =============================================================================
from typing import Any as _ReportAny, Dict as _ReportDict, List as _ReportList, Optional as _ReportOptional
import json as _report_json
import re as _report_re
from datetime import datetime as _report_datetime


def _survey_report_value(value: _ReportAny, default: str = "\u672a\u8bb0\u5f55") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text if text else default


def _survey_report_safe_title(title: str) -> str:
    safe = _report_re.sub(r"[^\w\u4e00-\u9fff-]+", "_", title or "survey_report")
    safe = _report_re.sub(r"_+", "_", safe).strip("_")
    return safe[:48] or "survey_report"


def _build_survey_report_evidence_package(
    plan_id: int,
    plan: _ReportDict[str, _ReportAny],
    recommendations: _ReportList[_ReportDict[str, _ReportAny]],
    observations: _ReportList[_ReportDict[str, _ReportAny]],
) -> _ReportDict[str, _ReportAny]:
    obs_by_rec = {obs.get("rec_id"): obs for obs in observations if obs.get("rec_id") is not None}
    completed_recs = [
        rec for rec in recommendations
        if str(rec.get("status") or "").lower() in {"done", "completed", "finished", "\u5df2\u5b8c\u6210"}
        or rec.get("rec_id") in obs_by_rec
    ]

    priority_counts: _ReportDict[str, int] = {}
    category_counts: _ReportDict[str, int] = {}
    for rec in recommendations:
        priority = _survey_report_value(rec.get("priority"), "\u672a\u5206\u7ea7")
        category = _survey_report_value(rec.get("category"), "\u672a\u5206\u7c7b")
        priority_counts[priority] = priority_counts.get(priority, 0) + 1
        category_counts[category] = category_counts.get(category, 0) + 1

    site_brief: _ReportDict[str, _ReportAny] = {}
    try:
        brief_result = get_site_survey_brief()
        if isinstance(brief_result, dict) and brief_result.get("status") == "success":
            site_brief = brief_result.get("brief") or {}
    except Exception as exc:
        site_brief = {"status": "unavailable", "error": str(exc)}

    compact_tasks = []
    for rec in recommendations:
        obs = obs_by_rec.get(rec.get("rec_id"), {})
        compact_tasks.append({
            "rec_id": rec.get("rec_id"),
            "target_type": rec.get("target_type"),
            "target_id": rec.get("target_id"),
            "target_name": rec.get("target_name"),
            "subplot_id": rec.get("subplot_id"),
            "priority": rec.get("priority"),
            "category": rec.get("category"),
            "reason": rec.get("reason"),
            "suggested_action": rec.get("suggested_action"),
            "status": rec.get("status"),
            "evidence_data": rec.get("evidence_data"),
            "field_observation": {
                "health_status": obs.get("health_status"),
                "pest_signs": obs.get("pest_signs"),
                "phenophase": obs.get("phenophase"),
                "notes": obs.get("notes"),
                "photos": obs.get("photos"),
                "recorded_at": obs.get("recorded_at"),
            } if obs else None,
        })

    return {
        "plan": {
            "plan_id": plan_id,
            "title": plan.get("title"),
            "user_request": plan.get("user_request"),
            "status": plan.get("status"),
            "summary": plan.get("summary"),
            "ai_analysis": plan.get("ai_analysis"),
            "created_at": plan.get("created_at"),
            "updated_at": plan.get("updated_at"),
        },
        "site_brief": site_brief,
        "execution_stats": {
            "total_tasks": len(recommendations),
            "completed_tasks": len(completed_recs),
            "pending_tasks": max(len(recommendations) - len(completed_recs), 0),
            "field_observation_count": len(observations),
            "priority_counts": priority_counts,
            "category_counts": category_counts,
        },
        "tasks": compact_tasks,
    }


_LAST_SURVEY_AGENT_REPORT_ERROR = ""


def _set_survey_agent_report_error(message: str) -> None:
    global _LAST_SURVEY_AGENT_REPORT_ERROR
    _LAST_SURVEY_AGENT_REPORT_ERROR = str(message or "").strip()[:1000]



def _normalize_report_audience(audience: str = "technical") -> str:
    value = str(audience or "technical").strip().lower()
    aliases = {
        "detail": "technical",
        "detailed": "technical",
        "expert": "technical",
        "field": "technical",
        "technical": "technical",
        "leader": "leader",
        "brief": "leader",
        "summary": "leader",
        "management": "leader",
    }
    return aliases.get(value, "technical")


def _report_audience_cn(audience: str) -> str:
    return "\u9886\u5bfc\u7b80\u62a5" if audience == "leader" else "\u8be6\u7ec6\u8c03\u67e5\u62a5\u544a"


def _generate_survey_report_with_agent(evidence_package: _ReportDict[str, _ReportAny], audience: str = "technical") -> _ReportOptional[str]:
    _set_survey_agent_report_error("")
    try:
        from agent import run_agent_chat
    except Exception as exc:
        _set_survey_agent_report_error(f"cannot import agent.run_agent_chat: {exc}")
        return None

    audience = _normalize_report_audience(audience)
    audience_name = _report_audience_cn(audience)
    if audience == "leader":
        structure = """
# \u91ce\u5916\u8c03\u67e5\u7b80\u62a5

## \u4e00\u3001\u6838\u5fc3\u7ed3\u8bba
\u75283\u52305\u6761\u6982\u62ec\u672c\u6b21\u8c03\u67e5\u4e3a\u4ec0\u4e48\u91cd\u8981\u3001\u91cd\u70b9\u5bf9\u8c61\u662f\u4ec0\u4e48\u3001\u4e3b\u8981\u8981\u89e3\u51b3\u4ec0\u4e48\u95ee\u9898\u3002

## \u4e8c\u3001\u8c03\u67e5\u5b89\u6392
\u6982\u62ec\u8c03\u67e5\u8303\u56f4\u3001\u5bf9\u8c61\u6570\u91cf\u3001\u5b8c\u6210\u8fdb\u5ea6\u548c\u4f18\u5148\u7ea7\uff0c\u4e0d\u5806\u780c\u660e\u7ec6\u3002

## \u4e09\u3001\u5173\u6ce8\u91cd\u70b9
\u8bf4\u660e\u73b0\u573a\u6700\u9700\u8981\u5173\u6ce8\u7684\u73b0\u8c61\u548c\u6570\u636e\u7f3a\u53e3\u3002

## \u56db\u3001\u9884\u671f\u4ea7\u51fa
\u8bf4\u660e\u8c03\u67e5\u7ed3\u675f\u540e\u80fd\u5f62\u6210\u4ec0\u4e48\u5224\u65ad\u3001\u56fe\u8868\u6216\u7ba1\u7406\u4f9d\u636e\u3002

## \u4e94\u3001\u5efa\u8bae\u4e8b\u9879
\u7ed9\u51fa\u9700\u8981\u534f\u8c03\u3001\u5b89\u6392\u6216\u51b3\u7b56\u7684\u4e8b\u9879\u3002
""".strip()
        required_sections = ["\u6838\u5fc3\u7ed3\u8bba", "\u8c03\u67e5\u5b89\u6392", "\u5173\u6ce8\u91cd\u70b9", "\u9884\u671f\u4ea7\u51fa"]
    else:
        structure = """
# \u91ce\u5916\u8c03\u67e5\u8be6\u7ec6\u62a5\u544a

## \u4e00\u3001\u8c03\u67e5\u76ee\u6807
\u8bf4\u660e\u672c\u6b21\u4e3a\u4ec0\u4e48\u8c03\u67e5\u8fd9\u4e9b\u5bf9\u8c61\uff1a\u8981\u6838\u67e5\u4ec0\u4e48\u73b0\u8c61\u3001\u8865\u8db3\u4ec0\u4e48\u8bc1\u636e\u3001\u670d\u52a1\u4ec0\u4e48\u5206\u6790\u95ee\u9898\u3002

## \u4e8c\u3001\u9009\u62e9\u4f9d\u636e
\u8bf4\u660e\u6837\u65b9\u3001\u5355\u6728\u6216\u6811\u79cd\u88ab\u7eb3\u5165\u65b9\u6848\u7684\u4f9d\u636e\uff0c\u533a\u5206\u7528\u6237\u8981\u6c42\u3001\u6570\u636e\u5e93\u4e8b\u5b9e\u3001\u6307\u6807\u7b5b\u67e5\u548c\u73b0\u573a\u8bb0\u5f55\u3002

## \u4e09\u3001\u6838\u67e5\u4efb\u52a1
\u628a\u73b0\u573a\u8981\u505a\u7684\u4e8b\u60c5\u5199\u6e05\u695a\uff1a\u770b\u4ec0\u4e48\u3001\u91cf\u4ec0\u4e48\u3001\u62cd\u4ec0\u4e48\u3001\u8bb0\u5f55\u4ec0\u4e48\u3002\u53ef\u7528\u8868\u683c\u5448\u73b0\u91cd\u70b9\u5bf9\u8c61\u3002

## \u56db\u3001\u9884\u671f\u5224\u65ad
\u8bf4\u660e\u8c03\u67e5\u5b8c\u6210\u540e\u80fd\u591f\u652f\u6301\u54ea\u4e9b\u5224\u65ad\u3001\u56fe\u8868\u6216\u540e\u7eed\u5206\u6790\uff0c\u4e5f\u8bf4\u660e\u4e0d\u80fd\u76f4\u63a5\u8bc1\u660e\u54ea\u4e9b\u56e0\u679c\u7ed3\u8bba\u3002

## \u4e94\u3001\u8865\u5145\u5efa\u8bae
\u7ed9\u51fa\u540e\u7eed\u590d\u6d4b\u3001\u8865\u91c7\u5b57\u6bb5\u3001\u6570\u636e\u66f4\u65b0\u548c\u98ce\u9669\u8fb9\u754c\u3002
""".strip()
        required_sections = ["\u8c03\u67e5\u76ee\u6807", "\u9009\u62e9\u4f9d\u636e", "\u6838\u67e5\u4efb\u52a1", "\u9884\u671f\u5224\u65ad"]

    payload = _report_json.dumps(evidence_package, ensure_ascii=False, default=str)
    prompt = f"""
You are a field ecological survey report agent. Generate one Markdown report from the evidence package.

Report audience: {audience_name}
Use this report structure exactly:

{structure}

Writing focus:
- The report must answer why these objects are investigated, not only describe what is in the plan.
- Explain what questions the survey can answer after completion.
- Keep leader brief concise and decision-oriented; keep technical report actionable for field staff.

Rules:
1. Use only facts in the evidence package. Do not invent observations, photos, counts, species, locations, or conclusions.
2. Distinguish database facts, generated-plan rationale, field observations, unresolved questions, and expected outputs.
3. If evidence only supports screening or association, use cautious wording such as "\u9700\u6838\u67e5", "\u5173\u6ce8\u4fe1\u53f7", "\u5bf9\u6bd4\u5bf9\u8c61", or "\u8bc1\u636e\u4e0d\u8db3".
4. Use Simplified Chinese.
5. Output Markdown only. Do not output JSON. Do not use code blocks.
6. Avoid dumping all raw records into the body.

Evidence package JSON:
```json
{payload}
```
""".strip()

    try:
        result = run_agent_chat(
            question=prompt,
            session_id=f"survey_report_{evidence_package.get('plan', {}).get('plan_id', 'unknown')}_{audience}",
            client_id="survey_report_generator",
            context={"current_page": "survey_report", "report_mode": True, "report_audience": audience, "context_policy": "auto"},
            options={"max_tool_rounds": 0, "history_limit": 0},
        )
    except TypeError as exc:
        try:
            result = run_agent_chat(prompt)
        except Exception as inner_exc:
            _set_survey_agent_report_error(f"run_agent_chat compatibility call failed: {inner_exc}; original TypeError: {exc}")
            return None
    except Exception as exc:
        _set_survey_agent_report_error(f"run_agent_chat failed: {exc}")
        return None

    if isinstance(result, dict):
        answer = result.get("answer") or result.get("content") or result.get("text")
    else:
        answer = str(result) if result is not None else ""
    answer = (answer or "").strip()
    if not answer:
        _set_survey_agent_report_error("agent returned empty report")
        return None
    if len(answer) < 120 or "#" not in answer or not all(section in answer for section in required_sections):
        _set_survey_agent_report_error(f"agent report does not match {audience_name} Markdown protocol; length={len(answer)}")
        return None
    return answer


def _generate_survey_report_fallback(evidence_package: _ReportDict[str, _ReportAny], audience: str = "technical") -> str:
    audience = _normalize_report_audience(audience)
    plan = evidence_package.get("plan") or {}
    stats = evidence_package.get("execution_stats") or {}
    site_brief = evidence_package.get("site_brief") or {}
    tasks = evidence_package.get("tasks") or []
    data_snapshot = (site_brief.get("data_snapshot") or {}) if isinstance(site_brief, dict) else {}
    observed_tasks = [task for task in tasks if task.get("field_observation")]
    user_request = _survey_report_value(plan.get("user_request"))
    ai_analysis = _survey_report_value(plan.get("ai_analysis"), "\u672c\u65b9\u6848\u7531\u667a\u80fd\u4f53\u6839\u636e\u7528\u6237\u9700\u6c42\u548c\u5df2\u6709\u6837\u5730\u6570\u636e\u751f\u6210\u3002")

    if audience == "leader":
        lines = [
            "# \u91ce\u5916\u8c03\u67e5\u7b80\u62a5",
            "",
            f"**\u751f\u6210\u65f6\u95f4**\uff1a{_report_datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"**\u8c03\u67e5\u65b9\u6848**\uff1a{_survey_report_value(plan.get('title'))}",
            "",
            "## \u4e00\u3001\u6838\u5fc3\u7ed3\u8bba",
            "",
            f"- \u672c\u6b21\u8c03\u67e5\u56f4\u7ed5\u7528\u6237\u95ee\u9898\u201c{user_request}\u201d\u7ec4\u7ec7\uff0c\u91cd\u70b9\u662f\u4e3a\u73b0\u573a\u6838\u67e5\u548c\u540e\u7eed\u5224\u65ad\u8865\u8db3\u8bc1\u636e\u3002",
            f"- \u65b9\u6848\u5305\u542b {stats.get('total_tasks', 0)} \u9879\u8c03\u67e5\u4efb\u52a1\uff0c\u5df2\u5b8c\u6210 {stats.get('completed_tasks', 0)} \u9879\uff0c\u5f85\u5b8c\u6210 {stats.get('pending_tasks', 0)} \u9879\u3002",
            f"- \u5f53\u524d\u5df2\u6709\u73b0\u573a\u8bb0\u5f55 {stats.get('field_observation_count', 0)} \u6761\uff0c\u62a5\u544a\u7ed3\u8bba\u9700\u4ee5\u73b0\u573a\u6838\u67e5\u7ed3\u679c\u66f4\u65b0\u3002",
            "",
            "## \u4e8c\u3001\u8c03\u67e5\u5b89\u6392",
            "",
            "| \u6307\u6807 | \u6570\u503c |",
            "|---|---:|",
            f"| \u8c03\u67e5\u4efb\u52a1\u6570 | {stats.get('total_tasks', 0)} |",
            f"| \u5df2\u5b8c\u6210 | {stats.get('completed_tasks', 0)} |",
            f"| \u5f85\u5b8c\u6210 | {stats.get('pending_tasks', 0)} |",
            f"| \u73b0\u573a\u8bb0\u5f55 | {stats.get('field_observation_count', 0)} |",
            "",
            "## \u4e09\u3001\u5173\u6ce8\u91cd\u70b9",
            "",
            ai_analysis,
            "",
            "## \u56db\u3001\u9884\u671f\u4ea7\u51fa",
            "",
            "- \u660e\u786e\u91cd\u70b9\u5bf9\u8c61\u662f\u5426\u771f\u5b9e\u5b58\u5728\u3001\u72b6\u6001\u662f\u5426\u4e0e\u7cfb\u7edf\u8bb0\u5f55\u4e00\u81f4\u3002",
            "- \u4e3a\u540e\u7eed\u6797\u5206\u7ed3\u6784\u3001\u5f02\u5e38\u6728\u590d\u6838\u3001\u6837\u65b9\u5bf9\u6bd4\u6216\u8865\u6d4b\u8ba1\u5212\u63d0\u4f9b\u8bc1\u636e\u3002",
            "- \u5f62\u6210\u53ef\u5f52\u6863\u7684\u73b0\u573a\u8bb0\u5f55\uff0c\u652f\u6301\u540e\u7eed\u66f4\u65b0\u6570\u636e\u5e93\u548c\u518d\u5206\u6790\u3002",
            "",
            "## \u4e94\u3001\u5efa\u8bae\u4e8b\u9879",
            "",
            "- \u4f18\u5148\u5b8c\u6210\u9ad8\u4f18\u5148\u7ea7\u548c\u4fe1\u606f\u7f3a\u53e3\u8f83\u5927\u7684\u5bf9\u8c61\u3002",
            "- \u8c03\u67e5\u7ed3\u675f\u540e\u91cd\u65b0\u751f\u6210\u8be6\u7ec6\u62a5\u544a\uff0c\u7528\u73b0\u573a\u8bb0\u5f55\u66ff\u4ee3\u63a8\u6d4b\u6027\u63cf\u8ff0\u3002",
            "- \u5f53\u524d\u62a5\u544a\u4e0d\u76f4\u63a5\u4f5c\u4e3a\u539f\u56e0\u5224\u5b9a\u6216\u7ecf\u8425\u51b3\u7b56\u7ed3\u8bba\u3002",
        ]
        return "\n".join(lines)

    lines = [
        "# \u91ce\u5916\u8c03\u67e5\u8be6\u7ec6\u62a5\u544a",
        "",
        f"**\u751f\u6210\u65f6\u95f4**\uff1a{_report_datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"**\u8c03\u67e5\u65b9\u6848**\uff1a{_survey_report_value(plan.get('title'))}",
        "",
        "## \u4e00\u3001\u8c03\u67e5\u76ee\u6807",
        "",
        f"\u672c\u6b21\u8c03\u67e5\u56f4\u7ed5\u7528\u6237\u9700\u6c42\u201c{user_request}\u201d\u5c55\u5f00\uff0c\u76ee\u6807\u662f\u901a\u8fc7\u73b0\u573a\u590d\u6838\u8865\u8db3\u6570\u636e\u5e93\u8bc1\u636e\uff0c\u786e\u8ba4\u63a8\u8350\u5bf9\u8c61\u7684\u771f\u5b9e\u72b6\u6001\uff0c\u5e76\u4e3a\u540e\u7eed\u5206\u6790\u63d0\u4f9b\u53ef\u9760\u8f93\u5165\u3002",
        f"\u65b9\u6848\u5171\u5305\u542b {stats.get('total_tasks', 0)} \u9879\u4efb\u52a1\uff0c\u5df2\u5b8c\u6210 {stats.get('completed_tasks', 0)} \u9879\uff0c\u5f85\u5b8c\u6210 {stats.get('pending_tasks', 0)} \u9879\uff0c\u5df2\u6709\u73b0\u573a\u8bb0\u5f55 {stats.get('field_observation_count', 0)} \u6761\u3002",
        "",
        "## \u4e8c\u3001\u9009\u62e9\u4f9d\u636e",
        "",
        ai_analysis,
        "",
    ]
    if data_snapshot:
        lines.extend([
            "\u7cfb\u7edf\u6570\u636e\u57fa\u7840\u5305\u62ec\uff1a",
            f"- \u4e54\u6728\u8bb0\u5f55\uff1a{_survey_report_value(data_snapshot.get('tree_count'))}",
            f"- \u6837\u65b9\u6570\u91cf\uff1a{_survey_report_value(data_snapshot.get('subplot_count'))}",
            f"- \u6811\u79cd\u6570\u91cf\uff1a{_survey_report_value(data_snapshot.get('species_count'))}",
            "",
        ])
    lines.extend([
        "## \u4e09\u3001\u6838\u67e5\u4efb\u52a1",
        "",
        "| \u5e8f\u53f7 | \u5bf9\u8c61 | \u6837\u65b9 | \u4f18\u5148\u7ea7 | \u5165\u9009\u7406\u7531 | \u73b0\u573a\u6838\u67e5\u91cd\u70b9 |",
        "|---:|---|---|---|---|---|",
    ])
    for index, task in enumerate(tasks[:30], start=1):
        target = _survey_report_value(task.get("target_name") or task.get("target_id"))
        action = _survey_report_value(task.get("suggested_action"), "\u6838\u67e5\u5bf9\u8c61\u72b6\u6001\u3001\u4f4d\u7f6e\u3001\u5065\u5eb7\u72b6\u51b5\u548c\u5f02\u5e38\u73b0\u8c61")
        lines.append(
            f"| {index} | {target} | {_survey_report_value(task.get('subplot_id'))} | "
            f"{_survey_report_value(task.get('priority'))} | {_survey_report_value(task.get('reason'))} | {action} |"
        )
    lines.extend(["", "## \u56db\u3001\u9884\u671f\u5224\u65ad", ""])
    lines.extend([
        "\u8c03\u67e5\u5b8c\u6210\u540e\uff0c\u53ef\u4ee5\u652f\u6301\u4ee5\u4e0b\u5224\u65ad\uff1a",
        "- \u63a8\u8350\u5bf9\u8c61\u662f\u5426\u771f\u5b9e\u5b58\u5728\uff0c\u73b0\u72b6\u662f\u5426\u4e0e\u7cfb\u7edf\u8bb0\u5f55\u4e00\u81f4\u3002",
        "- \u91cd\u70b9\u6837\u65b9\u6216\u5355\u6728\u662f\u5426\u9700\u8981\u7ee7\u7eed\u590d\u6d4b\u3001\u8865\u6d4b\u6216\u4ece\u65b9\u6848\u4e2d\u79fb\u9664\u3002",
        "- \u540e\u7eed\u662f\u5426\u53ef\u4ee5\u5f00\u5c55\u6837\u65b9\u5bf9\u6bd4\u3001\u6811\u79cd\u72b6\u6001\u5206\u6790\u3001\u5f02\u5e38\u6728\u590d\u6838\u6216\u56fe\u8868\u6c47\u603b\u3002",
        "- \u54ea\u4e9b\u95ee\u9898\u4ecd\u7f3a\u5c11\u8bc1\u636e\uff0c\u4e0d\u80fd\u5f62\u6210\u539f\u56e0\u6027\u7ed3\u8bba\u3002",
        "",
    ])
    lines.extend(["## \u4e94\u3001\u8865\u5145\u5efa\u8bae", ""])
    if observed_tasks:
        lines.append("\u5df2\u6709\u73b0\u573a\u8bb0\u5f55\u6458\u8981\uff1a")
        for task in observed_tasks[:12]:
            obs = task.get("field_observation") or {}
            lines.append(f"- {_survey_report_value(task.get('target_name') or task.get('target_id'))}\uff1a{_survey_report_value(obs.get('notes'))}")
        lines.append("")
    else:
        lines.extend([
            "\u5f53\u524d\u73b0\u573a\u8bb0\u5f55\u4ecd\u4e0d\u8db3\uff0c\u62a5\u544a\u5b9a\u4f4d\u4e3a\u8c03\u67e5\u524d\u4efb\u52a1\u8bf4\u660e\u3002\u73b0\u573a\u5e94\u91cd\u70b9\u8bb0\u5f55\u5bf9\u8c61\u662f\u5426\u5b58\u5728\u3001\u5065\u5eb7\u72b6\u6001\u3001\u5f02\u5e38\u73b0\u8c61\u3001\u7167\u7247\u3001\u4f4d\u7f6e\u548c\u5fc5\u8981\u8865\u6d4b\u5b57\u6bb5\u3002",
            "",
        ])
    lines.extend([
        "- \u5bf9\u5df2\u5b8c\u6210\u4efb\u52a1\uff0c\u7ed3\u5408\u5907\u6ce8\u3001\u7167\u7247\u548c\u6570\u636e\u5e93\u6307\u6807\u8fdb\u884c\u590d\u6838\u3002",
        "- \u5bf9\u672a\u5b8c\u6210\u4efb\u52a1\uff0c\u7ee7\u7eed\u4fdd\u7559\u5728\u540e\u7eed\u8c03\u67e5\u6e05\u5355\u4e2d\u3002",
        "- \u8bc1\u636e\u4e0d\u8db3\u65f6\u4e0d\u76f4\u63a5\u5f62\u6210\u539f\u56e0\u5224\u65ad\uff0c\u5e94\u8f6c\u5316\u4e3a\u8865\u6d4b\u5b57\u6bb5\u6216\u590d\u6d4b\u4efb\u52a1\u3002",
        "",
        "*\u62a5\u544a\u7531 ForestryAgent \u91ce\u5916\u8c03\u67e5\u6a21\u5757\u81ea\u52a8\u751f\u6210\u3002*",
    ])
    return "\n".join(lines)

def _normalize_report_formats(formats: _ReportAny = None) -> _ReportList[str]:
    if formats is None:
        return ["md"]
    if isinstance(formats, str):
        raw_items = formats.split(",")
    elif isinstance(formats, (list, tuple, set)):
        raw_items = list(formats)
    else:
        raw_items = [formats]
    normalized = []
    for item in raw_items:
        fmt = str(item or "").strip().lower().lstrip(".")
        if fmt == "markdown":
            fmt = "md"
        if fmt in {"md", "docx", "pdf"} and fmt not in normalized:
            normalized.append(fmt)
    return normalized or ["md"]



def _markdown_plain_text(text: _ReportAny) -> str:
    value = "" if text is None else str(text)
    value = _report_re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", value)
    value = _report_re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", value)
    value = value.replace("**", "").replace("__", "").replace("`", "")
    value = _report_re.sub(r"<[^>]+>", "", value)
    return value.strip()

def _markdown_cells(line: str) -> _ReportList[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_markdown_separator(line: str) -> bool:
    cells = _markdown_cells(line)
    return bool(cells) and all(set(cell.replace(":", "").strip()) <= {"-"} for cell in cells)


def _convert_markdown_to_docx(markdown_path: Path, output_path: Path) -> _ReportDict[str, _ReportAny]:
    try:
        from docx import Document
    except Exception as exc:
        return {"status": "error", "message": f"python-docx is required for docx export: {exc}"}

    doc = Document()
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            index += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(_markdown_plain_text(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(_markdown_plain_text(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(_markdown_plain_text(stripped[2:]), level=1)
        elif stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                if not _is_markdown_separator(lines[index]):
                    table_lines.append(lines[index])
                index += 1
            if table_lines:
                rows = [_markdown_cells(row) for row in table_lines]
                col_count = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=col_count)
                table.style = "Table Grid"
                for row_idx, row in enumerate(rows):
                    for col_idx in range(col_count):
                        table.cell(row_idx, col_idx).text = _markdown_plain_text(row[col_idx]) if col_idx < len(row) else ""
            continue
        elif stripped.startswith(('- ', '* ', '+ ')):
            doc.add_paragraph(_markdown_plain_text(stripped[2:]), style="List Bullet")
        else:
            doc.add_paragraph(_markdown_plain_text(stripped))
        index += 1
    doc.save(str(output_path))
    return {"status": "success", "file": output_path.name}


def _convert_markdown_to_pdf(markdown_path: Path, output_path: Path) -> _ReportDict[str, _ReportAny]:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from xml.sax.saxutils import escape as xml_escape
    except Exception as exc:
        return {"status": "error", "message": f"reportlab is required for pdf export: {exc}"}

    font_name = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception:
        font_name = "Helvetica"

    styles = getSampleStyleSheet()
    normal = ParagraphStyle("SurveyNormal", parent=styles["Normal"], fontName=font_name, fontSize=10, leading=15)
    h1 = ParagraphStyle("SurveyH1", parent=normal, fontSize=18, leading=24, spaceAfter=10)
    h2 = ParagraphStyle("SurveyH2", parent=normal, fontSize=14, leading=20, spaceBefore=8, spaceAfter=6)
    h3 = ParagraphStyle("SurveyH3", parent=normal, fontSize=12, leading=18, spaceBefore=6, spaceAfter=4)

    story = []
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped:
            story.append(Spacer(1, 0.25 * cm))
            index += 1
            continue
        if stripped.startswith("### "):
            story.append(Paragraph(xml_escape(_markdown_plain_text(stripped[4:])), h3))
        elif stripped.startswith("## "):
            story.append(Paragraph(xml_escape(_markdown_plain_text(stripped[3:])), h2))
        elif stripped.startswith("# "):
            story.append(Paragraph(xml_escape(_markdown_plain_text(stripped[2:])), h1))
        elif stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                if not _is_markdown_separator(lines[index]):
                    table_lines.append(lines[index])
                index += 1
            rows = [[Paragraph(xml_escape(_markdown_plain_text(cell)), normal) for cell in _markdown_cells(row)] for row in table_lines]
            if rows:
                table = Table(rows, repeatRows=1)
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(table)
                story.append(Spacer(1, 0.2 * cm))
            continue
        elif stripped.startswith(('- ', '* ', '+ ')):
            story.append(Paragraph(xml_escape(f"? {_markdown_plain_text(stripped[2:])}"), normal))
        else:
            story.append(Paragraph(xml_escape(_markdown_plain_text(stripped)), normal))
        index += 1

    doc = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=1.8*cm, rightMargin=1.8*cm, topMargin=1.6*cm, bottomMargin=1.6*cm)
    doc.build(story)
    return {"status": "success", "file": output_path.name}


def _export_survey_report_files(markdown_path: Path, formats: _ReportAny = None) -> _ReportDict[str, _ReportAny]:
    requested = _normalize_report_formats(formats)
    files = {"md": markdown_path.name}
    errors = {}
    if "docx" in requested:
        docx_path = markdown_path.with_suffix(".docx")
        result = _convert_markdown_to_docx(markdown_path, docx_path)
        if result.get("status") == "success":
            files["docx"] = docx_path.name
        else:
            errors["docx"] = result.get("message") or "docx export failed"
    if "pdf" in requested:
        pdf_path = markdown_path.with_suffix(".pdf")
        result = _convert_markdown_to_pdf(markdown_path, pdf_path)
        if result.get("status") == "success":
            files["pdf"] = pdf_path.name
        else:
            errors["pdf"] = result.get("message") or "pdf export failed"
    return {"requested_formats": requested, "files": files, "export_errors": errors}


def export_existing_report(report_file: str, formats: _ReportAny) -> _ReportDict[str, _ReportAny]:
    safe_name = Path(str(report_file or "")).name
    if not safe_name:
        return {"status": "error", "message": "report_file is required"}
    markdown_path = REPORT_DIR / safe_name
    if markdown_path.suffix.lower() != ".md":
        return {"status": "error", "message": "Only Markdown report files can be exported"}
    if not markdown_path.exists():
        return {"status": "not_found", "message": f"Report file not found: {safe_name}"}
    export_result = _export_survey_report_files(markdown_path, formats=formats)
    files = export_result.get("files", {"md": safe_name})
    requested = export_result.get("requested_formats", [])
    errors = export_result.get("export_errors", {})
    missing = [fmt for fmt in requested if fmt != "md" and not files.get(fmt)]
    status = "success" if not missing and not errors else "error"
    return {
        "status": status,
        "report_file": safe_name,
        "files": files,
        "requested_formats": requested,
        "export_errors": errors,
        "docx_file": files.get("docx"),
        "pdf_file": files.get("pdf"),
        "message": "; ".join(f"{fmt}: {errors.get(fmt, 'export file was not created')}" for fmt in missing) if missing else "",
    }



def _normalize_id_list(values: _ReportAny) -> _ReportList[str]:
    if values is None:
        return []
    if isinstance(values, str):
        raw = [part.strip() for part in values.replace("?", ",").split(",")]
    elif isinstance(values, (list, tuple, set)):
        raw = [str(item).strip() for item in values]
    else:
        raw = [str(values).strip()]
    return [item for item in raw if item]


def _build_comprehensive_report_evidence(plan_ids: _ReportAny = None, subplot_ids: _ReportAny = None) -> _ReportDict[str, _ReportAny]:
    selected_plan_ids = [int(x) for x in _normalize_id_list(plan_ids) if str(x).strip().isdigit()]
    selected_subplot_ids = set(_normalize_id_list(subplot_ids))
    plans = []
    if selected_plan_ids:
        for pid in selected_plan_ids:
            result = get_plan(pid)
            if result.get("status") == "success" and result.get("plan"):
                plans.append(result["plan"])
    else:
        for plan in list_plans(limit=200).get("plans", []):
            result = get_plan(int(plan.get("plan_id")))
            if result.get("status") == "success" and result.get("plan"):
                plans.append(result["plan"])

    compact_plans = []
    all_tasks = []
    for plan in plans:
        recs = plan.get("recommendations") or []
        if selected_subplot_ids:
            recs = [rec for rec in recs if str(rec.get("subplot_id") or "") in selected_subplot_ids]
        if not recs and selected_subplot_ids:
            continue
        observations_result = get_plan_observations(int(plan.get("plan_id")))
        observations = observations_result.get("observations", []) if isinstance(observations_result, dict) else []
        obs_by_rec = {obs.get("rec_id"): obs for obs in observations if obs.get("rec_id") is not None}
        compact_recs = []
        for rec in recs:
            obs = obs_by_rec.get(rec.get("rec_id"))
            compact = {
                "plan_id": plan.get("plan_id"),
                "rec_id": rec.get("rec_id"),
                "target_type": rec.get("target_type"),
                "target_id": rec.get("target_id"),
                "target_name": rec.get("target_name"),
                "subplot_id": rec.get("subplot_id"),
                "species": rec.get("species"),
                "priority": rec.get("priority"),
                "category": rec.get("category"),
                "reason": rec.get("reason"),
                "suggested_action": rec.get("suggested_action"),
                "status": rec.get("status"),
                "field_observation": obs,
            }
            compact_recs.append(compact)
            all_tasks.append(compact)
        compact_plans.append({
            "plan_id": plan.get("plan_id"),
            "title": plan.get("title"),
            "user_request": plan.get("user_request"),
            "ai_analysis": plan.get("ai_analysis"),
            "status": plan.get("status"),
            "created_at": plan.get("created_at"),
            "task_count": len(compact_recs),
        })

    subplot_set = sorted({str(task.get("subplot_id")) for task in all_tasks if task.get("subplot_id")})
    species_set = sorted({str(task.get("species")) for task in all_tasks if task.get("species")})
    completed = [task for task in all_tasks if str(task.get("status") or "").lower() in {"done", "completed", "finished", "\u5df2\u5b8c\u6210"} or task.get("field_observation")]
    priority_counts: _ReportDict[str, int] = {}
    category_counts: _ReportDict[str, int] = {}
    for task in all_tasks:
        priority = _survey_report_value(task.get("priority"), "\u672a\u5206\u7ea7")
        category = _survey_report_value(task.get("category"), "\u672a\u5206\u7c7b")
        priority_counts[priority] = priority_counts.get(priority, 0) + 1
        category_counts[category] = category_counts.get(category, 0) + 1
    return {
        "scope": {
            "selected_plan_ids": selected_plan_ids,
            "selected_subplot_ids": sorted(selected_subplot_ids),
            "actual_plan_count": len(compact_plans),
            "actual_subplot_count": len(subplot_set),
            "actual_species_count": len(species_set),
        },
        "plans": compact_plans,
        "tasks": all_tasks[:300],
        "summary_stats": {
            "total_tasks": len(all_tasks),
            "completed_tasks": len(completed),
            "pending_tasks": max(len(all_tasks) - len(completed), 0),
            "field_observation_count": len([task for task in all_tasks if task.get("field_observation")]),
            "subplot_ids": subplot_set[:100],
            "species": species_set[:100],
            "priority_counts": priority_counts,
            "category_counts": category_counts,
        },
    }


def _generate_comprehensive_report_fallback(evidence: _ReportDict[str, _ReportAny], audience: str = "leader") -> str:
    audience = _normalize_report_audience(audience)
    stats = evidence.get("summary_stats") or {}
    scope = evidence.get("scope") or {}
    plans = evidence.get("plans") or []
    tasks = evidence.get("tasks") or []
    title = "\u9636\u6bb5\u6027\u91ce\u5916\u8c03\u67e5\u7efc\u5408\u7b80\u62a5" if audience == "leader" else "\u9636\u6bb5\u6027\u91ce\u5916\u8c03\u67e5\u7efc\u5408\u62a5\u544a"
    lines = [
        f"# {title}",
        "",
        f"**\u751f\u6210\u65f6\u95f4**?{_report_datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"**\u7eb3\u5165\u65b9\u6848**?{scope.get('actual_plan_count', 0)} ?",
        f"**\u6d89\u53ca\u6837\u65b9**?{scope.get('actual_subplot_count', 0)} ?",
        "",
        "## \u4e00\u3001\u603b\u4f53\u6982\u51b5",
        "",
        f"\u672c\u62a5\u544a\u6574\u5408\u6240\u9009\u8c03\u67e5\u65b9\u6848\u548c\u6837\u65b9\u8bb0\u5f55\u751f\u6210\uff0c\u5171\u7eb3\u5165 {stats.get('total_tasks', 0)} \u9879\u8c03\u67e5\u4efb\u52a1\uff0c\u5df2\u5b8c\u6210 {stats.get('completed_tasks', 0)} \u9879\uff0c\u5f85\u5b8c\u6210 {stats.get('pending_tasks', 0)} \u9879\uff0c\u73b0\u573a\u8bb0\u5f55 {stats.get('field_observation_count', 0)} \u6761\u3002",
        "",
        "| \u6307\u6807 | \u6570\u503c |",
        "|---|---:|",
        f"| \u65b9\u6848\u6570 | {scope.get('actual_plan_count', 0)} |",
        f"| \u6837\u65b9\u6570 | {scope.get('actual_subplot_count', 0)} |",
        f"| \u6811\u79cd\u6570 | {scope.get('actual_species_count', 0)} |",
        f"| \u4efb\u52a1\u6570 | {stats.get('total_tasks', 0)} |",
        f"| \u5df2\u5b8c\u6210 | {stats.get('completed_tasks', 0)} |",
        f"| \u73b0\u573a\u8bb0\u5f55 | {stats.get('field_observation_count', 0)} |",
        "",
        "## \u4e8c\u3001\u4e3b\u8981\u53d1\u73b0??",
        "",
    ]
    cats = stats.get("category_counts") or {}
    if cats:
        for key, value in sorted(cats.items(), key=lambda item: item[1], reverse=True)[:8]:
            lines.append(f"- {key}?{value} ??")
    else:
        lines.append("- \u5f53\u524d\u6ca1\u6709\u8db3\u591f\u7684\u5206\u7c7b\u4fe1\u606f\u5f62\u6210\u7a33\u5b9a\u53d1\u73b0\u3002")
    lines.extend(["", "## \u4e09\u3001\u65b9\u6848\u6765\u6e90??", "", "| \u65b9\u6848 | \u7528\u6237\u9700\u6c42 | \u4efb\u52a1\u6570 |", "|---|---|---:|"])
    for plan in plans[:30]:
        lines.append(f"| {_survey_report_value(plan.get('title'))} | {_survey_report_value(plan.get('user_request'))} | {plan.get('task_count', 0)} |")
    lines.extend(["", "## \u56db\u3001\u73b0\u573a\u9a8c\u8bc1??", ""])
    observed = [task for task in tasks if task.get("field_observation")]
    if observed:
        for task in observed[:15]:
            obs = task.get("field_observation") or {}
            lines.append(f"- {_survey_report_value(task.get('target_name') or task.get('target_id'))}?{_survey_report_value(obs.get('notes'))}")
    else:
        lines.append("- \u5f53\u524d\u73b0\u573a\u8bb0\u5f55\u8f83\u5c11\uff0c\u7efc\u5408\u62a5\u544a\u4e3b\u8981\u53cd\u6620\u8c03\u67e5\u8ba1\u5212\u548c\u5f85\u9a8c\u8bc1\u95ee\u9898\uff0c\u4e0d\u80fd\u76f4\u63a5\u5f62\u6210\u6700\u7ec8\u73b0\u573a\u7ed3\u8bba\u3002")
    lines.extend([
        "",
        "## \u4e94\u3001\u9636\u6bb5\u5224\u65ad??",
        "",
        "- \u6240\u9009\u65b9\u6848\u53ef\u4ee5\u5171\u540c\u56de\u7b54\uff1a\u54ea\u4e9b\u6837\u65b9\u6216\u5bf9\u8c61\u88ab\u53cd\u590d\u7eb3\u5165\u8c03\u67e5\uff0c\u54ea\u4e9b\u95ee\u9898\u7c7b\u578b\u9700\u8981\u4f18\u5148\u6838\u67e5\u3002",
        "- \u5b8c\u6210\u73b0\u573a\u8bb0\u5f55\u540e\uff0c\u53ef\u4ee5\u8fdb\u4e00\u6b65\u5224\u65ad\u63a8\u8350\u5bf9\u8c61\u662f\u5426\u771f\u5b9e\u3001\u72b6\u6001\u662f\u5426\u4e00\u81f4\u3001\u662f\u5426\u9700\u8981\u8865\u6d4b\u6216\u79fb\u9664\u3002",
        "- \u5f53\u524d\u62a5\u544a\u4e0d\u76f4\u63a5\u8bc1\u660e\u539f\u56e0\u673a\u5236\u6216\u957f\u671f\u8d8b\u52bf\uff0c\u53ea\u4f5c\u4e3a\u9636\u6bb5\u6027\u8c03\u67e5\u7ec4\u7ec7\u548c\u8bc1\u636e\u6c47\u603b\u3002",
        "",
        "## \u516d\u3001\u540e\u7eed\u5b89\u6392??",
        "",
        "- \u4f18\u5148\u8865\u9f50\u5f85\u5b8c\u6210\u4efb\u52a1\u7684\u73b0\u573a\u8bb0\u5f55\u3001\u7167\u7247\u3001\u4f4d\u7f6e\u548c\u5fc5\u8981\u6d4b\u91cf\u5b57\u6bb5\u3002",
        "- \u5bf9\u591a\u65b9\u6848\u91cd\u590d\u51fa\u73b0\u7684\u6837\u65b9\u6216\u5bf9\u8c61\uff0c\u5efa\u8bae\u4f5c\u4e3a\u4e0b\u4e00\u8f6e\u590d\u6838\u91cd\u70b9\u3002",
        "- \u8c03\u67e5\u7ed3\u675f\u540e\u91cd\u65b0\u751f\u6210\u7efc\u5408\u62a5\u544a\uff0c\u7528\u73b0\u573a\u7ed3\u679c\u66f4\u65b0\u9636\u6bb5\u5224\u65ad\u3002",
    ])
    return "\n".join(lines)


def _generate_comprehensive_report_with_agent(evidence: _ReportDict[str, _ReportAny], audience: str = "leader") -> _ReportOptional[str]:
    _set_survey_agent_report_error("")
    try:
        from agent import run_agent_chat
    except Exception as exc:
        _set_survey_agent_report_error(f"cannot import agent.run_agent_chat: {exc}")
        return None
    audience = _normalize_report_audience(audience)
    audience_name = _report_audience_cn(audience)
    payload = _report_json.dumps(evidence, ensure_ascii=False, default=str)
    prompt = f"""
You are generating an integrated field survey report from selected survey plans and selected subplots.
Audience: {audience_name}

The report must answer:
1. Why these plans/subplots are integrated.
2. What common issues or repeated survey targets appear.
3. What field observations already support.
4. What can be concluded after completion.
5. What cannot be concluded yet.

Use Markdown only, Simplified Chinese, no code blocks. Do not invent facts.
For leader audience, be concise and decision-oriented. For technical audience, include more task evidence.

Evidence JSON:
```json
{payload}
```
""".strip()
    try:
        result = run_agent_chat(
            question=prompt,
            session_id=f"survey_comprehensive_report_{int(time.time())}_{audience}",
            client_id="survey_report_generator",
            context={"current_page": "survey_comprehensive_report", "report_mode": True, "report_audience": audience, "context_policy": "auto"},
            options={"max_tool_rounds": 0, "history_limit": 0},
        )
    except Exception as exc:
        _set_survey_agent_report_error(f"run_agent_chat failed: {exc}")
        return None
    answer = result.get("answer") if isinstance(result, dict) else str(result or "")
    answer = (answer or "").strip()
    if len(answer) < 120 or "#" not in answer:
        _set_survey_agent_report_error("agent returned invalid comprehensive report")
        return None
    return answer


def generate_comprehensive_report(plan_ids: _ReportAny = None, subplot_ids: _ReportAny = None, formats: _ReportAny = None, mode: str = "agent", allow_fallback: bool = True, audience: str = "leader") -> _ReportDict[str, _ReportAny]:
    audience = _normalize_report_audience(audience)
    mode = str(mode or "agent").strip().lower()
    evidence = _build_comprehensive_report_evidence(plan_ids=plan_ids, subplot_ids=subplot_ids)
    if not evidence.get("tasks"):
        return {"status": "not_found", "message": "\u672a\u627e\u5230\u53ef\u7eb3\u5165\u7efc\u5408\u62a5\u544a\u7684\u8c03\u67e5\u4efb\u52a1\uff0c\u8bf7\u9009\u62e9\u81f3\u5c11\u4e00\u4e2a\u6709\u4efb\u52a1\u7684\u65b9\u6848\u6216\u6837\u65b9\u3002"}
    report_text = ""
    report_mode = mode
    if mode in {"agent", "auto"}:
        report_mode = "agent"
        report_text = _generate_comprehensive_report_with_agent(evidence, audience=audience) or ""
    if not report_text and (mode == "template" or allow_fallback):
        report_mode = "template_fallback" if mode != "template" else "template"
        report_text = _generate_comprehensive_report_fallback(evidence, audience=audience)
    if not report_text:
        return {"status": "error", "message": "\u7efc\u5408\u62a5\u544a\u751f\u6210\u5931\u8d25\u3002"}
    safe_title = f"comprehensive_{audience}_{_report_datetime.now().strftime('%Y%m%d_%H%M%S')}"
    filename = f"survey_report_{safe_title}.md"
    path = REPORT_DIR / filename
    path.write_text(report_text, encoding="utf-8")
    export_result = _export_survey_report_files(path, formats=formats)
    return {
        "status": "success",
        "report": report_text,
        "report_file": filename,
        "report_mode": report_mode,
        "report_audience": audience,
        "files": export_result.get("files", {"md": filename}),
        "export_errors": export_result.get("export_errors", {}),
        "stats": evidence.get("summary_stats") or {},
        "scope": evidence.get("scope") or {},
    }

def generate_report(plan_id: int, formats: _ReportAny = None, mode: str = "agent", allow_fallback: bool = False, audience: str = "technical") -> _ReportDict[str, _ReportAny]:
    result = _get_plan_full(plan_id)
    if not result:
        return {"status": "not_found", "message": f"\u672a\u627e\u5230\u8c03\u67e5\u65b9\u6848 {plan_id}"}

    plan = result["plan"]
    recommendations = plan.pop("recommendations", [])
    observations_result = get_plan_observations(plan_id)
    observations = observations_result.get("observations", []) if isinstance(observations_result, dict) else []

    evidence_package = _build_survey_report_evidence_package(
        plan_id=plan_id,
        plan=plan,
        recommendations=recommendations,
        observations=observations,
    )

    audience = _normalize_report_audience(audience)
    mode = str(mode or "agent").strip().lower()
    if mode not in {"agent", "template", "auto"}:
        return {"status": "error", "message": f"\u4e0d\u652f\u6301\u7684\u62a5\u544a\u751f\u6210\u6a21\u5f0f: {mode}"}

    report_mode = mode
    report_text = ""
    if mode in {"agent", "auto"}:
        report_mode = "agent"
        report_text = _generate_survey_report_with_agent(evidence_package, audience=audience) or ""

    if not report_text and mode == "template":
        report_mode = "template"
        report_text = _generate_survey_report_fallback(evidence_package, audience=audience)

    if not report_text and mode == "auto" and allow_fallback:
        report_mode = "template_fallback"
        report_text = _generate_survey_report_fallback(evidence_package, audience=audience)

    if not report_text:
        return {
            "status": "error",
            "plan_id": plan_id,
            "report_mode": "agent_failed" if mode in {"agent", "auto"} else mode,
            "message": "\u667a\u80fd\u4f53\u62a5\u544a\u751f\u6210\u5931\u8d25\uff0c\u672a\u81ea\u52a8\u4f7f\u7528\u6a21\u677f\u62fc\u63a5\u3002\u4f60\u53ef\u4ee5\u91cd\u8bd5\uff0c\u6216\u9009\u62e9\u6a21\u677f\u62a5\u544a\u3002",
            "stats": evidence_package.get("execution_stats") or {},
        }

    safe_title = _survey_report_safe_title(plan.get("title") or f"plan_{plan_id}")
    filename = f"survey_report_{plan_id}_{audience}_{safe_title}.md"
    path = REPORT_DIR / filename
    path.write_text(report_text, encoding="utf-8")
    export_result = _export_survey_report_files(path, formats=formats)

    stats = evidence_package.get("execution_stats") or {}
    with _get_conn() as conn:
        conn.execute(
            """
            UPDATE field_survey_plans
            SET summary = ?,
                latest_report_text = ?,
                latest_report_file = ?,
                latest_report_mode = ?,
                latest_report_generated_at = ?,
                updated_at = ?
            WHERE plan_id = ?
            """,
            (
                f"\u62a5\u544a\u5df2\u751f\u6210\uff1a{filename}\uff1b\u751f\u6210\u6a21\u5f0f\uff1a{report_mode}\uff1b\u73b0\u573a\u8bb0\u5f55 {stats.get('field_observation_count', 0)} \u6761\u3002",
                report_text,
                filename,
                report_mode,
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                plan_id,
            ),
        )
        conn.commit()

    return {
        "status": "success",
        "plan_id": plan_id,
        "report": report_text,
        "report_file": filename,
        "report_mode": report_mode,
        "report_audience": audience,
        "files": export_result.get("files", {"md": filename}),
        "requested_formats": export_result.get("requested_formats", ["md"]),
        "export_errors": export_result.get("export_errors", {}),
        "docx_file": export_result.get("files", {}).get("docx"),
        "pdf_file": export_result.get("files", {}).get("pdf"),
        "stats": stats,
    }

