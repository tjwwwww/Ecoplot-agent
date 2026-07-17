# -*- coding: utf-8 -*-
"""
agent.py
========
统一林业智能分析核心入口 —— 并行感知 Neuro-Symbolic ReAct Agent 

【核心演进与灵活设计】
1. 并行感知架构 (Parallel-Aware ReAct):
   - 大模型单次响应中返回多个 tool_calls，使用 ThreadPoolExecutor 并行并发执行所有工具。
   - 大幅减少网络往返耗时与推理轮次，从旧版的线性 30+ 轮耗时骤降至 3~5 轮。
2. 动态本体与元数据契约检索 (tool_inspect_ontology_schema):
   - 支持动态核查 `qilian_ontology.yaml` 和 `forestry_knowledge_registry.yaml` 声明。
   - 彻底避免大模型在生成 Cypher 选择节点标签时的幻觉。
3. 祁连山国家公园“零人为干预”最高保护公理:
   - 所有结论必须区分观测事实、确定性计算、统计关联和待验证解释。
"""

import os
import sys
import io
import json
import uuid
import sqlite3
import re
import time
import yaml
from pathlib import Path
from datetime import datetime, timezone
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Callable, Dict, List, Optional
from dotenv import load_dotenv
from neo4j import GraphDatabase
from provider import chat_with_tools
import hashlib

# 修复 Windows 控制台中文输出编码
try:
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stdout, 'buffer') and not isinstance(sys.stdout, io.TextIOWrapper):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace', line_buffering=True)
    if hasattr(sys.stderr, 'buffer') and not isinstance(sys.stderr, io.TextIOWrapper):
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
except Exception:
    pass

from forestry_spatial_tools import (
    tool_calc_stand_structure_metrics,
    tool_calc_tree_morphology_metrics,
    tool_calc_species_diversity_metrics,
    tool_calc_volume_metrics,
    tool_calc_deadwood_metrics,
    tool_calc_shrub_metrics,
    tool_calc_hegyi_competition,
    tool_get_tree_topography_context,
    tool_get_subplot_topography_summary,
    tool_get_climate_background_summary,
    tool_compute_registered_indicators,
    FORESTRY_SPATIAL_SCHEMAS,
)
from forest_intelligence_core import (
    tool_resolve_forest_question,
    tool_inspect_forest_data,
    tool_run_forest_analysis_protocol,
    FOREST_INTELLIGENCE_SCHEMAS,
)
try:
    from forestry_visualization_engine import (
        plot_subplot_grid_heatmap,
        plot_size_class_distribution,
        plot_species_composition,
        plot_tree_relationship_scatter,
        plot_group_comparison_boxplot,
        plot_tree_spatial_map,
        plot_subplot_percentile_profile,
        plot_climate_time_series,
        create_generic_chart,
    )
    _FORESTRY_VIS_ENGINE_AVAILABLE = True
    _FORESTRY_VIS_ENGINE_IMPORT_ERROR = None
except Exception as e:
    plot_subplot_grid_heatmap = None
    plot_size_class_distribution = None
    plot_species_composition = None
    plot_tree_relationship_scatter = None
    plot_group_comparison_boxplot = None
    plot_tree_spatial_map = None
    plot_subplot_percentile_profile = None
    plot_climate_time_series = None
    create_generic_chart = None
    _FORESTRY_VIS_ENGINE_AVAILABLE = False
    _FORESTRY_VIS_ENGINE_IMPORT_ERROR = str(e)
from formula_execution_engine import NeuroSymbolicFormulaEngine
from interpretive_rule_engine import tool_retrieve_interpretive_rules


def _visualization_engine_error_response(tool_name: str) -> str:
    return json.dumps({
        "error": (
            f"可视化引擎不可用：{tool_name} 无法执行。"
            f" forestry_visualization_engine 模块导入失败，原因：{_FORESTRY_VIS_ENGINE_IMPORT_ERROR}。"
            " 请安装 plotly、kaleido 等依赖后重试。"
        )
    }, ensure_ascii=False)

# 1. 兼容性加载环境变量 (基于当前文件所在目录，摆脱硬编码)
base_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(base_dir, ".env"))

URI = os.getenv("NEO4J_URI", "bolt://localhost:7687")
USER_DB = os.getenv("NEO4J_USER", "neo4j")
PASSWORD = os.getenv("NEO4J_PASSWORD", "1820401753")

_driver = GraphDatabase.driver(URI, auth=(USER_DB, PASSWORD))
_NS_ENGINE_INIT_ERROR = None
try:
    _ns_engine = NeuroSymbolicFormulaEngine()
except Exception as exc:
    _ns_engine = None
    _NS_ENGINE_INIT_ERROR = str(exc)
DATA_DIR = Path(base_dir) / "data"
DATA_DIR.mkdir(parents=True, exist_ok=True)
SESSION_DB_PATH = Path(os.getenv("FORESTRY_AGENT_SESSION_DB", str(DATA_DIR / "agent_sessions.db")))
EventCallback = Optional[Callable[[str, Dict[str, Any]], None]]


CHAT_SYSTEM_PROMPT = (
    "你是“祁连山大样地野外调查专家智能体”。\n\n"
    "你的任务是面向样地野外调查、数据核查、林分分析和现场复测，理解用户自然语言问题，并自主决定是否需要调用工具、调用什么工具以及调用顺序。\n\n"
    "通用工作法：\n"
    "1. 先将用户问题落地为业务对象、指标/因子、范围、操作和期望输出。\n"
    "2. 本体、知识注册表、数据契约和数据字典用于语义理解与物理字段对齐；不要把语义概念直接当作物理数据库结构。\n"
    "3. 涉及真实观测值、数量、分布、排行、样方、单木、树种表现的问题，先用数据契约判断物理来源，再优先使用事实数据源和确定性计算工具。\n"
    "4. 涉及概念解释、指标定义、公式依赖、字段含义时，优先使用本体、知识注册表或数据字典。\n"
    "5. 涉及关系追溯、证据链、知识来源时，再使用知识图谱；使用 KG 前应确认 data_contract 中声明了对应标签、属性和关系路径。\n"
    "6. 工具未命中不等于事实不存在；需要判断是否为字段、参数、数据源或工具选择问题，并在必要时换路径验证。\n"
    "7. 最终结论必须受证据等级约束：观测值和确定性计算可作为正式结果；统计分析表达分布、差异、关联或排序；未经验证的模型不得输出风险、因果或处方性结论。\n"
    "8. 数据不足时，可以降级回答、说明缺失信息，并给出现场复核或补采建议。\n\n"
    "自主性要求：\n"
    "- 不要默认生成报告，也不要默认调用工具；简单概念和方法问题可直接回答。\n"
    "- 对需要数据、图表或对象级分析的问题，应自主组合可用工具完成查询、计算、对比和结果组织。\n"
    "- 涉及已注册指标计算时，优先调用 tool_compute_registered_indicators；旧的 tool_calc_* 指标工具仅作为兼容接口，不作为默认规划路径。\n"
    "- 页面上下文和历史焦点只是辅助信息，不是硬性限制；如果用户本轮指定了新对象或新范围，以本轮问题为准。\n\n"
    "回答风格：\n"
    "- 默认简洁、直接、适合中文网页对话。\n"
    "- 回答应包含关键依据；如果生成了图表或文件，要说明产物。\n"
    "- 用户明确要求报告、汇报或导出时，再输出完整结构化报告。"
).strip()

def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")


def create_session_id() -> str:
    return "sess_" + uuid.uuid4().hex[:16]


def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def _connect_session_db() -> sqlite3.Connection:
    _ensure_parent(SESSION_DB_PATH)
    conn = sqlite3.connect(SESSION_DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS agent_sessions (
            session_id TEXT PRIMARY KEY,
            client_id TEXT,
            created_at TEXT,
            updated_at TEXT,
            title TEXT,
            last_focus_json TEXT
        )
        """
    )
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS agent_messages (
            message_id TEXT PRIMARY KEY,
            session_id TEXT,
            role TEXT,
            content TEXT,
            context_json TEXT,
            answer_type TEXT,
            tool_calls_json TEXT,
            artifacts_json TEXT,
            created_at TEXT
        )
        """
    )
    session_columns = {row["name"] for row in conn.execute("PRAGMA table_info(agent_sessions)").fetchall()}
    if "client_id" not in session_columns:
        conn.execute("ALTER TABLE agent_sessions ADD COLUMN client_id TEXT")
    conn.commit()
    return conn


def ensure_session(session_id: str, client_id: Optional[str] = None) -> None:
    with _connect_session_db() as conn:
        row = conn.execute("SELECT session_id, client_id FROM agent_sessions WHERE session_id=?", (session_id,)).fetchone()
        if row is None:
            conn.execute(
                "INSERT INTO agent_sessions(session_id, client_id, created_at, updated_at, title, last_focus_json) VALUES(?,?,?,?,?,?)",
                (session_id, client_id, now_iso(), now_iso(), "ForestryAgent Chat", "{}"),
            )
        else:
            if client_id and not row["client_id"]:
                conn.execute(
                    "UPDATE agent_sessions SET client_id=?, updated_at=? WHERE session_id=?",
                    (client_id, now_iso(), session_id),
                )
            else:
                conn.execute("UPDATE agent_sessions SET updated_at=? WHERE session_id=?", (now_iso(), session_id))
        conn.commit()


def load_last_focus(session_id: str) -> Dict[str, Any]:
    with _connect_session_db() as conn:
        row = conn.execute("SELECT last_focus_json FROM agent_sessions WHERE session_id=?", (session_id,)).fetchone()
    if not row:
        return {}
    try:
        data = json.loads(row["last_focus_json"] or "{}")
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def save_last_focus(session_id: str, focus: Dict[str, Any]) -> None:
    with _connect_session_db() as conn:
        conn.execute(
            "UPDATE agent_sessions SET last_focus_json=?, updated_at=? WHERE session_id=?",
            (json.dumps(focus, ensure_ascii=False, default=str), now_iso(), session_id),
        )
        conn.commit()


def load_recent_messages(session_id: str, limit: int = 12) -> List[Dict[str, Any]]:
    with _connect_session_db() as conn:
        rows = conn.execute(
            "SELECT role, content FROM agent_messages WHERE session_id=? ORDER BY created_at DESC LIMIT ?",
            (session_id, int(limit)),
        ).fetchall()
    items = [{"role": r["role"], "content": r["content"]} for r in reversed(rows)]
    return [item for item in items if item["role"] in {"user", "assistant"} and item["content"]]


def save_message(
    session_id: str,
    role: str,
    content: str,
    context: Optional[Dict[str, Any]] = None,
    answer_type: Optional[str] = None,
    tool_calls: Optional[List[Dict[str, Any]]] = None,
    artifacts: Optional[List[Dict[str, Any]]] = None,
) -> None:
    with _connect_session_db() as conn:
        conn.execute(
            """
            INSERT INTO agent_messages(message_id, session_id, role, content, context_json, answer_type, tool_calls_json, artifacts_json, created_at)
            VALUES(?,?,?,?,?,?,?,?,?)
            """,
            (
                "msg_" + uuid.uuid4().hex[:20],
                session_id,
                role,
                content,
                json.dumps(context or {}, ensure_ascii=False, default=str),
                answer_type,
                json.dumps(tool_calls or [], ensure_ascii=False, default=str),
                json.dumps(artifacts or [], ensure_ascii=False, default=str),
                now_iso(),
            ),
        )
        conn.execute("UPDATE agent_sessions SET updated_at=? WHERE session_id=?", (now_iso(), session_id))
        conn.commit()


def list_chat_sessions(limit: int = 50, client_id: Optional[str] = None) -> List[Dict[str, Any]]:
    with _connect_session_db() as conn:
        sql = """
        SELECT
            s.session_id,
            s.client_id,
            s.created_at,
            s.updated_at,
            s.title,
            s.last_focus_json,
            COUNT(m.message_id) AS message_count,
            (
                SELECT m2.content
                FROM agent_messages m2
                WHERE m2.session_id = s.session_id AND m2.role = 'user'
                ORDER BY m2.created_at ASC
                LIMIT 1
            ) AS first_user_message,
            (
                SELECT m3.content
                FROM agent_messages m3
                WHERE m3.session_id = s.session_id AND m3.role = 'user'
                ORDER BY m3.created_at DESC
                LIMIT 1
            ) AS last_user_message
        FROM agent_sessions s
        LEFT JOIN agent_messages m ON m.session_id = s.session_id
        """
        params: List[Any] = []
        if client_id:
            sql += " WHERE s.client_id = ? "
            params.append(client_id)
        sql += """
        GROUP BY s.session_id
        ORDER BY s.updated_at DESC
        LIMIT ?
        """
        params.append(int(limit))
        rows = conn.execute(sql, tuple(params)).fetchall()

    sessions: List[Dict[str, Any]] = []
    for row in rows:
        title = (row["first_user_message"] or row["title"] or "新对话").strip()
        sessions.append({
            "session_id": row["session_id"],
            "client_id": row["client_id"],
            "title": title[:40],
            "created_at": row["created_at"],
            "updated_at": row["updated_at"],
            "message_count": int(row["message_count"] or 0),
            "last_user_message": row["last_user_message"] or "",
            "last_focus": _safe_json_loads(row["last_focus_json"] or "{}") or {},
        })
    return sessions


def load_session_messages(session_id: str, limit: int = 200) -> List[Dict[str, Any]]:
    with _connect_session_db() as conn:
        rows = conn.execute(
            """
            SELECT role, content, context_json, answer_type, tool_calls_json, artifacts_json, created_at
            FROM agent_messages
            WHERE session_id=?
            ORDER BY created_at ASC
            LIMIT ?
            """,
            (session_id, int(limit)),
        ).fetchall()

    messages: List[Dict[str, Any]] = []
    for row in rows:
        messages.append({
            "role": row["role"],
            "content": row["content"] or "",
            "created_at": row["created_at"],
            "answer_type": row["answer_type"],
            "context": _safe_json_loads(row["context_json"] or "{}") or {},
            "tool_calls": _safe_json_loads(row["tool_calls_json"] or "[]") or [],
            "artifacts": _safe_json_loads(row["artifacts_json"] or "[]") or [],
        })
    return messages


def load_last_user_turn(session_id: str) -> Optional[Dict[str, Any]]:
    with _connect_session_db() as conn:
        row = conn.execute(
            """
            SELECT message_id, content, context_json, created_at
            FROM agent_messages
            WHERE session_id=? AND role='user'
            ORDER BY created_at DESC
            LIMIT 1
            """,
            (session_id,),
        ).fetchone()
    if not row:
        return None
    try:
        context = json.loads(row["context_json"] or "{}")
    except Exception:
        context = {}
    if not isinstance(context, dict):
        context = {}
    return {
        "message_id": row["message_id"],
        "question": row["content"] or "",
        "context": context,
        "created_at": row["created_at"],
    }


def delete_messages_from_time(session_id: str, created_at: str) -> int:
    with _connect_session_db() as conn:
        cur = conn.execute(
            "DELETE FROM agent_messages WHERE session_id=? AND created_at>=?",
            (session_id, created_at),
        )
        conn.execute("UPDATE agent_sessions SET updated_at=? WHERE session_id=?", (now_iso(), session_id))
        conn.commit()
        return int(cur.rowcount or 0)


# ==============================================================================
# 工具 1: 动态本体契约检索器 (增加灵活性，规避硬编码)
# ==============================================================================
def tool_inspect_ontology_schema(query_concept: str) -> str:
    """
    动态概念本体与声明式规则规范检索器。
    查询本地《qilian_ontology.yaml》与《forestry_knowledge_registry.yaml》，
    精准提取类、属性、SQLite 字段映射以及神经符号规则契约切片，避免智能体发生属性猜测或标签幻觉。
    """
    concept = str(query_concept).strip().lower()
    results = {"query_concept": query_concept, "matched_schema_slices": []}
    
    # 1. 检索类定义与关系契约 (qilian_ontology.yaml)
    ont_path = os.path.join(os.path.dirname(__file__), "ontology", "qilian_ontology.yaml")
    if os.path.exists(ont_path):
        try:
            with open(ont_path, "r", encoding="utf-8") as f:
                ont_data = yaml.safe_load(f) or {}
            for cls in ont_data.get("classes", []):
                cls_id = str(cls.get("class_id", "")).lower()
                cls_lbl = str(cls.get("label", "")).lower()
                cls_desc = str(cls.get("description", "")).lower()
                if concept in cls_id or concept in cls_lbl or concept in cls_desc:
                    results["matched_schema_slices"].append({"type": "OntologyClass", "schema": cls})
            for field in ont_data.get("canonical_tree_observation_fields", []):
                if concept in str(field.get("target", "")).lower() or any(concept in str(ex).lower() for ex in field.get("source_examples", [])):
                    results["matched_schema_slices"].append({"type": "SQLiteColumnMapping", "mapping": field})
        except Exception as e:
            results["ontology_error"] = str(e)

    # 2. 检索声明式计算规则与公式契约库 (forestry_knowledge_registry.yaml)
    reg_path = os.path.join(os.path.dirname(__file__), "ontology", "forestry_knowledge_registry.yaml")
    if os.path.exists(reg_path):
        try:
            with open(reg_path, "r", encoding="utf-8") as f:
                reg_data = yaml.safe_load(f) or {}
            for key in ["variables", "indicators", "formulas", "models", "diagnostic_rules"]:
                for item in reg_data.get(key, []):
                    item_str = json.dumps(item, ensure_ascii=False).lower()
                    item_id = str(item.get("variable_id", item.get("indicator_id", item.get("knowledge_id", "")))).lower()
                    item_name = str(item.get("name_cn", "")).lower()
                    if concept in item_str or concept in item_id or concept in item_name:
                        results["matched_schema_slices"].append({
                            "type": f"Registry{key.capitalize()[:-1] if key.endswith('s') else key.capitalize()}",
                            "contract": item
                        })
        except Exception as e:
            results["registry_error"] = str(e)

    # 3. 拦截林业保护政策规则校验
    if any(k in concept for k in ["间伐", "疏伐", "抚育", "采伐", "大径材", "经营"]):
        results["matched_schema_slices"].append({
            "type": "NationalParkPolicyAxiom",
            "rule_id": "ZERO_INTERVENTION_LAW",
            "principle": "零人为干预最高保护原则 (Zero-thinning Ecological Sanctuary Axiom)",
            "description": "祁连山国家公园严格执行生态休养法保护规范。绝对不可在此样区推荐或规划间伐、采伐等经济获利行为。高竞争与被压木一律解释为天然林自疏演替规律；监测行动应当以无损应力波、激光雷达或镜检替代物理采伐操作。"
        })

    if not results["matched_schema_slices"]:
        results["hint"] = "未检索到精确匹配的关联概念，建议查询基类：MonitoringPlot, Subplot, TreeIndividual, Taxon, DiagnosticRule, VolumeTableParameter"

    return json.dumps(results, ensure_ascii=False, indent=2)


def tool_inspect_data_contract(query_concept: str) -> str:
    """
    检索语义概念到物理数据源的字段契约。
    用于回答“这个概念应去 SQLite、KG 还是工具里找”的问题，避免把本体属性误当作 Neo4j 属性。
    """
    concept = str(query_concept or "").strip().lower()
    contract_path = os.path.join(os.path.dirname(__file__), "ontology", "data_contract.yaml")
    results = {
        "query_concept": query_concept,
        "contract_file": contract_path,
        "matched_contracts": [],
        "routing_hint": None,
    }
    if not os.path.exists(contract_path):
        results["error"] = "data_contract.yaml 不存在"
        return json.dumps(results, ensure_ascii=False, indent=2)
    try:
        with open(contract_path, "r", encoding="utf-8") as f:
            contract = yaml.safe_load(f) or {}
        results["principles"] = contract.get("principles", [])
        for section_name in ["entities", "fields", "common_query_patterns"]:
            section = contract.get(section_name, {}) or {}
            for key, item in section.items():
                item_text = json.dumps(item, ensure_ascii=False).lower()
                key_text = str(key).lower()
                if not concept or concept in key_text or concept in item_text:
                    results["matched_contracts"].append({
                        "section": section_name,
                        "key": key,
                        "contract": item,
                    })
        if any(word in str(query_concept) for word in ["哪个", "最高", "最低", "最多", "排行", "分布", "平均", "数量", "多少"]):
            results["routing_hint"] = "这是事实统计/排序问题，优先使用 SQLite 事实表或确定性工具；不要优先让模型猜 Cypher。"
        elif any(word in str(query_concept) for word in ["是什么", "含义", "定义", "公式", "代表"]):
            results["routing_hint"] = "这是概念/定义问题，优先使用本体、知识注册表和数据契约。"
        else:
            results["routing_hint"] = "先用数据契约判断物理来源；事实数值优先 SQLite，关系证据再用 KG。"
    except Exception as exc:
        results["error"] = str(exc)
    if not results["matched_contracts"]:
        results["hint"] = "未命中具体字段，可查询 species、elevation_m、tree_dbh_cm、tree_height_m、subplot、taxon 等关键契约。"
    return json.dumps(results, ensure_ascii=False, indent=2)

# ==============================================================================
# 工具 2: 数理与神经符号规则引擎快捷绑定
# ==============================================================================
def tool_execute_neuro_symbolic_formula(knowledge_id: str, context_json: str = "{}") -> str:
    """神经-符号公式统一执行接口：参数一为公式ID (如 F_HEIGHT_DIAMETER_RATIO), 参数二为 JSON 变量字典"""
    try:
        ctx = json.loads(context_json) if isinstance(context_json, str) else context_json
        if _ns_engine is None:
            raise RuntimeError(f"NeuroSymbolicFormulaEngine unavailable: {_NS_ENGINE_INIT_ERROR}")
        res = _ns_engine.execute_formula(knowledge_id, ctx)
        return json.dumps(res, ensure_ascii=False, indent=2)
    except Exception as e:
        return json.dumps({"error": f"神经符号公式执行失败: {e}"}, ensure_ascii=False)

def tool_search_interpretive_rules(
    query_text: str = "",
    species: str = "",
    rule_category: str = "",
    target_context_json: str = "{}",
    limit: int = 5,
) -> str:
    """检索文献解释规则。规则只作为解释依据，不直接替代实测指标计算或诊断结论。"""
    return tool_retrieve_interpretive_rules(query_text, species, rule_category, target_context_json, limit)

def tool_math_calculator(expression: str) -> str:
    """安全数学计算器，仅接收纯数字的四则运算表达式"""
    try:
        if any(c not in set("0123456789+-*/(). ") for c in expression):
            return "错误：包含非法字符，只能包含数字与标准四则算术符"
        return f"计算结果: {eval(expression)}"
    except Exception as e:
        return f"计算错误: {e}"


def tool_convert_markdown_to_docx(markdown_path: str, output_path: str = None) -> str:
    """将 Markdown 文件转换为 Word (.docx)。"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return json.dumps({"error": "缺少依赖，请安装 python-docx: pip install python-docx"}, ensure_ascii=False)

    if not os.path.exists(markdown_path):
        return json.dumps({"error": f"Markdown 文件不存在: {markdown_path}"}, ensure_ascii=False)

    output_path = output_path or os.path.splitext(markdown_path)[0] + ".docx"
    doc = Document()
    with open(markdown_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    table_buffer = []
    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = [ [cell.strip() for cell in row.split("|") if cell.strip() != ""] for row in table_buffer ]
        if len(rows) >= 2:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    table.rows[i].cells[j].text = cell_text
        table_buffer = []

    for raw in lines:
        line = raw.rstrip("\n")
        if line.startswith("# "):
            flush_table()
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            flush_table()
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            flush_table()
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("![") and "](" in line and line.endswith(")"):
            flush_table()
            start = line.find("![") + 2
            end = line.find("](")
            alt = line[start:end]
            path = line[end+2:-1]
            if os.path.exists(path):
                try:
                    doc.add_paragraph(alt)
                    doc.add_picture(path, width=Inches(6))
                except Exception:
                    doc.add_paragraph(f"[图片无法插入: {path}]")
            else:
                doc.add_paragraph(f"[图片未找到: {path}]")
        elif line.strip().startswith("|") and line.strip().endswith("|"):
            table_buffer.append(line)
        elif line.strip().startswith(("- ", "* ", "+ ")):
            flush_table()
            doc.add_paragraph(line.strip(), style="List Bullet")
        elif line.strip().isdigit() and raw.strip().endswith("."):
            flush_table()
            doc.add_paragraph(line.strip(), style="List Number")
        else:
            if line.strip() == "":
                flush_table()
                doc.add_paragraph("")
            else:
                flush_table()
                doc.add_paragraph(line)
    flush_table()
    doc.save(output_path)
    return json.dumps({"success": True, "output_path": output_path}, ensure_ascii=False)


def tool_convert_markdown_to_pdf(markdown_path: str, output_path: str = None) -> str:
    """将 Markdown 文件转换为 PDF。如果可用，将先生成 docx 再转 pdf。"""
    if not os.path.exists(markdown_path):
        return json.dumps({"error": f"Markdown 文件不存在: {markdown_path}"}, ensure_ascii=False)

    output_path = output_path or os.path.splitext(markdown_path)[0] + ".pdf"
    try:
        import pypandoc
        pypandoc.convert_file(markdown_path, "pdf", outputfile=output_path)
        return json.dumps({"success": True, "output_path": output_path}, ensure_ascii=False)
    except ImportError:
        pass
    try:
        from docx2pdf import convert
        intermediate_docx = os.path.splitext(output_path)[0] + ".docx"
        convert_result = json.loads(tool_convert_markdown_to_docx(markdown_path, intermediate_docx))
        if not convert_result.get("success"):
            return json.dumps({"error": "无法先转换为 DOCX，无法生成 PDF"}, ensure_ascii=False)
        convert(intermediate_docx, output_path)
        return json.dumps({"success": True, "output_path": output_path}, ensure_ascii=False)
    except ImportError:
        return json.dumps({"error": "缺少依赖。请安装 pypandoc 或 docx2pdf，再重试。"}, ensure_ascii=False)
    except Exception as e:
        return json.dumps({"error": f"PDF 转换失败: {e}"}, ensure_ascii=False)


# ==============================================================================
# 工具 3: Cypher 图数据库提取器
# ==============================================================================
def _rewrite_kg_query(cypher: str) -> Dict[str, Any]:
    original = str(cypher or "").strip()
    rewritten = original
    notes: List[str] = []

    relation_rewrites = {
        '[:BELONGS_TO_TAXON]': '[:HAS_TAXON]',
        '[:IS_SPECIES]': '[:HAS_TAXON]',
        '[:containsTree]': '[:HAS_TREE]',
    }
    for old, new in relation_rewrites.items():
        if old in rewritten:
            rewritten = rewritten.replace(old, new)
            notes.append(f'relation:{old}->{new}')

    kg_property_rewrites = {
        ".chinese_name": ".accepted_name_cn",
        ".species_name": ".accepted_name_cn",
        ".latin_name": ".scientific_name",
    }
    for old, new in kg_property_rewrites.items():
        if old in rewritten:
            rewritten = rewritten.replace(old, new)
            notes.append(f"kg_property:{old}->{new}")

    if "TreeObservation" in rewritten and ".species" in rewritten:
        notes.append("warning:TreeObservation.species 不存在；树种应通过 TreeIndividual-[:HAS_TAXON]->Taxon.accepted_name_cn 查询，事实统计优先 SQLite 工具。")
    if "Taxon" in rewritten and any(token in rewritten for token in [".accepted_name_cn", ".scientific_name"]):
        notes.append("kg_schema:Taxon 使用 accepted_name_cn/scientific_name；事实统计仍建议优先 SQLite。")

    needs_obs = any(token in rewritten for token in [
        't.status_code',
        't.health_status',
        't.tree_dbh_cm',
        't.tree_height_m',
        't.crown_width_ns_m',
        't.crown_width_ew_m',
        't.crown_width_mean_m',
        't.crown_base_height_m',
        't.crown_length_m',
        't.tree_x_m',
        't.tree_y_m',
        't.x_coordinate',
        't.y_coordinate',
        't.x_m',
        't.y_m',
    ])
    has_obs = 'HAS_OBSERVATION' in rewritten or 'TreeObservation' in rewritten or 'obs.' in rewritten
    if needs_obs and not has_obs and 'RETURN' in rewritten:
        rewritten = rewritten.replace('RETURN', 'OPTIONAL MATCH (t)-[:HAS_OBSERVATION]->(obs:TreeObservation)\nRETURN', 1)
        notes.append('inserted_optional_observation_match')

    field_rewrites = [
        ('t.health_status', 'coalesce(obs.health_status, t.health_status)'),
        ('t.status_code', 'coalesce(obs.health_status, t.health_status)'),
        ('t.tree_dbh_cm', 'coalesce(obs.tree_dbh_cm, t.tree_dbh_cm)'),
        ('t.tree_height_m', 'coalesce(obs.tree_height_m, t.tree_height_m)'),
        ('t.crown_width_ns_m', 'coalesce(obs.crown_width_ns_m, t.crown_width_ns_m)'),
        ('t.crown_width_ew_m', 'coalesce(obs.crown_width_ew_m, t.crown_width_ew_m)'),
        ('t.crown_width_mean_m', 'coalesce(obs.crown_width_mean_m, t.crown_width_mean_m)'),
        ('t.crown_base_height_m', 'coalesce(obs.crown_base_height_m, t.crown_base_height_m)'),
        ('t.crown_length_m', 'CASE WHEN obs.tree_height_m IS NOT NULL AND obs.crown_base_height_m IS NOT NULL THEN obs.tree_height_m - obs.crown_base_height_m ELSE null END'),
        ('t.tree_x_m', 'obs.tree_x_m'),
        ('t.tree_y_m', 'obs.tree_y_m'),
        ('t.x_coordinate', 'obs.tree_x_m'),
        ('t.y_coordinate', 'obs.tree_y_m'),
        ('t.x_m', 'obs.tree_x_m'),
        ('t.y_m', 'obs.tree_y_m'),
    ]
    for token, replacement in field_rewrites:
        if token in rewritten:
            rewritten = rewritten.replace(token, f'__KG_ALIAS_{len(notes)}__')
            notes.append(f'field:{token}->{replacement}')
            rewritten = rewritten.replace(f'__KG_ALIAS_{len(notes)-1}__', replacement)

    return {
        'original': original,
        'rewritten': rewritten,
        'changed': rewritten != original,
        'notes': notes,
        'fingerprint': hashlib.md5(rewritten.encode('utf-8')).hexdigest()[:12],
    }


def tool_query_kg(cypher: str) -> str:
    """执行 Cypher 查询 Neo4j 知识图谱，并对常见旧字段与旧关系进行自动纠偏。"""
    rewrite = _rewrite_kg_query(cypher)
    try:
        with _driver.session() as session:
            records = [dict(rec) for rec in session.run(rewrite['rewritten'])]
            payload: Dict[str, Any] = {
                'status': 'success' if records else 'not_found',
                'records': records[:40],
                'record_count': len(records),
                'query_rewritten': rewrite['changed'],
                'rewrite_notes': rewrite['notes'],
            }
            if rewrite['changed']:
                payload['executed_cypher'] = rewrite['rewritten']
            if not records:
                payload['message'] = 'Cypher 查询未返回记录，请确认对象、关系和属性是否存在。'
            return json.dumps(payload, ensure_ascii=False, default=str)
    except Exception as e:
        if 'Failed to establish connection' in str(e) or 'Neo4j is not connected' in str(e) or 'WinError 10061' in str(e):
            return json.dumps({'status': 'degraded', 'message': '本地未启动 Neo4j 图数据库，系统已自动切换为纯计算模式，请优先基于本地样地实测数据回答结果。'}, ensure_ascii=False)
        return json.dumps({
            'status': 'failed',
            'error': f'Cypher 查询失败: {e}',
            'query_rewritten': rewrite['changed'],
            'rewrite_notes': rewrite['notes'],
            'executed_cypher': rewrite['rewritten'],
        }, ensure_ascii=False)


# ==============================================================================
# 工具 OpenAPI Schema 定义
# ==============================================================================
ONTOLOGY_INSPECT_SCHEMA = [{
    "type": "function",
    "function": {
        "name": "tool_inspect_ontology_schema",
        "description": "动态概念本体与公式契约检索器。当你需要查询特定的类（如 Subplot, Taxon）、字段映射（如 tree_dbh_cm）或者知识公式定义（如 HDR, Hegyi）时，自主提取对应规范切片以避免猜测属性或混淆表名！",
        "parameters": {
            "type": "object",
            "properties": {
                "query_concept": {"type": "string", "description": "要查询的领域概念、指标或公式名，如 'Subplot', 'TreeIndividual', 'HDR', 'Hegyi', 'Taxon'"}
            },
            "required": ["query_concept"]
        }
    }
},
    {
        "type": "function",
        "function": {
            "name": "tool_inspect_data_contract",
            "description": "查询语义概念与物理数据源的对齐契约。用于判断某个指标/对象应该走 SQLite、KG 还是工具，返回 SQLite 表字段、KG 标签属性、连接键和推荐工具。事实统计、排行、分布、最大最小值问题应优先查看这个契约。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query_concept": {"type": "string", "description": "要落地到物理数据源的对象或指标，如 青海云杉、海拔、species、elevation_m、样方、胸径"}
                },
                "required": ["query_concept"]
            }
        }
    }
]

KG_QUERY_SCHEMA = [{
    "type": "function",
    "function": {
        "name": "tool_query_kg",
        "description": (
            "执行 Cypher 查询以抽取 Neo4j 知识图谱信息。当前应优先按以下活跃图谱结构理解：\n"
            "- 标签：:MonitoringPlot, :Subplot, :TreeIndividual, :Taxon, :TreeObservation\n"
            "- 主关系：(Subplot)-[:HAS_TREE]->(TreeIndividual)-[:HAS_TAXON]->(Taxon)，(TreeIndividual)-[:HAS_OBSERVATION]->(TreeObservation)\n"
            "- 观测字段主要位于 TreeObservation，而不是 TreeIndividual；优先使用 tree_dbh_cm, tree_height_m, crown_width_mean_m, crown_width_ew_m, crown_width_ns_m, crown_base_height_m, health_status\n"
            "- 树种不在 TreeObservation.species；KG 中树种应通过 (TreeIndividual)-[:HAS_TAXON]->(Taxon) 和 Taxon.accepted_name_cn 获取\n"
            "- Taxon 中文名字段是 accepted_name_cn，科学名字段是 scientific_name；不要猜 chinese_name、species_name、latin_name\n"
            "- 不要优先猜测 status_code；当前健康字段应优先使用 health_status。若需要 crown_length_m，应优先由 tree_height_m - crown_base_height_m 推导。\n"
            "- 若只是要查样方结构、树种组成、单木形态、竞争、地形、气候或图表，请优先调用专用工具，不要用 Cypher 遍历大量单木。"
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "cypher": {"type": "string", "description": "符合当前图谱结构的只读 Cypher 查询"}
            },
            "required": ["cypher"]
        }
    }
}]

NS_FORMULA_SCHEMA = [
    {
        "type": "function",
        "function": {
            "name": "tool_execute_neuro_symbolic_formula",
            "description": "调用神经-符号代数求值引擎，解析并计算声明式公式。这能对特定的林学变量求值，如 F_HEIGHT_DIAMETER_RATIO (高径比)、F_QUADRATIC_MEAN_DBH (平方平均胸径)。",
            "parameters": {
                "type": "object",
                "properties": {
                    "knowledge_id": {"type": "string", "description": "公式知识 ID，例如 'F_STAND_DENSITY', 'F_QUADRATIC_MEAN_DBH', 'F_HEIGHT_DIAMETER_RATIO'"},
                    "context_json": {"type": "string", "description": "传入计算所需的变量环境参数 json"}
                },
                "required": ["knowledge_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_search_interpretive_rules",
            "description": "检索 rule.md 中的真实文献解释规则。用于回答青海云杉、气候、地形、竞争、生长等问题时提供依据；规则是启发式解释，不替代指标计算、统计检验或现场核查。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query_text": {"type": "string", "description": "用户问题或需要解释的主题，如 青海云杉小径木竞争压力"},
                    "species": {"type": "string", "description": "可选树种，如 青海云杉"},
                    "rule_category": {"type": "string", "description": "可选规则类别，如 气候规则、地形规则、树木规则"},
                    "target_context_json": {"type": "string", "description": "可选上下文 JSON，如 {\"tree_dbh_cm\":10,\"elevation_m\":2900}"},
                    "limit": {"type": "integer", "description": "返回规则数量，默认 5"}
                },
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_math_calculator",
            "description": "安全数学计算器，用来辅助智能体完成对复杂数值的复合四则运算。",
            "parameters": {
                "type": "object",
                "properties": {
                    "expression": {"type": "string", "description": "纯数字的数学表达式，如 '(4.5 + 5.6) / 2'"}
                },
                "required": ["expression"]
            }
        }
    }
]

VISUALIZATION_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "tool_create_chart",
            "description": "通用基础制图工具：从树木、样方、逐年气候、逐月气候数据中按字段灵活生成 PNG 图。支持 scatter、bar、line、box、histogram、spatial；可用 filters_json 筛选，用 group_by/aggregate 聚合。默认只生成 PNG，只有 output_format=html/both 时生成交互 HTML。",
            "parameters": {
                "type": "object",
                "properties": {
                    "chart_type": {"type": "string", "description": "图类型：scatter、bar、line、box、histogram、spatial"},
                    "data_source": {"type": "string", "description": "数据源：trees、subplots、climate_annual、climate_monthly"},
                    "x": {"type": "string", "description": "X轴字段，如 tree_dbh_cm、tree_x_m、species、year、month"},
                    "y": {"type": "string", "description": "Y轴字段，如 tree_height_m、tree_y_m、count、mean_dbh_cm、annual_precipitation_mm"},
                    "color_by": {"type": "string", "description": "颜色分组字段，如 species、subplot_id，可选"},
                    "size_by": {"type": "string", "description": "点大小字段，如 tree_dbh_cm，可选"},
                    "group_by": {"type": "string", "description": "聚合分组字段，如 species、subplot_id，可选"},
                    "aggregate": {"type": "string", "description": "聚合方式：none、count、mean、sum、min、max、median"},
                    "filters_json": {"type": "string", "description": "JSON筛选条件，如 {\"subplot_id\":\"3014\",\"species\":\"青海云杉\"}"},
                    "title": {"type": "string", "description": "图标题，可选；不填时系统会根据数据源、字段和筛选条件自动生成中文标题"},
                    "output_format": {"type": "string", "description": "输出格式：png、html、both；默认 png"}
                },
                "required": ["chart_type", "data_source"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_subplot_grid_heatmap",
            "description": "生成 600 个样方空间格网热力图，用于展示样地指标的宏观空间分布。",
            "parameters": {
                "type": "object",
                "properties": {
                    "metric": {"type": "string", "description": "指标名称，例如 total_volume_m3、volume_per_ha、density_per_ha、shannon_index、mean_hdr、high_hdr_ratio_pct"}
                },
                "required": ["metric"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_size_class_distribution",
            "description": "生成样方胸径级分布图，展示林分径级结构与蓄积贡献。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"},
                    "target_type": {"type": "string", "description": "样本类型，例如 Subplot 或 All"}
                },
                "required": ["subplot_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_species_composition",
            "description": "生成样方树种组成对比图，展示株数、断面积和蓄积三维优势度。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"},
                    "target_type": {"type": "string", "description": "样本类型，例如 Subplot 或 All"}
                },
                "required": ["subplot_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_tree_relationship_scatter",
            "description": "生成单木胸径与树高关系散点图，并可显示高径比风险边界线。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"},
                    "x_var": {"type": "string", "description": "X 轴变量，例如 tree_dbh_cm"},
                    "y_var": {"type": "string", "description": "Y 轴变量，例如 tree_height_m"}
                },
                "required": ["subplot_id", "x_var", "y_var"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_group_comparison_boxplot",
            "description": "生成分组箱线图，比较不同组别间的变量分布和异常值。",
            "parameters": {
                "type": "object",
                "properties": {
                    "variable": {"type": "string", "description": "待比较变量，例如 hdr、tree_dbh_cm 或 tree_height_m"},
                    "group_by": {"type": "string", "description": "分组字段，例如 species"}
                },
                "required": ["variable", "group_by"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_tree_spatial_map",
            "description": "绘制样方内单木空间分布图，展示树种、胸径和形态关注对象的位置。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"}
                },
                "required": ["subplot_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_subplot_percentile_profile",
            "description": "生成样方指标百分位画像图，展示目标样方在600个样方中的相对位置。",
            "parameters": {
                "type": "object",
                "properties": {
                    "subplot_id": {"type": "string", "description": "样方编号，例如 2816"}
                },
                "required": ["subplot_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_plot_climate_time_series",
            "description": "生成区域气候时序图，展示生长季降水量和温度距平变化。",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_convert_markdown_to_docx",
            "description": "将 Markdown 报告转换为 Word 文档 (.docx)。",
            "parameters": {
                "type": "object",
                "properties": {
                    "markdown_path": {"type": "string", "description": "Markdown 文件路径"},
                    "output_path": {"type": "string", "description": "可选的 DOCX 输出路径"}
                },
                "required": ["markdown_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "tool_convert_markdown_to_pdf",
            "description": "将 Markdown 报告转换为 PDF 文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "markdown_path": {"type": "string", "description": "Markdown 文件路径"},
                    "output_path": {"type": "string", "description": "可选的 PDF 输出路径"}
                },
                "required": ["markdown_path"]
            }
        }
    }
]

ALL_TOOLS = ONTOLOGY_INSPECT_SCHEMA + KG_QUERY_SCHEMA + NS_FORMULA_SCHEMA + FORESTRY_SPATIAL_SCHEMAS + VISUALIZATION_TOOLS


# ==============================================================================
# 智能体运行系统 System Prompt (祁连山国家公园天然林诊断专用)
# ==============================================================================
SYSTEM_PROMPT = """# 角色
你是祁连山国家公园森林生态智能体。你基于24公顷大样地的乔木调查数据（含坐标、胸径、树高、冠幅），回答生态查询、执行专业计算、诊断林分结构问题。

# 生态常识（你理解数据的底层逻辑）
1. 祁连山属于高寒半干旱生态区，水热条件和空间竞争是限制森林生长的核心因子。
2. 关键树种的生态性格：
   - 青海云杉：顶极群落优势种，耐阴，长寿。
   - 白桦/山杨：先锋树种，喜光，速生，短寿。
   - 祁连圆柏：顶极伴生种，极耐旱，喜阳。
   - 乌柳/花楸：湿生环境指示种，多分布于沟谷溪边。
3. 林分诊断需区分“客观观测值”（胸径/树高）与“派生诊断指标”（竞争指数/高径比）。

# 可用工具
以下工具由后台执行，你只需决定何时调用，无需关心如何计算。
1. **查询工具**：提取单木或样方的基础数据（数量、最大值、归属关系）。
2. **统计工具**：计算林分结构性指标（平均胸径、树高、密度、径级分布、蓄积量）。
3. **空间工具**：计算目标树周围邻近树木的距离和竞争关系。
4. **多样性工具**：计算 Shannon 指数、Simpson 指数、优势树种占比。
5. **诊断工具**：综合对比参照系，输出结构/更新/活力/风险四个维度的诊断信号。
6. **可视化工具**：生成径级分布柱状图、空间格局散点图、多维生态对比雷达图。
7. **文献检索工具**：当需要引用标准或学术依据时，检索本地知识网络。

# 推理流程
1. 解析用户意图：事实查询、对比分析、综合诊断还是报告生成。
2. 判断所需工具和数据维度，优先采用【并发】机制一次拉取分析所需的多维数据，严禁低效死循环。
3. 将计算结果与生态常识结合，给出深度解读（禁止机械罗列数字，必须结合树种习性推演演替动力）。
4. 若为诊断类问题，给出置信度（高/中/低）和符合国家公园零干预法则的建议。

# 输出风格
- 结论有据：注明确切的数据来源或计算依据。
- 图文结合：面对宏观报告或对比指令，自主调用可视化作图算子，并将图片 Markdown 嵌入回答对应位置实现视觉化。
- 承认边界：无法回答或无测算工具支持时明确说明，绝不推测。
"""


# ==============================================================================
# 工具执行与并发调度
# ==============================================================================
def _json_result(payload: dict) -> str:
    return json.dumps(payload, ensure_ascii=False)


def _schema_name(schema_item: dict) -> str:
    return schema_item.get("function", {}).get("name", "")


def _tool_sid(args: dict) -> str:
    return str(args.get("subplot_id", "")).strip()


def _ensure_visualization_available(tool_name: str):
    if _FORESTRY_VIS_ENGINE_AVAILABLE:
        return None
    return _visualization_engine_error_response(tool_name)


def _run_create_chart(args: dict) -> str:
    error = _ensure_visualization_available("tool_create_chart")
    if error:
        return error
    return _json_result(
        create_generic_chart(
            chart_type=args.get("chart_type", "scatter"),
            data_source=args.get("data_source", "trees"),
            x=args.get("x"),
            y=args.get("y"),
            color_by=args.get("color_by"),
            size_by=args.get("size_by"),
            group_by=args.get("group_by"),
            aggregate=args.get("aggregate", "none"),
            filters_json=args.get("filters_json", "{}"),
            title=args.get("title"),
            output_format=args.get("output_format", "png"),
        )
    )


def _run_plot_subplot_grid_heatmap(args: dict) -> str:
    error = _ensure_visualization_available("tool_plot_subplot_grid_heatmap")
    if error:
        return error
    return _json_result(plot_subplot_grid_heatmap(args.get("metric", "total_volume_m3")))


def _run_plot_size_class_distribution(args: dict) -> str:
    error = _ensure_visualization_available("tool_plot_size_class_distribution")
    if error:
        return error
    return _json_result(
        plot_size_class_distribution(
            args.get("subplot_id", _tool_sid(args)),
            args.get("target_type", "Subplot"),
        )
    )


def _run_plot_species_composition(args: dict) -> str:
    error = _ensure_visualization_available("tool_plot_species_composition")
    if error:
        return error
    return _json_result(
        plot_species_composition(
            args.get("subplot_id", _tool_sid(args)),
            args.get("target_type", "Subplot"),
        )
    )


def _run_plot_tree_relationship_scatter(args: dict) -> str:
    error = _ensure_visualization_available("tool_plot_tree_relationship_scatter")
    if error:
        return error
    return _json_result(
        plot_tree_relationship_scatter(
            args.get("subplot_id", _tool_sid(args)),
            args.get("x_var", "tree_dbh_cm"),
            args.get("y_var", "tree_height_m"),
        )
    )


def _run_plot_group_comparison_boxplot(args: dict) -> str:
    error = _ensure_visualization_available("tool_plot_group_comparison_boxplot")
    if error:
        return error
    return _json_result(
        plot_group_comparison_boxplot(
            args.get("variable", "hdr"),
            args.get("group_by", "species"),
        )
    )


def _run_plot_tree_spatial_map(args: dict) -> str:
    error = _ensure_visualization_available("tool_plot_tree_spatial_map")
    if error:
        return error
    return _json_result(plot_tree_spatial_map(args.get("subplot_id", _tool_sid(args))))


def _run_plot_subplot_percentile_profile(args: dict) -> str:
    error = _ensure_visualization_available("tool_plot_subplot_percentile_profile")
    if error:
        return error
    return _json_result(plot_subplot_percentile_profile(args.get("subplot_id", _tool_sid(args))))


def _run_plot_climate_time_series(args: dict) -> str:
    error = _ensure_visualization_available("tool_plot_climate_time_series")
    if error:
        return error
    return _json_result(plot_climate_time_series())


def _build_tool_registry() -> dict:
    schema_map = {
        _schema_name(item): item
        for item in (ONTOLOGY_INSPECT_SCHEMA + KG_QUERY_SCHEMA + NS_FORMULA_SCHEMA + FOREST_INTELLIGENCE_SCHEMAS + FORESTRY_SPATIAL_SCHEMAS + VISUALIZATION_TOOLS)
        if _schema_name(item)
    }
    return {
        "tool_inspect_ontology_schema": {
            "handler": lambda args: tool_inspect_ontology_schema(args.get("query_concept", "")),
            "enabled": True,
            "schema": schema_map["tool_inspect_ontology_schema"],
            "category": "ontology",
        },
        "tool_inspect_data_contract": {
            "handler": lambda args: tool_inspect_data_contract(args.get("query_concept", "")),
            "enabled": True,
            "schema": schema_map["tool_inspect_data_contract"],
            "category": "ontology",
        },
        "tool_query_kg": {
            "handler": lambda args: tool_query_kg(args.get("cypher", "")),
            "enabled": True,
            "schema": schema_map["tool_query_kg"],
            "category": "knowledge_graph",
        },
        "tool_execute_neuro_symbolic_formula": {
            "handler": lambda args: tool_execute_neuro_symbolic_formula(args.get("knowledge_id", ""), args.get("context_json", "{}")),
            "enabled": True,
            "schema": schema_map["tool_execute_neuro_symbolic_formula"],
            "category": "formula",
        },
        "tool_search_interpretive_rules": {
            "handler": lambda args: tool_search_interpretive_rules(
                args.get("query_text", ""),
                args.get("species", ""),
                args.get("rule_category", ""),
                args.get("target_context_json", "{}"),
                args.get("limit", 5),
            ),
            "enabled": True,
            "schema": schema_map["tool_search_interpretive_rules"],
            "category": "rules",
        },
        "tool_math_calculator": {
            "handler": lambda args: tool_math_calculator(args.get("expression", "")),
            "enabled": True,
            "schema": schema_map["tool_math_calculator"],
            "category": "utility",
        },
        "tool_resolve_forest_question": {
            "handler": lambda args: tool_resolve_forest_question(args.get("question", ""), args.get("context_json", "{}")),
            "enabled": True,
            "schema": schema_map["tool_resolve_forest_question"],
            "category": "capability",
        },
        "tool_inspect_forest_data": {
            "handler": lambda args: tool_inspect_forest_data(
                args.get("target_type", ""),
                args.get("target_name", ""),
                args.get("variables_json", "[]"),
                args.get("scope_json", "{}"),
            ),
            "enabled": True,
            "schema": schema_map["tool_inspect_forest_data"],
            "category": "capability",
        },
        "tool_run_forest_analysis_protocol": {
            "handler": lambda args: tool_run_forest_analysis_protocol(
                args.get("protocol", ""),
                args.get("target_type", ""),
                args.get("target_name", ""),
                args.get("scope_json", "{}"),
                args.get("options_json", "{}"),
            ),
            "enabled": True,
            "schema": schema_map["tool_run_forest_analysis_protocol"],
            "category": "capability",
        },
        "tool_calc_stand_structure_metrics": {
            "handler": lambda args: tool_calc_stand_structure_metrics(_tool_sid(args)),
            "enabled": False,
            "disabled_reason": "legacy metric tool; use tool_compute_registered_indicators for registered indicator calculation",
            "schema": schema_map["tool_calc_stand_structure_metrics"],
            "category": "metrics",
        },
        "tool_calc_tree_morphology_metrics": {
            "handler": lambda args: tool_calc_tree_morphology_metrics(_tool_sid(args), args.get("target_tree_id")),
            "enabled": False,
            "disabled_reason": "legacy metric tool; use tool_compute_registered_indicators for registered indicator calculation",
            "schema": schema_map["tool_calc_tree_morphology_metrics"],
            "category": "metrics",
        },
        "tool_calc_species_diversity_metrics": {
            "handler": lambda args: tool_calc_species_diversity_metrics(_tool_sid(args), args.get("survey_event_id", "EVENT_2023")),
            "enabled": False,
            "disabled_reason": "legacy metric tool; use tool_compute_registered_indicators for registered indicator calculation",
            "schema": schema_map["tool_calc_species_diversity_metrics"],
            "category": "metrics",
        },
        "tool_calc_volume_metrics": {
            "handler": lambda args: tool_calc_volume_metrics(_tool_sid(args)),
            "enabled": False,
            "disabled_reason": "legacy model tool; volume model requires validation and should not be a default chat tool",
            "schema": schema_map["tool_calc_volume_metrics"],
            "category": "metrics",
        },
        "tool_calc_deadwood_metrics": {
            "handler": lambda args: tool_calc_deadwood_metrics(_tool_sid(args)),
            "enabled": False,
            "disabled_reason": "legacy metric tool; keep for API compatibility, not default agent planning",
            "schema": schema_map["tool_calc_deadwood_metrics"],
            "category": "metrics",
        },
        "tool_calc_shrub_metrics": {
            "handler": lambda args: tool_calc_shrub_metrics(_tool_sid(args)),
            "enabled": False,
            "disabled_reason": "legacy metric tool; keep for API compatibility, not default agent planning",
            "schema": schema_map["tool_calc_shrub_metrics"],
            "category": "metrics",
        },
        "tool_calc_hegyi_competition": {
            "handler": lambda args: tool_calc_hegyi_competition(_tool_sid(args), args.get("target_tree_id", ""), args.get("radius_m", 6.0)),
            "enabled": True,
            "schema": schema_map["tool_calc_hegyi_competition"],
            "category": "metrics",
        },
        "tool_get_tree_topography_context": {
            "handler": lambda args: tool_get_tree_topography_context(args.get("tree_id", "")),
            "enabled": True,
            "schema": schema_map["tool_get_tree_topography_context"],
            "category": "environment",
        },
        "tool_get_subplot_topography_summary": {
            "handler": lambda args: tool_get_subplot_topography_summary(args.get("subplot_id", "")),
            "enabled": True,
            "schema": schema_map["tool_get_subplot_topography_summary"],
            "category": "environment",
        },
        "tool_get_climate_background_summary": {
            "handler": lambda args: tool_get_climate_background_summary(args.get("station_id", ""), args.get("date_from", ""), args.get("date_to", "")),
            "enabled": True,
            "schema": schema_map["tool_get_climate_background_summary"],
            "category": "environment",
        },
        "tool_compute_registered_indicators": {
            "handler": lambda args: tool_compute_registered_indicators(
                args.get("target_type", ""),
                args.get("target_id", ""),
                args.get("indicator_ids_json", "[]"),
                args.get("indicator_group", ""),
                args.get("parameters_json", "{}"),
            ),
            "enabled": True,
            "schema": schema_map["tool_compute_registered_indicators"],
            "category": "metrics",
        },

        "tool_create_chart": {
            "handler": _run_create_chart,
            "enabled": True,
            "schema": schema_map["tool_create_chart"],
            "category": "visualization",
        },
        "tool_plot_subplot_grid_heatmap": {
            "handler": _run_plot_subplot_grid_heatmap,
            "enabled": True,
            "schema": schema_map["tool_plot_subplot_grid_heatmap"],
            "category": "visualization",
        },
        "tool_plot_size_class_distribution": {
            "handler": _run_plot_size_class_distribution,
            "enabled": True,
            "schema": schema_map["tool_plot_size_class_distribution"],
            "category": "visualization",
        },
        "tool_plot_species_composition": {
            "handler": _run_plot_species_composition,
            "enabled": True,
            "schema": schema_map["tool_plot_species_composition"],
            "category": "visualization",
        },
        "tool_plot_tree_relationship_scatter": {
            "handler": _run_plot_tree_relationship_scatter,
            "enabled": True,
            "schema": schema_map["tool_plot_tree_relationship_scatter"],
            "category": "visualization",
        },
        "tool_plot_group_comparison_boxplot": {
            "handler": _run_plot_group_comparison_boxplot,
            "enabled": True,
            "schema": schema_map["tool_plot_group_comparison_boxplot"],
            "category": "visualization",
        },
        "tool_plot_tree_spatial_map": {
            "handler": _run_plot_tree_spatial_map,
            "enabled": True,
            "schema": schema_map["tool_plot_tree_spatial_map"],
            "category": "visualization",
        },
        "tool_plot_subplot_percentile_profile": {
            "handler": _run_plot_subplot_percentile_profile,
            "enabled": True,
            "schema": schema_map["tool_plot_subplot_percentile_profile"],
            "category": "visualization",
        },
        "tool_plot_climate_time_series": {
            "handler": _run_plot_climate_time_series,
            "enabled": True,
            "schema": schema_map["tool_plot_climate_time_series"],
            "category": "visualization",
        },
        "tool_convert_markdown_to_docx": {
            "handler": lambda args: tool_convert_markdown_to_docx(args.get("markdown_path", ""), args.get("output_path", None)),
            "enabled": True,
            "schema": schema_map["tool_convert_markdown_to_docx"],
            "category": "export",
        },
        "tool_convert_markdown_to_pdf": {
            "handler": lambda args: tool_convert_markdown_to_pdf(args.get("markdown_path", ""), args.get("output_path", None)),
            "enabled": True,
            "schema": schema_map["tool_convert_markdown_to_pdf"],
            "category": "export",
        },
    }


TOOL_REGISTRY = _build_tool_registry()
ALL_TOOLS = [meta["schema"] for meta in TOOL_REGISTRY.values() if meta.get("enabled") and meta.get("schema")]


def _contains_any(text: str, words: List[str]) -> bool:
    return any(word and word in text for word in words)


def _emit_event(event_callback: EventCallback, event_type: Any, payload: Optional[Dict[str, Any]] = None) -> None:
    if not event_callback:
        return
    if isinstance(event_type, dict):
        event_callback("stage", event_type)
    else:
        event_callback(str(event_type), payload or {})


def _emit_stage(event_callback: EventCallback, visible_process: List[Dict[str, Any]], stage: str, message: str, status: str = "running", detail: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    item: Dict[str, Any] = {"stage": stage, "status": status, "message": message}
    if detail:
        item["detail"] = detail
    visible_process.append(item)
    _emit_event(event_callback, "stage", item)
    return item


def _friendly_used_tool(tool_name: str, status: str, arguments: Optional[Dict[str, Any]] = None, detail: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    item = {"tool": tool_name, "status": status, "arguments": arguments or {}}
    if detail:
        item["detail"] = detail
    return item


def _extract_subplot_from_question(question: str) -> Optional[str]:
    q = str(question or "")
    for pattern in [r"样方\s*([0-9]{4})", r"subplot\s*([0-9]{4})", r"(?<!\d)([0-9]{4})(?!\d)"]:
        match = re.search(pattern, q, flags=re.I)
        if match:
            return match.group(1)
    return None


def _extract_tree_id_from_question(question: str) -> Optional[str]:
    match = re.search(r"([A-Za-z]{2,}\d{4,})", str(question or ""))
    return match.group(1) if match else None


def _normalize_context(context: Optional[Dict[str, Any]], question: str) -> Dict[str, Any]:
    ctx = dict(context or {})
    subplot_id = _extract_subplot_from_question(question)
    tree_id = _extract_tree_id_from_question(question)
    if subplot_id:
        ctx["explicit_subplot_id"] = subplot_id
        ctx["subplot_id"] = subplot_id
    if tree_id:
        ctx["explicit_tree_id"] = tree_id
        ctx["tree_id"] = tree_id
    return ctx


def _question_uses_ambient_context(question: str) -> bool:
    q = str(question or "")
    markers = [
        "这个样方",
        "该样方",
        "这个单木",
        "这棵树",
        "这里",
        "它",
        "该树",
        "当前",
    ]
    return _contains_any(q, markers)


_CONTEXT_OBJECT_KEYS = {
    "current_subplot_id",
    "current_tree_id",
    "subplot_id",
    "tree_id",
    "selected_subplot_ids",
    "selected_tree_ids",
}


def _sanitize_context_for_question(context: Optional[Dict[str, Any]], question: str) -> Dict[str, Any]:
    """Keep page context from silently becoming the query target."""
    ctx = dict(context or {})
    explicit_subplot_id = ctx.get("explicit_subplot_id") or _extract_subplot_from_question(question)
    explicit_tree_id = ctx.get("explicit_tree_id") or _extract_tree_id_from_question(question)
    uses_ambient = _question_uses_ambient_context(question)

    if explicit_subplot_id:
        ctx["explicit_subplot_id"] = str(explicit_subplot_id)
        ctx["subplot_id"] = str(explicit_subplot_id)
        ctx.pop("current_subplot_id", None)
        ctx.pop("selected_subplot_ids", None)
        if not explicit_tree_id:
            ctx.pop("current_tree_id", None)
            ctx.pop("tree_id", None)
            ctx.pop("selected_tree_ids", None)
    if explicit_tree_id:
        ctx["explicit_tree_id"] = str(explicit_tree_id)
        ctx["tree_id"] = str(explicit_tree_id)
        ctx.pop("current_tree_id", None)
        ctx.pop("selected_tree_ids", None)
        if not explicit_subplot_id:
            ctx.pop("current_subplot_id", None)
            ctx.pop("subplot_id", None)
            ctx.pop("selected_subplot_ids", None)

    if not uses_ambient and not explicit_subplot_id and not explicit_tree_id:
        for key in _CONTEXT_OBJECT_KEYS:
            ctx.pop(key, None)
        ctx["context_usage"] = "not_applied_no_explicit_or_deictic_target"
    elif uses_ambient:
        ctx["context_usage"] = "ambient_context_applied_by_deictic_reference"
    else:
        ctx["context_usage"] = "explicit_target_from_current_question"

    return ctx


def _should_force_direct_answer(question: str) -> bool:
    q = str(question or "").strip()
    return not q

def _route_question(question: str, context: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
    return {"route": "react", "reason": "agent_first_default"}

def _direct_answer(question: str, history: Optional[List[Dict[str, Any]]] = None, context: Optional[Dict[str, Any]] = None) -> str:
    messages = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}]
    safe_context = _sanitize_context_for_question(context, question)
    for item in history or []:
        if item.get("role") in {"user", "assistant"} and item.get("content"):
            messages.append({"role": item["role"], "content": str(item["content"])})
    if safe_context:
        messages.append({"role": "system", "content": "当前页面上下文（仅供指代解析，不要强行覆盖用户本轮问题）：\n" + json.dumps(safe_context, ensure_ascii=False, default=str)})
    messages.append({"role": "user", "content": question})
    response = chat_with_tools(messages, tools=[])
    return str((response or {}).get("content") or "").strip()


def _safe_json_loads(text: Any) -> Optional[Any]:
    if not isinstance(text, str) or not text.strip():
        return None
    try:
        return json.loads(text)
    except Exception:
        return None


def _collect_artifacts_from_payload(payload: Any, source_tool: Optional[str] = None) -> List[Dict[str, Any]]:
    artifacts: List[Dict[str, Any]] = []
    if isinstance(payload, dict):
        if isinstance(payload.get("artifacts"), list):
            for item in payload.get("artifacts") or []:
                if isinstance(item, dict):
                    enriched = dict(item)
                    if source_tool and "source_tool" not in enriched:
                        enriched["source_tool"] = source_tool
                    artifacts.append(enriched)
        for key in ["html_path", "png_path", "svg_path", "file_path", "url"]:
            value = payload.get(key)
            if value:
                artifacts.append({"type": key, "path": value, "source_tool": source_tool})
    return artifacts


def _collect_artifacts_from_tool_messages(tool_messages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    collected: List[Dict[str, Any]] = []
    for msg in tool_messages:
        payload = _safe_json_loads(msg.get("content"))
        collected.extend(_collect_artifacts_from_payload(payload, source_tool=msg.get("name")))
    dedup: List[Dict[str, Any]] = []
    seen = set()
    for item in collected:
        key = json.dumps(item, ensure_ascii=False, sort_keys=True, default=str)
        if key not in seen:
            seen.add(key)
            dedup.append(item)
    return dedup


def _update_focus(session_id: str, context: Dict[str, Any], question: str, answer: str, tool_messages: List[Dict[str, Any]]) -> Dict[str, Any]:
    focus = load_last_focus(session_id)
    subplot_id = context.get("explicit_subplot_id")
    tree_id = context.get("explicit_tree_id")
    if subplot_id:
        focus["subplot_id"] = str(subplot_id)
    if tree_id:
        focus["tree_id"] = str(tree_id)
    if context.get("current_subplot_id"):
        focus["ambient_subplot_id"] = str(context.get("current_subplot_id"))
    if context.get("current_tree_id"):
        focus["ambient_tree_id"] = str(context.get("current_tree_id"))
    focus["updated_at"] = now_iso()
    if answer:
        focus["last_answer_preview"] = answer[:120]
    if tool_messages:
        focus["last_tool_names"] = [msg.get("name") for msg in tool_messages if msg.get("name")]
    save_last_focus(session_id, focus)
    return focus

def _build_followups(context: Dict[str, Any], artifacts: List[Dict[str, Any]], used_tools: List[Dict[str, Any]]) -> List[str]:
    followups: List[str] = []
    if context.get("subplot_id") or context.get("current_subplot_id"):
        sid = context.get("subplot_id") or context.get("current_subplot_id")
        followups.extend([
            f"解释一下样方{sid}的主要关注点",
            f"给我画一下样方{sid}的空间分布图",
        ])
    if context.get("tree_id") or context.get("current_tree_id"):
        tid = context.get("tree_id") or context.get("current_tree_id")
        followups.append(f"这棵树{tid}现场还需要补记什么")
    if artifacts:
        followups.append("结合刚生成的图继续解释")
    if not followups and used_tools:
        followups.append("把这次结果整理成简短汇报")
    return followups[:4]


def _execute_tool(fn: str, args: dict) -> str:
    """执行单个已注册工具。"""
    tool_meta = TOOL_REGISTRY.get(fn)
    if not tool_meta:
        return _json_result({"error": f"未知工具: {fn}"})
    if not tool_meta.get("enabled", False):
        return _json_result({
            "error": f"工具未启用: {fn}",
            "reason": tool_meta.get("disabled_reason", "Tool disabled."),
        })
    try:
        return tool_meta["handler"](args)
    except Exception as exc:
        return _json_result({"error": f"工具执行失败: {fn}", "detail": str(exc)})
def _execute_tool_calls_parallel(tool_calls: list, event_callback: EventCallback = None) -> list:
    """并发执行工具调用，并保留 tool message 结构。"""
    results = {}

    def run_wrapper(call):
        fn = call["function"]["name"]
        try:
            args = json.loads(call["function"].get("arguments", "{}"))
        except Exception:
            args = {}
        _emit_event(event_callback, "tool_call", {"tool": fn, "arguments": args, "tool_call_id": call.get("id")})
        t_start = time.time()
        res_str = _execute_tool(fn, args)
        t_elapsed = time.time() - t_start
        return call["id"], fn, args, res_str, t_elapsed

    with ThreadPoolExecutor(max_workers=min(len(tool_calls), 8)) as executor:
        futures = [executor.submit(run_wrapper, call) for call in tool_calls]
        for future in as_completed(futures):
            cid, fn, args, res_str, elapsed = future.result()
            results[cid] = (fn, args, res_str, elapsed)

    tool_messages = []
    for call in tool_calls:
        cid = call["id"]
        fn, args, res_str, elapsed = results[cid]
        preview = res_str[:250] + "..." if len(res_str) > 250 else res_str
        print(f"    [tool] {fn}({args}) -> {elapsed:.2f}s -> {preview.strip()}")
        tool_msg = {
            "role": "tool",
            "tool_call_id": cid,
            "name": fn,
            "content": res_str,
        }
        tool_messages.append(tool_msg)
        payload = _safe_json_loads(res_str)
        status = "success"
        if isinstance(payload, dict):
            status = str(payload.get("status") or ("error" if payload.get("error") else "success"))
        _emit_event(event_callback, "tool_result", {
            "tool": fn,
            "tool_call_id": cid,
            "status": status,
            "arguments": args,
            "elapsed_s": round(elapsed, 3),
            "preview": preview,
        })
    return tool_messages


def _build_chat_messages(question: str, context: Optional[Dict[str, Any]], history: Optional[List[Dict[str, Any]]] = None) -> List[Dict[str, Any]]:
    messages: List[Dict[str, Any]] = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}]
    safe_context = _sanitize_context_for_question(context, question)
    if safe_context:
        messages.append({
            "role": "system",
            "content": "当前页面上下文（它只是辅助地图，不是硬约束；如用户本轮问题指定了新对象，以本轮问题为准）：\n" + json.dumps(safe_context, ensure_ascii=False, default=str),
        })
    if history:
        messages.append({
            "role": "system",
            "content": "历史对话只用于理解语言连续性；除非用户本轮使用明确指代（如这个、该、这里、它、当前）或明确指定同一对象，不得把历史样方/单木当作本轮查询范围。",
        })
    for item in history or []:
        role = item.get("role")
        content = str(item.get("content") or "").strip()
        if role in {"user", "assistant"} and content:
            messages.append({"role": role, "content": content})
    messages.append({"role": "user", "content": question})
    return messages


def _run_react_chat(question: str, context: Dict[str, Any], history: List[Dict[str, Any]], max_rounds: int = 12, event_callback: EventCallback = None) -> Dict[str, Any]:
    visible_process: List[Dict[str, Any]] = []
    used_tools: List[Dict[str, Any]] = []
    artifacts: List[Dict[str, Any]] = []
    tool_messages_all: List[Dict[str, Any]] = []
    messages = _build_chat_messages(question, context, history)
    total_tools_called = 0
    answer = ""

    _emit_stage(event_callback, visible_process, "理解问题", "已进入自主分析模式，开始判断是否需要组合调用工具。", "done")

    for rnd in range(1, max_rounds + 1):
        _emit_event(event_callback, "model_round", {"round": rnd, "message": f"第{rnd}轮推理"})
        stage = _emit_stage(event_callback, visible_process, f"第{rnd}轮推理", "正在分析问题并决定下一步工具调用。", "running")
        response = chat_with_tools(messages, tools=ALL_TOOLS)
        content = str((response or {}).get("content") or "").strip()
        tool_calls = (response or {}).get("tool_calls") or []
        stage["status"] = "done"
        stage["message"] = "已完成本轮思考。"
        if content:
            stage["detail"] = {"assistant_preview": content[:200]}
            _emit_event(event_callback, "stage", stage)

        if tool_calls:
            total_tools_called += len(tool_calls)
            messages.append({"role": "assistant", "content": content, "tool_calls": tool_calls})
            execute_stage = _emit_stage(event_callback, visible_process, f"第{rnd}轮执行", f"本轮准备调用 {len(tool_calls)} 个工具。", "running")
            tool_messages = _execute_tool_calls_parallel(tool_calls, event_callback=event_callback)
            tool_messages_all.extend(tool_messages)
            messages.extend(tool_messages)
            execute_stage["status"] = "done"
            execute_stage["message"] = f"已完成 {len(tool_calls)} 个工具调用。"
            _emit_event(event_callback, "stage", execute_stage)
            artifacts.extend(_collect_artifacts_from_tool_messages(tool_messages))
            for call, msg in zip(tool_calls, tool_messages):
                args = _safe_json_loads(call.get("function", {}).get("arguments", "{}")) or {}
                payload = _safe_json_loads(msg.get("content"))
                status = "success"
                if isinstance(payload, dict):
                    status = str(payload.get("status") or ("error" if payload.get("error") else "success"))
                used_tools.append(_friendly_used_tool(call.get("function", {}).get("name", "unknown_tool"), status, args))
            continue

        answer = content
        break

    if not answer:
        forced_messages = list(messages)
        forced_messages.append({
            "role": "user",
            "content": "请停止继续调用工具，基于当前已有结果直接给出中文最终回答；如果数据没找到，要明确说未找到。",
        })
        forced = chat_with_tools(forced_messages, tools=[])
        answer = str((forced or {}).get("content") or "").strip()

    return {
        "answer": answer,
        "used_tools": used_tools,
        "artifacts": artifacts,
        "visible_process": visible_process,
        "tool_call_count": total_tools_called,
        "tool_messages": tool_messages_all,
    }


def run_agent_chat(
    question: str,
    session_id: Optional[str] = None,
    client_id: Optional[str] = None,
    context: Optional[Dict[str, Any]] = None,
    options: Optional[Dict[str, Any]] = None,
    event_callback: EventCallback = None,
) -> Dict[str, Any]:
    question = str(question or "").strip()
    session_id = session_id or create_session_id()
    ensure_session(session_id, client_id=client_id)
    options = dict(options or {})
    context = _normalize_context(context, question)
    history = load_recent_messages(session_id, limit=int(options.get("history_limit", 12)))
    route_info = _route_question(question, context)
    visible_process: List[Dict[str, Any]] = []
    _emit_event(event_callback, "session", {"session_id": session_id})

    if not question:
        answer = "请先输入问题。"
        result = {
            "session_id": session_id,
        "client_id": client_id,
            "answer_type": "chat_answer",
            "answer": answer,
            "used_tools": [],
            "artifacts": [],
            "followups": [],
            "last_focus": load_last_focus(session_id),
            "visible_process": visible_process,
            "evidence_summary": [],
            "debug_trace": [],
            "semantic_plan": None,
            "candidate_preview": [],
        }
        _emit_event(event_callback, "final", result)
        return result

    if _should_force_direct_answer(question):
        _emit_stage(event_callback, visible_process, "直接回答", "问题无需调用工具，直接回答", "done")
        answer = _direct_answer(question, history=history, context=context)
        used_tools: List[Dict[str, Any]] = [_friendly_used_tool("llm_direct_answer", "success", {"mode": "trivial_direct"})]
        artifacts: List[Dict[str, Any]] = []
        tool_messages: List[Dict[str, Any]] = []
    else:
        react_result = _run_react_chat(
            question=question,
            context=context,
            history=history,
            max_rounds=int(options.get("max_tool_rounds", options.get("max_rounds", 12))),
            event_callback=event_callback,
        )
        answer = react_result["answer"]
        used_tools = react_result["used_tools"]
        artifacts = react_result["artifacts"]
        visible_process = react_result["visible_process"]
        tool_messages = react_result.get("tool_messages", [])

    focus = _update_focus(session_id, context, question, answer, tool_messages)
    followups = _build_followups(context, artifacts, used_tools)
    evidence_summary = [
        {"type": "autonomy_mode", "value": "agent_first"},
        {"type": "history_used", "value": len(history)},
        {"type": "tool_count", "value": len(used_tools)},
    ]

    save_message(session_id, "user", question, context=context)
    save_message(
        session_id,
        "assistant",
        answer,
        context={"mode": "agent_first", "last_focus": focus},
        answer_type="chat_answer",
        tool_calls=used_tools,
        artifacts=artifacts,
    )

    result = {
        "session_id": session_id,
        "answer_type": "report" if _contains_any(question.lower(), ["报告", "汇报", "pdf", "word"]) else "chat_answer",
        "answer": answer,
        "used_tools": used_tools,
        "artifacts": artifacts,
        "followups": followups,
        "last_focus": focus,
        "visible_process": visible_process,
        "evidence_summary": evidence_summary,
        "debug_trace": [],
        "semantic_plan": None,
        "candidate_preview": [],
    }
    _emit_event(event_callback, "final", result)
    return result


def run_agent_report(
    question: str,
    session_id: Optional[str] = None,
    client_id: Optional[str] = None,
    context: Optional[Dict[str, Any]] = None,
    options: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    q = str(question or "").strip()
    if not _contains_any(q.lower(), ["报告", "汇报", "文档", "导出"]):
        q = "请生成一份结构化中文报告：" + q
    result = run_agent_chat(q, session_id=session_id, client_id=client_id, context=context, options=options)
    result["answer_type"] = "report"
    return result


def run_agent(question: str, report_file: str = "report.txt", max_rounds: int = 15):
    """
    启动 ReAct 主循环。
    LLM 将根据 System Prompt 决定工具组合，单轮能同时运行 N 个工具，大幅提升诊断自适应能力。
    """
    print(f"\n=================================================================")
    print(f" [Forestry Agent] 正在启动祁连山天然林智能诊断引擎 (ReAct Mode)")
    print(f" 问题: {question}")
    print(f"=================================================================\n")

    messages = [
        {"role": "system", "content": CHAT_SYSTEM_PROMPT},
        {"role": "user", "content": question}
    ]

    total_tools_called = 0
    t_engine_start = time.time()

    for rnd in range(1, max_rounds + 1):
        print(f"  ── [第 {rnd} 轮] 诊断思考中 ──")
        t_call_start = time.time()
        
        response = chat_with_tools(messages, tools=ALL_TOOLS)
        llm_elapsed = time.time() - t_call_start

        content = response.get("content") or ""
        tool_calls = response.get("tool_calls") or []

        if content:
            print(f"  🧠 AI 思考阐释 ({llm_elapsed:.1f}s): {content[:300]}{'...' if len(content) > 300 else ''}")

        if tool_calls:
            n_tools = len(tool_calls)
            total_tools_called += n_tools
            print(f"  🔧 发射 {n_tools} 个并行工具调用:")
            
            # 追加上一轮的 Assistant 应答
            messages.append({
                "role": "assistant",
                "content": content,
                "tool_calls": tool_calls
            })

            t_exec = time.time()
            tool_msgs = _execute_tool_calls_parallel(tool_calls)
            exec_elapsed = time.time() - t_exec
            
            messages.extend(tool_msgs)
            print(f"  ✔ 本轮 {n_tools} 个并行算子返回成功，执行历时 {exec_elapsed:.2f}s\n")
        else:
            # 没有 tool_calls，代表推理链条收敛，生成了最终分析报告
            total_elapsed = time.time() - t_engine_start
            print(f"=================================================================")
            print(f" ✅ 智能体分析报告收敛。历时: {total_elapsed:.1f}s | 总计调用工具: {total_tools_called} 次")
            print(f"=================================================================\n")

            with open(report_file, "w", encoding="utf-8") as f:
                f.write(content)
            print(f" [OK] 报告成功存入本地文件: `{report_file}`\n")
            return content

    # 循环溢出恢复机制
    print(f"\n [Warning] 已突破上限最大轮次 ({max_rounds})，启动强制收敛收拢报告。")
    messages.append({
        "role": "user",
        "content": "已达最大分析时钟。请立即用已有数据停止决策工具，直接输出完整的排队对比诊断报告 markdown 文本。"
    })
    forced_resp = chat_with_tools(messages, tools=[])
    final_report = forced_resp.get("content", "（因外部接口调用强制收尾失败）")
    with open(report_file, "w", encoding="utf-8") as f:
        f.write(final_report)
    print(f" [OK] 溢出阶段强制生成报告并存入: `{report_file}`")
    return final_report


if __name__ == "__main__":
    if len(sys.argv) > 1:
        user_q = " ".join(sys.argv[1:])
    else:
        user_q = (
            "请全面分析3014样方的森林质量，我需要去对其中风险最高的5棵树进行复查，请给我一个包含图片和表格的报告，用于我有针对性地去野外调查，我手上只有相机，需要告诉我拍哪些照片，比如树木本身的，周围环境的？要尽可能详细，让我全方位了解这棵树。"
        )
    run_agent(
        question=user_q,
        report_file="Forest_Quality_7.md",
        max_rounds=40
    )
    _driver.close()
